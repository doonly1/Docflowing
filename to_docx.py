# -*- coding: utf-8 -*-
"""
文档转Word工具
支持 PDF、DOC、DOCX、TXT、HTML、MD 等格式提取文本，生成公文格式DOCX
"""

import os
import re
from logging_config import setup_logging, get_logger

setup_logging()
logger = get_logger(__name__)


def merge_lines_to_paragraphs(lines, char_positions=None):
    """将 PDF 提取的逐行文本合并为段落

    策略：
    1. 以缩进或结构性标记开头的行 → 新段落
    2. 已累积段落以句末标点结尾 → 在当前行前断段
    3. 其余情况 → 与上一段合并（行间不加空格，中文排版无空格需求）

    Args:
        lines: 逐行文本列表
        char_positions: 可选，每行首字符的 x 坐标列表（由 pdfplumber 提供），
                        用于精确检测缩进，优于正则匹配
    """
    if not lines:
        return lines

    # 句末标点：出现在段落末尾时，下一行应为新段落
    SENTENCE_ENDS = set('。！？…」）】』》"\u300b：')

    # 新段落起始模式（缩进由 char_positions 单独判断，此处只匹配结构标记）
    NEW_PARA_START = re.compile(
        r'^('
        r'[（(][一二三四五六七八九十百]+[)）]'     # （一）（二）
        r'|[一二三四五六七八九十]+[、]'              # 一、二、
        r'|第[一二三四五六七八九十百千]+[条章节款项]'  # 第一条
        r'|附\s*[件录则]'                          # 附件/附录/附则
        r')'
    )

    # 计算缩进阈值：首字符 x 坐标明显大于大多数行 → 视为缩进行
    indent_threshold = None
    if char_positions and len(char_positions) == len(lines):
        # 取所有 x 坐标，用众数（最小值群）作为正文左边界
        valid_x = [x for x in char_positions if x is not None and x >= 0]
        if valid_x:
            baseline_x = min(valid_x)  # 正文左边界
            indent_threshold = baseline_x + 15  # 超出左边界15pt视为缩进

    result = []
    current = lines[0]

    for i in range(1, len(lines)):
        line = lines[i]
        current_last = current.rstrip()[-1] if current.rstrip() else ''

        # 规则1a：当前行有缩进（基于字符坐标）→ 新段落
        if indent_threshold is not None:
            line_x = char_positions[i] if i < len(char_positions) else None
            if line_x is not None and line_x > indent_threshold:
                result.append(current)
                current = line
                continue

        # 规则1b：当前行以结构标记开头 → 新段落
        if NEW_PARA_START.match(line):
            result.append(current)
            current = line
            continue

        # 规则2：已累积段落以句末标点结尾 → 断段
        if current_last in SENTENCE_ENDS:
            result.append(current)
            current = line
            continue

        # 默认：合并到当前段落
        current += line

    if current:
        result.append(current)

    return result


def extract_text_from_pdf(file_path):
    """从PDF提取文本，并合并行成段落"""
    import pdfplumber

    all_lines = []
    all_char_x = []  # 每行首字符的 x 坐标（用于缩进检测）

    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if not page_text:
                continue

            # 获取字符级坐标用于缩进检测
            # 按 y 坐标分组，每组代表一个可视行，取每行最左侧 x 坐标
            line_x_map = {}  # 行序号 → 首字符 x 坐标
            chars = page.chars if hasattr(page, 'chars') else []
            if chars:
                from collections import defaultdict
                row_chars = defaultdict(list)
                for ch in chars:
                    if not ch.get('text', '').strip():
                        continue
                    # y 坐标四舍五入，容忍轻微偏移
                    y = round(ch['top'], 0)
                    row_chars[y].append(ch)

                # 按 y 排序（从上到下），为每行分配序号
                sorted_ys = sorted(row_chars.keys())
                row_x_list = []
                for y in sorted_ys:
                    # 该行最左侧字符的 x 坐标
                    min_x = min(c['x0'] for c in row_chars[y])
                    row_x_list.append(min_x)

                # 将 page_text 的行与 chars 行按序对齐
                text_lines = page_text.split('\n')
                ti = 0  # text_lines 索引
                ri = 0  # row_chars 索引
                while ti < len(text_lines) and ri < len(sorted_ys):
                    text_line = text_lines[ti].strip()
                    if not text_line:
                        ti += 1
                        continue
                    # 从 chars 构建该行的文本来比对
                    row_text = ''.join(
                        c['text'] for c in sorted(row_chars[sorted_ys[ri]],
                                                   key=lambda c: c['x0'])
                    ).strip()
                    if row_text and (text_line in row_text or row_text in text_line):
                        line_x_map[ti] = row_x_list[ri]
                        ti += 1
                        ri += 1
                    elif len(row_text) < len(text_line):
                        # chars 行较短，可能跨行，跳过 chars 行
                        ri += 1
                    else:
                        # text 行较短（可能是过滤行），跳过 text 行
                        ti += 1

            for ti, line in enumerate(page_text.split('\n')):
                line = line.strip()
                if not line:
                    continue
                # 过滤页码
                if re.match(r'^-?\d+-?$', line):
                    continue
                if re.match(r'^第\s*\d+\s*页$', line):
                    continue
                if line.isdigit() and len(line) <= 3:
                    continue
                if '版权所有' in line or '翻印必究' in line:
                    continue

                all_lines.append(line)
                all_char_x.append(line_x_map.get(ti))

    # 后处理：合并行成段落
    paragraphs = merge_lines_to_paragraphs(all_lines, all_char_x)
    return '\n'.join(paragraphs).strip()


def extract_text_from_docx(file_path):
    """从DOCX提取文本"""
    from docx import Document
    doc = Document(file_path)
    return '\n'.join(para.text for para in doc.paragraphs if para.text.strip())


def extract_text_from_doc(file_path):
    """从DOC提取文本（转为DOCX后提取）"""
    from doc_process import doc_to_docx
    workdir = os.path.dirname(file_path)
    # 先转换为docx
    doc_to_docx(workdir)
    # 重新查找转换后的docx文件
    basename = os.path.splitext(os.path.basename(file_path))[0]
    docx_path = os.path.join(workdir, basename + '.docx')
    if os.path.exists(docx_path):
        return extract_text_from_docx(docx_path)
    return ""


def extract_text_from_txt(file_path):
    """从TXT提取文本"""
    encodings = ['utf-8', 'gbk', 'gb2312']
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                return f.read().strip()
        except UnicodeDecodeError:
            continue
    return ""


def extract_text_from_html(file_path):
    """从HTML提取文本"""
    from bs4 import BeautifulSoup
    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
        soup = BeautifulSoup(f.read(), 'html.parser')
    # 移除script和style标签
    for tag in soup(['script', 'style']):
        tag.decompose()
    text = soup.get_text()
    # 清理多余空白
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    return '\n'.join(lines)


def extract_text_from_md(file_path):
    """从Markdown提取文本（去除标记）"""
    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
        content = f.read()
    # 移除代码块
    content = re.sub(r'```[\s\S]*?```', '', content)
    # 移除行内代码
    content = re.sub(r'`[^`]+`', '', content)
    # 移除图片
    content = re.sub(r'!\[.*?\]\(.*?\)', '', content)
    # 移除链接，保留文字
    content = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', content)
    # 移除标题标记
    content = re.sub(r'^#{1,6}\s+', '', content, flags=re.MULTILINE)
    # 移除粗体斜体标记
    content = re.sub(r'[*_]{1,3}([^*_]+)[*_]{1,3}', r'\1', content)
    # 移除引用标记
    content = re.sub(r'^>\s*', '', content, flags=re.MULTILINE)
    # 移除列表标记
    content = re.sub(r'^[\s]*[-*+]\s+', '', content, flags=re.MULTILINE)
    content = re.sub(r'^[\s]*\d+\.\s+', '', content, flags=re.MULTILINE)
    # 清理多余空白
    lines = [line.strip() for line in content.splitlines() if line.strip()]
    return '\n'.join(lines)


def extract_text(file_path):
    """根据文件扩展名提取文本"""
    ext = os.path.splitext(file_path)[1].lower()
    filename = os.path.basename(file_path)
    
    # 跳过临时文件
    if filename.startswith('~$'):
        return None
    
    extractors = {
        '.pdf': extract_text_from_pdf,
        '.docx': extract_text_from_docx,
        '.doc': extract_text_from_doc,
        '.txt': extract_text_from_txt,
        '.html': extract_text_from_html,
        '.htm': extract_text_from_html,
        '.md': extract_text_from_md,
    }
    
    if ext not in extractors:
        return None
    
    try:
        logger.info("正在提取: %s", filename)
        return extractors[ext](file_path)
    except Exception as e:
        logger.error("  提取失败: %s", e)
        return None


def generate_docx(file_path):
    """将单个文件转换为公文格式DOCX

    Args:
        file_path: 源文件路径，支持 PDF/DOC/DOCX/TXT/HTML/MD
    """
    text = extract_text(file_path)
    if not text:
        return False

    from docx import Document
    from doc_process import set_appendix, set_date, save_docx, set_headings
    from mystyle import clear_styles, add_my_styles, set_page

    workdir = os.path.dirname(file_path)
    base_name = os.path.splitext(os.path.basename(file_path))[0]

    doc = Document()
    set_page(doc)
    clear_styles(doc)
    add_my_styles(doc)

    for line in text.splitlines():
        doc.add_paragraph(line)

    set_headings(doc)
    set_appendix(doc)
    set_date(doc)
    save_docx(doc, f"{base_name}.docx", workdir)
    return True


def convert_folder(workdir):
    """转换文件夹中的所有支持的文档"""
    supported_ext = ('.pdf', '.docx', '.doc', '.txt', '.html', '.htm', '.md')
    
    files = [f for f in os.listdir(workdir) 
             if os.path.splitext(f)[1].lower() in supported_ext]
    
    # 跳过带4位数下划线前缀的docx文件（避免重复处理）
    original_count = len(files)
    files = [f for f in files if not (
        os.path.splitext(f)[1].lower() == '.docx' and 
        re.match(r'^\d{4}_', f)
    )]
    skipped = original_count - len(files)
    if skipped > 0:
        logger.info("已跳过 %s 可能处理过的docx文件", skipped)
    
    if not files:
        logger.warning("未找到支持的文档文件 (pdf/doc/docx/txt/html/md)")
        return
    
    success_count = sum(generate_docx(os.path.join(workdir, f)) for f in files)
    
    logger.info("%s", '=' * 50)
    logger.info("转换完成: 成功 %s, 失败 %s", success_count, len(files) - success_count)


if __name__ == '__main__':
    import sys
    paths = sys.argv[1:] if len(sys.argv) > 1 else [os.path.dirname(__file__)]
    for path in paths:
        if os.path.isfile(path):
            generate_docx(path)
        elif os.path.isdir(path):
            convert_folder(path)
