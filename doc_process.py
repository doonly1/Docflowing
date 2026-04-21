from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
import platform
import subprocess

from mystyle import para_fm, run_fm


def find_libreoffice():
    """查找 LibreOffice 可执行文件路径，未找到返回 None

    Returns:
        str or None: soffice 可执行文件路径
    """
    candidates = []
    system = platform.system()
    if system == 'Windows':
        candidates = [
            os.path.join(os.environ.get('PROGRAMFILES', r'C:\Program Files'),
                         'LibreOffice', 'program', 'soffice.exe'),
            os.path.join(os.environ.get('PROGRAMFILES(X86)', r'C:\Program Files (x86)'),
                         'LibreOffice', 'program', 'soffice.exe'),
        ]
        # 尝试从注册表查找
        try:
            import winreg
            key = winreg.OpenKey(winreg.HKEY_LOCAL_MACHINE,
                                 r'SOFTWARE\LibreOffice\LibreOffice', 0,
                                 winreg.KEY_READ | winreg.KEY_WOW64_64KEY)
            install_dir, _ = winreg.QueryValueEx(key, 'InstallDir')
            candidates.insert(0, os.path.join(install_dir, 'program', 'soffice.exe'))
        except Exception:
            pass
    elif system == 'Darwin':
        candidates = ['/Applications/LibreOffice.app/Contents/MacOS/soffice']
    else:  # Linux
        candidates = ['libreoffice', 'soffice',
                      '/usr/bin/libreoffice', '/usr/bin/soffice',
                      '/snap/bin/libreoffice']
    for cmd in candidates:
        try:
            if os.path.isabs(cmd):
                if os.path.isfile(cmd):
                    return cmd
            else:
                result = subprocess.run(['which', cmd], capture_output=True, text=True)
                if result.returncode == 0:
                    return cmd
        except Exception:
            continue
    return None


def libreoffice_install_hint():
    """返回当前平台的 LibreOffice 安装提示"""
    system = platform.system()
    if system == 'Windows':
        return '请安装 LibreOffice: https://www.libreoffice.org/download/'
    elif system == 'Darwin':
        return '请安装 LibreOffice: brew install --cask libreoffice'
    else:
        return '请安装 LibreOffice: sudo apt install libreoffice  (或 yum/dnf/pacman 等对应命令)'


def doc_to_docx(workdir):
    """将workdir目录下的.doc文件批量转换为.docx格式（跨平台）"""
    doc_files = [f for f in os.listdir(workdir)
                 if f.lower().endswith('.doc') and not f.startswith("~$")]
    if not doc_files:
        return

    # 优先尝试 win32com（Windows + Word 已安装时最快）
    try:
        from win32com import client
        word = client.Dispatch("Word.Application")
        word.Visible = False
        for file in doc_files:
            print('转化docx：{}'.format(file))
            file_path = os.path.join(workdir, file)
            doc = word.Documents.Open(file_path)
            doc.SaveAs("{}x".format(file_path), 12)
            doc.Close()
            try:
                os.remove(file_path)
            except:
                pass
        word.Quit()
        return
    except Exception:
        pass  # win32com 不可用，回退到 LibreOffice

    # 尝试 LibreOffice
    lo_cmd = find_libreoffice()
    if lo_cmd:
        for file in doc_files:
            print('转化docx(LibreOffice)：{}'.format(file))
        cmd = [lo_cmd, '--headless', '--convert-to', 'docx',
               '--outdir', workdir] + [os.path.join(workdir, f) for f in doc_files]
        try:
            subprocess.run(cmd, check=True, capture_output=True)
            # 清理原 .doc 文件
            for file in doc_files:
                try:
                    os.remove(os.path.join(workdir, file))
                except:
                    pass
        except subprocess.CalledProcessError as e:
            print(f'LibreOffice 转换失败: {e}')
        return

    # 都不可用
    print('警告：无法转换 .doc 文件 — 需要 Microsoft Word 或 LibreOffice')
    print(libreoffice_install_hint())


def save_docx(doc, doc_name, workdir=None):
    """
    保存Word文档
    
    Args:
        doc: Document对象
        doc_name: 文件名（可包含路径）
        workdir: 可选的保存目录，默认为None（当前目录）
    
    Returns:
        保存的完整文件路径，失败返回None
    """
    import re
    import zlib
    
    # 分离目录和文件名
    if workdir is None:
        # 如果doc_name包含路径，分离出来
        if os.path.dirname(doc_name):
            workdir = os.path.dirname(doc_name)
            doc_name = os.path.basename(doc_name)
        else:
            workdir = os.getcwd()
    
    # 清理文件名中的非法字符
    doc_name = re.sub(r'[\\/:*?"<>|]', '_', doc_name)
    
    # 确保扩展名正确
    if not doc_name.lower().endswith('.docx'):
        doc_name += '.docx'
    
    # 生成目标文件名哈希前缀（4位纯数字）
    hash_prefix = str(zlib.crc32(doc_name.encode()) % 10000).zfill(4)
    
    # 检查目标文件是否已存在（已另存过）
    save_path = os.path.normpath(os.path.join(workdir, hash_prefix + "_" + doc_name))
    if os.path.exists(save_path):
        return save_path
    
    try:
        doc.save(save_path)
        print(f"已另存: {save_path}")
        return save_path
    except Exception as e:
        print(f"另存失败: {e}")
        return None


def set_headings(doc):
    paras=doc.paragraphs
    title = ['决议','决定','命令','公报','公告','通告','意见',\
                '通知','通报','报告','请示','批复','议案','的函','纪要',\
                '计划','总结','申请','名单','制度','办法','规定','方案','要点']
    heading1=['一、','二、','三、','四、','五、','六、',\
              '七、','八、','九、','十、']
    heading2=['（一）','（二）','（三）','（四）','（五）',\
              '（六）','（七）','（八）','（九）','（十）']
    for para in paras:                          
        if paras.index(para)<3 and len(para.text)<60 and para.text.strip(" ")[-2:] in title:  #标题识别
            para.style = doc.styles['Title']
            para_fm(para,0,0,28.95,0,0,0,'C')
            for run in para.runs:
                run_fm(run,'方正小标宋简体',22,0,0,0,0)

        if 1<len(para.text)<30 and '：' in para.text[-1] \
           and len(paras[paras.index(para)-1].text)==0:  #主送识别
            para.style = doc.styles['unindent']
            para_fm(para,0,0,28.95,0,0,0,'J')
            for run in para.runs:
                run_fm(run,'仿宋',16,0,0,0,0)

        if para.text.strip(" ")[:2] in heading1 and len(para.text)<60 and '：' not in para.text:  #一级标题识别
            para.style = doc.styles['H1']
            for run in para.runs:
                run_fm(run,'黑体',16,0,0,0,0)

        if para.text.strip(" ")[:3] in heading2 and len(para.text)<60 and '：' not in para.text:   #二级标题识别
            para.style = doc.styles['H2']
            for run in para.runs:
                run_fm(run,'楷体',16,0,0,0,0)


def set_appendix(doc):
    paras=doc.paragraphs
    #2段落遍历完毕后，再次遍历到“附件”段落。设置附件格式，并获得其个数和内容。
    dx_s=[]      #进入循环前，确定卡点变量不被重复覆盖。
    for para in paras:
        dx='附件：'
        if dx in para.text[:5] and len(para.text)>3:
            dx_s.append(paras.index(para))	 #附件段落序列数

    n = 0  # 初始化n
    if len(dx_s) !=0:   #如果有附件
        dx_n_strs=[]
        for n in range(len(paras[dx_s[0]+1:])):		#往后每一段，遍历查找附件
            if str(n+1)+'.' in paras[dx_s[0]+n].text[:6]:  #判断n.在不在前6个字符内
                dx_n_str=paras[dx_s[0]+n].text[2:] #获得单个附件字符串
                for p in ['：','.','。']:	  #去掉标点
                    if p in dx_n_str:
                        dx_n_str=dx_n_str.replace(p,'')
                dx_n_strs.append(dx_n_str)	#获得所有附件的字符串

        if len(dx_n_strs)==0:
            dx_n_strs.append(paras[dx_s[0]].text[3:]) #只有一个附件时
            paras[dx_s[0]].style=doc.styles['Apdix']
            para_fm(paras[dx_s[0]],0,0,28.95,80,0,-48,'L')   #段落格式
			
        n=len(dx_n_strs)  #有n个附件
        # print('有{}个附件'.format(n))
    
    if len(dx_s) != 0 and n>0:    #如果有附件
        for i in range(n):
            for j in range(n):		  #遍历所有附件
                if str(i+1)+'.' in paras[dx_s[0]+j].text:
                    paras[dx_s[0]+j].style=doc.styles['Apdix 2']
                    para_f=paras[dx_s[0]+j].paragraph_format   #段落格式赋给para_f
                    para_f.alignment=WD_PARAGRAPH_ALIGNMENT.LEFT  #对齐方式
                    para_f.first_line_indent = Pt(0) 
                    para_f.left_indent = Pt(16*6)   #左缩进（Inches,Cm，Pt）需弥补悬挂负值
                    para_f.first_line_indent = Pt(-16)  #悬挂1个字符
        paras[dx_s[0]].style=doc.styles['Apdix 1']     #重设附件1.的格式
        para_fm(paras[dx_s[0]],0,0,28.95,16*6,0,0,'L')
        para_fm(paras[dx_s[0]],0,0,28.95,16*6,0,-16*4,'L')  #重设附件1.的格式

    #4设置顶格附件格式
    if len(dx_s) != 0:
        for para in paras[dx_s[0]:len(paras)-1]:
            if '附' in para.text and '件' in para.text and len(para.text)<5:   #找到顶格'附件'
                # print('第{}段有|附件|：{}'.format(paras.index(para)+1,para.text.strip('\n')))
                if '：' in para.text[-1]:
                    para.text=para.text[:-1]
                para.style=doc.styles['Blackbody']	#设置顶格附件样式
                para_fm(para,0,0,28.95,0,0,0,'L')
                for run in para.runs:
                    run_fm(run,'黑体',16,0,0,0,0)	#设置顶格附件格式
                                      

def set_date(doc):
    #3署名及日期设置 
    paras=doc.paragraphs
    for para in paras:
        if '年' in para.text and '月' in para.text\
           and '日' in para.text and len(para.text)<12:
            # print('第{}段有日期：{}'.format(paras.index(para)+1,para.text))
            para.style=doc.styles['dater']	#日期格式
            para_fm(para,0,0,28.95,0,16*4,0,'R')	#日期格式
            
            paras[paras.index(para)-1].style = doc.styles['Sign']   #日期上一段的样式
            para_fm_bef=paras[paras.index(para)-1].paragraph_format   #日期上一段的格式
            para_fm_bef.alignment=WD_PARAGRAPH_ALIGNMENT.RIGHT
            para_fm_bef.first_line_indent = Pt(0)
            b=len(paras[paras.index(para)-1].text)  # b是日期上方署名的字符数
            if len(para.text) in {9,10}:		  #判断日期的字符数
                para_fm_bef.right_indent = Pt((8-0.5*b)*16)
            elif len(para.text)==11:
                para_fm_bef.right_indent = Pt((8.5-0.5*b)*16)
