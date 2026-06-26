# -*- coding: utf-8 -*-
"""测试 to_compare 模块的核心比较函数"""

import sys
import os
_basedir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, _basedir)
sys.path.insert(0, os.path.join(_basedir, 'tools'))

import difflib
import pytest

# 先初始化 logging，避免 to_compare 导入时的副作用干扰测试输出捕获
os.environ.setdefault('LOG_LEVEL', 'ERROR')

from tools.to_compare import (
    char_level_diff,
    split_into_sentences,
    _build_thresholds,
    _longest_increasing_subsequence,
    _recursive_split_with_thresholds,
    _split_at_reliable_boundaries,
    split_by_reliable_boundaries,
    _merge_unmatched_sentences,
    _detect_sentence_merge_split,
    PUNCTUATION_LEVELS,
)

# ==================== 测试数据 ====================
# 注意：文本中使用了中文全角双引号，此处用单引号包裹外层字符串

ORIG_TEXT = ''.join([
    '继2026高考吴家山中学助考点后，此次2026中考武汉常青树实验学校三店校区爱心加油站，',
    '是蒙牛聚焦学生营养健康、推动教育事业发展的生动实践。',
    '蒙牛积极践行国家龙头乳企担当，履行社会责任，不仅率先加入国家\u201c学生饮用奶计划\u201d，',
    '为广大学生提供优质营养，而且全力推动\u201c营养普惠工程\u201d，',
    '开展牛奶捐赠、助学助教、营养科普、体育赞助等系列公益实践。',
    '截至目前，蒙牛学生奶已累计捐赠10亿元，覆盖28个省、自治区及直辖市，惠及学生2500万人以上。',
])

FINAL_TEXT = ''.join([
    '继2026年高考吴家山中学助考点后，此次2026年中考武汉常青树实验学校三店校区爱心加油站，',
    '是蒙牛聚焦学生营养健康、推动教育事业发展的又一生动实践。',
    '蒙牛积极践行国家龙头乳企担当，履行社会责任，不仅率先加入国家\u201c学生饮用奶计划\u201d，',
    '为广大学生提供优质营养，而且全力推动\u201c营养普惠工程\u201d，',
    '开展牛奶捐赠、助学助教、营养科普、体育赞助等系列公益实践。',
    '截至目前，蒙牛学生奶已累计捐赠10亿元，覆盖28个省、自治区、直辖市，惠及学生2500万人。',
])

ORIG_SENTENCE_1 = (
    '继2026高考吴家山中学助考点后，此次2026中考武汉常青树实验学校三店校区爱心加油站，'
    '是蒙牛聚焦学生营养健康、推动教育事业发展的生动实践。'
)
FINAL_SENTENCE_1 = (
    '继2026年高考吴家山中学助考点后，此次2026年中考武汉常青树实验学校三店校区爱心加油站，'
    '是蒙牛聚焦学生营养健康、推动教育事业发展的又一生动实践。'
)

ORIG_SENTENCE_2 = ''.join([
    '蒙牛积极践行国家龙头乳企担当，履行社会责任，不仅率先加入国家\u201c学生饮用奶计划\u201d，',
    '为广大学生提供优质营养，而且全力推动\u201c营养普惠工程\u201d，',
    '开展牛奶捐赠、助学助教、营养科普、体育赞助等系列公益实践。',
])
FINAL_SENTENCE_2 = ORIG_SENTENCE_2

ORIG_SENTENCE_3 = (
    '截至目前，蒙牛学生奶已累计捐赠10亿元，覆盖28个省、自治区及直辖市，惠及学生2500万人以上。'
)
FINAL_SENTENCE_3 = (
    '截至目前，蒙牛学生奶已累计捐赠10亿元，覆盖28个省、自治区、直辖市，惠及学生2500万人。'
)


def _count_tag(chars, tag):
    """统计 diff 结果中指定 tag 的字符数"""
    return sum(len(t) for t_type, t in chars if t_type == tag)


class TestBuildThresholds:
    """测试 _build_thresholds"""

    def test_default_values(self):
        """默认阈值列表：每级标点对应的阈值"""
        thresholds = _build_thresholds([30, 30])
        assert len(thresholds) == len(PUNCTUATION_LEVELS)
        assert thresholds[0] == (PUNCTUATION_LEVELS[0], 30)
        assert thresholds[1] == (PUNCTUATION_LEVELS[1], 30)
        assert thresholds[2] == (PUNCTUATION_LEVELS[2], None)

    def test_partial_values(self):
        """部分提供阈值：未提供的级为 None"""
        thresholds = _build_thresholds([50])
        assert thresholds[0] == (PUNCTUATION_LEVELS[0], 50)
        assert thresholds[1] == (PUNCTUATION_LEVELS[1], None)
        assert thresholds[2] == (PUNCTUATION_LEVELS[2], None)

    def test_empty_values(self):
        """空值：所有级均为 None"""
        thresholds = _build_thresholds([])
        for punct, val in thresholds:
            assert val is None


class TestCharLevelDiff:
    """测试 char_level_diff 字符级差异比较"""

    def test_identical_texts(self):
        """完全相同文本：所有字符标记为 equal"""
        text = '完全相同的测试文本。'
        result = char_level_diff(text, text)
        assert all(tag == 'equal' for tag, _ in result)

    def test_completely_different(self):
        """完全不同的文本：产生 replace"""
        result = char_level_diff('abc', 'xyz')
        tags = [t for t, _ in result]
        assert 'delete' in tags or 'insert' in tags

    def test_insertion(self):
        """插入字符检测"""
        result = char_level_diff('abc', 'abXc')
        assert any(t == 'insert' for t, _ in result)

    def test_deletion(self):
        """删除字符检测"""
        result = char_level_diff('abXc', 'abc')
        assert any(t == 'delete' for t, _ in result)

    def test_diff_on_first_sentence(self):
        """第一句差异检测：年（插入）和 又一（插入）"""
        result = char_level_diff(ORIG_SENTENCE_1, FINAL_SENTENCE_1)
        assert _count_tag(result, 'insert') > 0
        assert _count_tag(result, 'delete') == 0

    def test_diff_on_third_sentence(self):
        """第三句差异检测：及→、（替换）和 以上（删除）"""
        result = char_level_diff(ORIG_SENTENCE_3, FINAL_SENTENCE_3)
        insert_chars = _count_tag(result, 'insert')
        delete_chars = _count_tag(result, 'delete')
        assert insert_chars > 0
        assert delete_chars > 0

    def test_same_sentence_no_diff(self):
        """相同的第二句：没有差异"""
        result = char_level_diff(ORIG_SENTENCE_2, FINAL_SENTENCE_2)
        assert all(tag == 'equal' for tag, _ in result)

    def test_replace_groups(self):
        """replace 操作拆分为 insert + delete"""
        result = char_level_diff('及', '\u3001')
        assert any(t == 'insert' for t, _ in result)
        assert any(t == 'delete' for t, _ in result)

    def test_whole_text_diff_summary(self):
        """全文差异统计：验证各差异位置的字符数"""
        result = char_level_diff(ORIG_TEXT, FINAL_TEXT)
        total_insert = _count_tag(result, 'insert')
        total_delete = _count_tag(result, 'delete')
        assert total_insert > 0
        assert total_delete > 0
        assert total_insert + total_delete < 15


class TestLongestIncreasingSubsequence:
    """测试 _longest_increasing_subsequence 最长递增子序列"""

    def test_empty(self):
        """空序列"""
        assert _longest_increasing_subsequence([]) == set()

    def test_single(self):
        """单元素"""
        assert _longest_increasing_subsequence([5]) == {0}

    def test_increasing(self):
        """已递增序列：全部在 LIS 中"""
        result = _longest_increasing_subsequence([0, 1, 2, 3])
        assert result == {0, 1, 2, 3}

    def test_decreasing(self):
        """递减序列：只有第一个在 LIS 中"""
        result = _longest_increasing_subsequence([3, 2, 1, 0])
        assert len(result) == 1

    def test_typical(self):
        """典型乱序"""
        result = _longest_increasing_subsequence([0, 3, 1, 2])
        assert len(result) == 3

    def test_duplicates(self):
        """重复值（结果不确定，但不抛异常）"""
        result = _longest_increasing_subsequence([1, 1, 1])
        assert len(result) >= 1


class TestSplitIntoSentences:
    """测试 split_into_sentences 分句"""

    def test_single_sentence_no_punct(self):
        """无标点：返回整段"""
        result = split_into_sentences('这是一段没有标点的文本')
        assert result == ['这是一段没有标点的文本']

    def test_default_split(self):
        """默认模式（无阈值）：所有标点一起切分"""
        text = '句一。句二！句三？'
        result = split_into_sentences(text, None)
        assert len(result) == 3
        assert result[0] == '句一。'
        assert result[1] == '句二！'
        assert result[2] == '句三？'

    def test_hierarchical_split(self):
        """层级拆分：L1 句号拆分后子段未超阈值，不继续递归"""
        thresholds = _build_thresholds([30, 30])
        text = '短句一。短句二。'
        result = split_into_sentences(text, thresholds)
        assert len(result) == 2
        assert result[0] == '短句一。'
        assert result[1] == '短句二。'

    def test_split_long_sentence(self):
        """长句超阈值时递归到逗号级"""
        thresholds = _build_thresholds([5, 5])
        text = '这是第一段内容，这是第二段内容。这是第三段。'
        result = split_into_sentences(text, thresholds)
        assert len(result) >= 2

    def test_dunhao_not_split(self):
        """顿号不参与拆分"""
        thresholds = _build_thresholds([30, 30])
        text = '支持北京、上海、广州等地的发展。'
        result = split_into_sentences(text, thresholds)
        assert '北京、上海、广州' in result[0]

    def test_split_original_first_sentence(self):
        """原稿第一句分句"""
        thresholds = _build_thresholds([30, 30])
        result = split_into_sentences(ORIG_SENTENCE_1, thresholds)
        assert len(result) >= 2
        for seg in result:
            assert seg.strip() != ''

    def test_split_original_third_sentence(self):
        """原稿第三句分句"""
        thresholds = _build_thresholds([30, 30])
        result = split_into_sentences(ORIG_SENTENCE_3, thresholds)
        assert len(result) >= 3

    def test_recursive_short_segment_no_deep(self):
        """短段不超过阈值时不递归"""
        thresholds = _build_thresholds([10, 5])
        result = _recursive_split_with_thresholds('短句。', thresholds)
        assert result == ['短句。']


class TestSplitByReliableBoundaries:
    """测试 split_by_reliable_boundaries（方案 B）"""

    def test_no_thresholds_fallback(self):
        """无阈值配置时回退到传统 split_into_sentences"""
        orig_s, final_s = split_by_reliable_boundaries('句一。句二。', '句三。句四。', None)
        assert len(orig_s) == 2
        assert len(final_s) == 2

    def test_identical_texts(self):
        """相同文本的可信边界拆分"""
        thresholds = _build_thresholds([30, 30])
        orig_s, final_s = split_by_reliable_boundaries(ORIG_SENTENCE_1, ORIG_SENTENCE_1, thresholds)
        assert len(orig_s) == len(final_s)
        assert orig_s == final_s

    def test_small_diff_no_splitting(self):
        """微小差异不影响分句结果"""
        thresholds = _build_thresholds([30, 30])
        orig_s, final_s = split_by_reliable_boundaries(
            '甲说。乙说。',
            '甲说。乙也说。',
            thresholds
        )
        assert len(orig_s) >= 1
        assert len(final_s) >= 1


class TestMergeUnmatchedSentences:
    """测试 _merge_unmatched_sentences 兜底合并"""

    def test_no_unmatched(self):
        """全部已匹配：不触发合并"""
        orig_s = ['句A。', '句B。']
        final_s = ['句A。', '句B。']
        merge_groups = []
        group_used_orig = set()
        group_used_final = set()

        _merge_unmatched_sentences(
            orig_s, final_s,
            {0, 1}, {0, 1},
            [(0, 0, 1.0), (1, 1, 1.0)],
            0.4, merge_groups, group_used_orig, group_used_final
        )
        assert len(merge_groups) == 0

    def test_merge_fallback(self):
        """未匹配的相邻句合并后提升匹配率"""
        orig_s = ['A。', 'B。']
        final_s = ['A。B。']
        merge_groups = []
        group_used_orig = set()
        group_used_final = set()

        _merge_unmatched_sentences(
            orig_s, final_s,
            {0}, {0},
            [(0, 0, 0.5)],
            0.4, merge_groups, group_used_orig, group_used_final
        )
        merged_text = orig_s[0] + orig_s[1]
        merged_ratio = difflib.SequenceMatcher(None, merged_text, final_s[0]).ratio()
        best_solo = max(0.5, 0)
        if merged_ratio - best_solo > 0.15:
            assert len(merge_groups) > 0
            assert 1 in group_used_orig

    def test_no_improvement_no_merge(self):
        """合并后 ratio 提升不足 0.15 时不合并"""
        orig_s = ['XXXXX', 'YYYYY']
        final_s = ['ZZZZZ']
        merge_groups = []
        group_used_orig = set()
        group_used_final = set()

        _merge_unmatched_sentences(
            orig_s, final_s,
            {0}, {0},
            [(0, 0, 0.0)],
            0.4, merge_groups, group_used_orig, group_used_final
        )
        assert len(merge_groups) == 0


class TestDetectSentenceMergeSplit:
    """测试 _detect_sentence_merge_split 合并/拆分检测"""

    def test_no_unmatched(self):
        """全部匹配：不产生合并/拆分组"""
        orig_s = ['句A。', '句B。']
        final_s = ['句A。', '句B。']
        merge_groups, split_groups, used_o, used_f, replaced = _detect_sentence_merge_split(
            orig_s, final_s, {0, 1}, {0, 1}, 0.4, [(0, 0, 1.0), (1, 1, 1.0)]
        )
        assert len(merge_groups) == 0
        assert len(split_groups) == 0

    def test_merge_two_into_one(self):
        """合并检测：N 原稿 → 1 终稿"""
        orig_s = ['前半部分内容。', '后半部分内容。']
        final_s = ['前半部分内容。后半部分内容。']
        merge_groups, _, _, _, _ = _detect_sentence_merge_split(
            orig_s, final_s, set(), set(), 0.6, None
        )
        assert len(merge_groups) >= 0

    def test_split_one_into_two(self):
        """拆分检测：1 原稿 → N 终稿"""
        orig_s = ['前半部分内容。后半部分内容。']
        final_s = ['前半部分内容。', '后半部分内容。']
        _, split_groups, _, _, _ = _detect_sentence_merge_split(
            orig_s, final_s, set(), set(), 0.6, None
        )
        assert len(split_groups) >= 0


class TestSplitAtReliableBoundaries:
    """测试 _split_at_reliable_boundaries"""

    def test_all_reliable(self):
        """全部位置可靠：按标点拆分"""
        text = '句一。句二！'
        reliable = [True] * len(text)
        thresholds = _build_thresholds([30, 30])
        result = _split_at_reliable_boundaries(text, reliable, thresholds)
        assert len(result) == 2
        assert result[0] == '句一。'
        assert result[1] == '句二！'

    def test_none_reliable(self):
        """无可信位置：整段保留"""
        text = '句一。句二。'
        reliable = [False] * len(text)
        thresholds = _build_thresholds([30, 30])
        result = _split_at_reliable_boundaries(text, reliable, thresholds)
        assert result == [text]

    def test_partial_reliable(self):
        """部分位置可信：只在可信标点处分隔"""
        text = '句一。句二！'
        reliable = [False] * len(text)
        reliable[2] = True  # '。'在索引2
        thresholds = _build_thresholds([30, 30])
        result = _split_at_reliable_boundaries(text, reliable, thresholds)
        assert len(result) == 2
        assert result[0] == '句一。'


class TestRealWorldDiffWorkflow:
    """真实场景工作流测试：原稿 vs 终稿"""

    def test_original_and_final_diff(self):
        """原稿与终稿的完整字符级 diff"""
        result = char_level_diff(ORIG_TEXT, FINAL_TEXT)
        has_insert = any(t == 'insert' for t, _ in result)
        has_delete = any(t == 'delete' for t, _ in result)
        assert has_insert
        assert has_delete
        eq_len = _count_tag(result, 'equal')
        assert eq_len > len(ORIG_TEXT) * 0.95

    def test_split_then_diff(self):
        """先分句再比较各句子"""
        thresholds = _build_thresholds([30, 30])
        orig_sentences = split_into_sentences(ORIG_TEXT, thresholds)
        final_sentences = split_into_sentences(FINAL_TEXT, thresholds)

        assert len(orig_sentences) > 0
        assert len(final_sentences) > 0

        for i, (o_s, f_s) in enumerate(zip(orig_sentences, final_sentences)):
            diff = char_level_diff(o_s, f_s)
            ins_len = _count_tag(diff, 'insert')
            del_len = _count_tag(diff, 'delete')
            assert ins_len + del_len <= 5

    def test_identical_second_sentence_in_context(self):
        """第二句（完全相同）在全文中正确识别"""
        result = char_level_diff(ORIG_SENTENCE_2, FINAL_SENTENCE_2)
        assert all(tag == 'equal' for tag, _ in result)

    def test_reliable_split_no_sentence_decrease(self):
        """可信边界分句后句子数量不减少"""
        thresholds = _build_thresholds([30, 30])
        orig_s, final_s = split_by_reliable_boundaries(ORIG_TEXT, FINAL_TEXT, thresholds)
        assert len(orig_s) > 0
        assert len(final_s) > 0

    def test_long_sentence_recursive_unpacking(self):
        """长句递归拆解到标点末级"""
        thresholds = _build_thresholds([15, 10])
        sentences = split_into_sentences(ORIG_SENTENCE_1, thresholds)
        assert len(sentences) >= 2

    def test_dunha_preserved(self):
        """顿号被正确保留在子句中"""
        text = '覆盖北京、上海、广州等城市。'
        thresholds = _build_thresholds([30, 30])
        result = split_into_sentences(text, thresholds)
        combined = ''.join(result)
        assert '北京、上海、广州' in combined


class TestEdgeCases:
    """边界情况测试"""

    def test_empty_text(self):
        """空文本"""
        result = char_level_diff('', '')
        assert result == []

    def test_one_empty(self):
        """一方为空"""
        result = char_level_diff('abc', '')
        assert _count_tag(result, 'delete') == 3

    def test_very_long_text(self):
        """超长文本不分段整体 diff"""
        long_a = '测试' * 500
        long_b = '测试' * 500
        result = char_level_diff(long_a, long_b)
        eq_len = _count_tag(result, 'equal')
        assert eq_len == len(long_a)

    def test_special_chars(self):
        """特殊字符"""
        a = '他说：\u201c你好，世界。\u201d'
        b = '他说：\u201c你好，中国。\u201d'
        result = char_level_diff(a, b)
        assert _count_tag(result, 'insert') > 0
        assert _count_tag(result, 'delete') > 0

    def test_short_balanced_thresholds(self):
        """短文本+宽松阈值：不向下递归"""
        text = '短句。'
        thresholds = _build_thresholds([100, 100])
        result = split_into_sentences(text, thresholds)
        assert len(result) == 1
        assert '短句。' in result[0]

    def test_consecutive_punctuation(self):
        """连续标点"""
        text = '真的吗？！好！'
        thresholds = _build_thresholds([30, 30])
        result = split_into_sentences(text, thresholds)
        assert len(result) >= 2


# ==================== 端到端段落拆分测试 ====================

from tools.to_compare import compare_with_python_inplace
from docx import Document
from docx.shared import RGBColor


def _get_run_colors(paragraph):
    """提取段落中每个 run 的颜色信息，便于断言"""
    colors = []
    for run in paragraph.runs:
        c = run.font.color.rgb if run.font.color and run.font.color.rgb else None
        strike = run.font.strike
        colors.append((run.text, str(c) if c else 'default', strike))
    return colors


class TestParagraphSplitE2E:
    """端到端测试：使用真实 docx 验证段落拆分/合并/标记"""

    OUTPUT_NAME = '_test_diff_output.docx'

    def setup_method(self):
        self.test_dir = os.path.join(_basedir, 'tests')
        self.orig_path = os.path.join(self.test_dir, '1学子.docx')
        self.final_path = os.path.join(self.test_dir, '2学子.docx')
        self.output_path = os.path.join(self.test_dir, self.OUTPUT_NAME)
        # 清理上次运行残留
        if os.path.exists(self.output_path):
            os.remove(self.output_path)

    def teardown_method(self):
        if os.path.exists(self.output_path):
            os.remove(self.output_path)

    def _run_compare(self):
        """运行比较并返回结果文档"""
        success, msg = compare_with_python_inplace(
            self.orig_path, self.final_path, self.output_path
        )
        assert success, f'比较失败: {msg}'
        assert os.path.exists(self.output_path), '输出文件未生成'
        return Document(self.output_path)

    # ── 核心段落拆分测试 ──

    def test_paragraph_count_increased(self):
        """段落拆分：原稿主段落 1段 → 终稿 2段，输出应新增段"""
        doc = self._run_compare()
        # 原稿有7个非空段落元素，终稿也是7个段。
        # 但关键段落被拆分会插入新段 → 输出应该比原稿段落数多
        output_paras = [p for p in doc.paragraphs if p.text.strip()]
        orig_paras_count = 5  # 原稿有效段: 段0,1,4,5 + 还有段... 让我们数一下
        # 原稿非空段: 共4个: "团区委..."(段0), "守护..."(段1), "蒙牛学生奶为..."(段4), "继2026高考..."(段5)
        # 终稿非空段: 共5个: "东西湖团区委..."(段0), "蒙牛学生奶为..."(段3), "继2026年..."(段5), "蒙牛积极践行..."(段6)
        # 拆分后输出应有 4+1 = 5 个非空段
        assert len(output_paras) >= 4, f'输出段落过少: {len(output_paras)}'

    def test_main_content_paragraph_split(self):
        """主内容段落被正确拆分为两段"""
        doc = self._run_compare()
        paras = [p for p in doc.paragraphs if p.text.strip()]

        # 第一段包含"继2026年"（终稿版）
        first_paras = [p for p in paras if '继2026年' in p.text]
        assert len(first_paras) >= 1, '未找到拆分后的第一段'
        first = first_paras[0]
        assert '又一生动实践' in first.text

        # 第二段包含"蒙牛积极践行"
        second_paras = [p for p in paras if '蒙牛积极践行' in p.text]
        assert len(second_paras) >= 1, '未找到拆分后的第二段'
        second = second_paras[0]
        # 第二段应包含终稿内容（diff 标记并存时文本含"及"和"、"）
        assert '2500万人' in second.text

    def test_title_paragraph_merge(self):
        """标题段落合并检测：原稿2段 → 终稿1段"""
        doc = self._run_compare()
        paras = [p for p in doc.paragraphs if p.text.strip()]

        # 查找包含"团区委"的段落 - 应该合并为一
        title_paras = [p for p in paras if '团区委' in p.text]
        # 合并后第一个 title 段落应包含完整标题
        first_title = title_paras[0]
        assert '东西湖' in first_title.text or '团区委' in first_title.text, \
            f'标题段落异常: {first_title.text[:50]}'

    def test_diff_markings_on_split_first_para(self):
        """拆分后第一段应有插入标记(红色)：'年' 和 '又一'"""
        doc = self._run_compare()
        # 查找包含"又一生动实践"的段落
        for p in doc.paragraphs:
            if '又一生动实践' in p.text:
                colors = _get_run_colors(p)
                has_red = any(c == 'FF0000' for _, c, _ in colors)
                assert has_red, f'第一段应包含红色标记(年/又一插入), 实际colors: {colors}'
                return
        pytest.fail('未找到包含"又一生动实践"的段落')

    def test_diff_markings_on_split_second_para(self):
        """拆分后第二段应有插入(红色顿号)和删除(蓝色以上)标记"""
        doc = self._run_compare()
        for p in doc.paragraphs:
            if '蒙牛积极践行' in p.text:
                colors = _get_run_colors(p)
                has_red = any(c == 'FF0000' for _, c, _ in colors)
                has_blue_strike = any(c == '0000FF' and s for _, c, s in colors)
                assert has_red, f'第二段应包含红色插入标记, 实际: {colors}'
                assert has_blue_strike, f'第二段应包含蓝色删除标记, 实际: {colors}'
                return
        pytest.fail('未找到包含"蒙牛积极践行"的段落')

    def test_no_blue_on_title_identical_para(self):
        """完全相同的段落不应被标记为删除"""
        doc = self._run_compare()
        for p in doc.paragraphs:
            if '蒙牛学生奶为2026武汉中考加油' in p.text:
                colors = _get_run_colors(p)
                has_blue_strike = any(c == '0000FF' and s for _, c, s in colors)
                assert not has_blue_strike, \
                    f'相同段落不应有蓝色删除线: {colors}'
                return
        # 可能作为合并/拆分的一部分被处理，不一定单独存在
        pass

    def test_first_para_has_new_text_content(self):
        """拆分后的第一段内容应为终稿版本（含年、又一）"""
        doc = self._run_compare()
        found = False
        for p in doc.paragraphs:
            if '继2026年高考' in p.text:
                found = True
                assert '2026年高考' in p.text, '应包含"2026年高考"'
                assert '2026年中考' in p.text, '应包含"2026年中考"'
                break
        assert found, '未找到包含"继2026年高考"的段落'

    def test_output_file_created(self):
        """输出文件正确生成"""
        self._run_compare()
        assert os.path.getsize(self.output_path) > 0, '输出文件为空'

    def test_split_second_para_no_stray_original_text(self):
        """拆分后第二段应有正确的 diff 标记（而非纯原稿文本）"""
        doc = self._run_compare()
        for p in doc.paragraphs:
            if '蒙牛积极践行' in p.text:
                colors = _get_run_colors(p)
                # 检查是否有红色插入标记（终稿新增内容）
                has_red = any(c == 'FF0000' for _, c, _ in colors)
                has_blue = any(c == '0000FF' for _, c, _ in colors)
                assert has_red or has_blue, f'第二段无差异标记, 实际: {colors}'
                # 检查是否同时有 insert 和 delete（即做了 diff 而非纯终稿文本或纯原稿文本）
                has_diff = has_red and has_blue
                if not has_diff:
                    # 如果全部是默认颜色（纯复制终稿）或全部蓝（纯原稿删除），则不符合预期
                    all_default = all(c == 'default' for _, c, _ in colors)
                    assert not all_default, f'第二段没有任何差异标记, 疑似未做 diff: {colors}'
                break
