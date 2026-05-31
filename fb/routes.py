import os
import sys
import uuid
import time
import zipfile
import io
import shutil
import subprocess
import json
import requests

from flask import Blueprint, request, jsonify, send_file, Response, stream_with_context, g
from functools import wraps

from server.auth import login_required
from fb.database import get_db, get_visible_fb_ids, get_user_role
from tools.tool_defs import TOOL_SCRIPTS, TOOL_EXTENSIONS
from server import get_p2p_discovery, get_node_identity

fb_bp = Blueprint('fb', __name__, url_prefix='/api/fb')

PERMISSION_LEVELS = {'view': 0, 'edit': 1, 'manage': 2}


def _is_remote_fb(filebase_id):
    """检查文件库是否为远程文件库（共享自其他节点）"""
    from p2p.models import RemoteFilebaseStore
    store = RemoteFilebaseStore()
    info = store.get(filebase_id)
    return info is not None, info


_node_identity = None


def _get_node_identity():
    global _node_identity
    if _node_identity is None:
        from p2p.node import NodeIdentity
        _node_identity = NodeIdentity()
        _node_identity.load_or_create()
    return _node_identity


def _ensure_local_fb_route(f):
    """装饰器：对路由函数包装，如果文件库是远程则自动代理"""
    @wraps(f)
    def wrapper(*args, **kwargs):
        fb_id = kwargs.get('filebase_id') or kwargs.get('fb_id', '')
        if fb_id:
            is_remote, remote_info = _is_remote_fb(fb_id)
            if is_remote and remote_info:
                g.is_remote_fb = True
                g.remote_fb_info = remote_info
            else:
                g.is_remote_fb = False
        return f(*args, **kwargs)
    return wrapper


def _check_fb_permission(filebase_id, user_id, required_level):
    """检查用户对指定文件库的权限"""
    if not user_id:
        return False
    # 检查是否为管理员
    if get_user_role(user_id) == 'admin':
        return True
    db = get_db()
    row = db.execute(
        "SELECT permission_level FROM filebase_permissions WHERE filebase_id = ? AND user_id = ?",
        (filebase_id, user_id)
    ).fetchone()
    if row:
        actual = PERMISSION_LEVELS.get(row['permission_level'], -1)
        return actual >= PERMISSION_LEVELS.get(required_level, 0)
    kb_row = db.execute(
        "SELECT owner_id FROM filebases WHERE id = ?", (filebase_id,)
    ).fetchone()
    if kb_row and kb_row['owner_id'] == user_id:
        return True
    return False


def _require_fb_permission(required_level):
    """FB 专属权限校验装饰器，需在 @login_required 之后使用"""
    def decorator(f):
        @wraps(f)
        def decorated(*args, **kwargs):
            if request.method == 'OPTIONS':
                return f(*args, **kwargs)

            user_id = g.user_id
            if not user_id:
                return jsonify({'success': False, 'message': '未登录，请先登录'}), 401

            fb_id = kwargs.pop('fb_id', None)
            if fb_id:
                kwargs['filebase_id'] = fb_id

            filebase_id = kwargs.get('filebase_id')
            if filebase_id and not _check_fb_permission(filebase_id, user_id, required_level):
                is_remote, remote_info = _is_remote_fb(filebase_id)
                if not is_remote:
                    return jsonify({'success': False, 'message': '权限不足'}), 403

            return f(*args, **kwargs)
        return decorated
    return decorator


# ==================== 文件库 CRUD ====================

def _get_user_workspace(user_id=None):
    # 获取系统标准桌面目录（支持自定义位置）
    try:
        import platform
        system = platform.system()
        desktop = None

        if system == 'Windows':
            # Windows: 从注册表读取桌面位置
            import winreg
            key = winreg.OpenKey(winreg.HKEY_CURRENT_USER, r'Software\Microsoft\Windows\CurrentVersion\Explorer\User Shell Folders')
            desktop, _ = winreg.QueryValueEx(key, 'Desktop')
            winreg.CloseKey(key)
            # 展开环境变量（如 %USERPROFILE%）
            desktop = os.path.expandvars(desktop)
        else:
            # Linux/macOS: 使用 XDG 环境变量或标准位置
            desktop = os.environ.get('XDG_DESKTOP_DIR')
            if not desktop:
                # 检查 ~/.config/user-dirs.dirs
                user_dirs = os.path.join(os.path.expanduser('~'), '.config', 'user-dirs.dirs')
                if os.path.exists(user_dirs):
                    with open(user_dirs, 'r', encoding='utf-8') as f:
                        for line in f:
                            if line.strip().startswith('XDG_DESKTOP_DIR'):
                                desktop = line.split('=', 1)[1].strip().strip('"')
                                desktop = os.path.expandvars(desktop)
                                break
                if not desktop:
                    home = os.path.expanduser('~')
                    # 回退到硬编码检查
                    for name in ('Desktop', '桌面'):
                        candidate = os.path.join(home, name)
                        if os.path.isdir(candidate):
                            desktop = candidate
                            break

        if desktop and os.path.isdir(desktop):
            return desktop
    except Exception:
        pass

    # 都找不到时回退到用户家目录
    return os.path.expanduser('~')


def _get_trash_dir():
    """获取回收站目录，放在用户主目录下的 .trash 中"""
    trash_dir = os.path.join(os.path.expanduser('~'), '.trash')
    os.makedirs(trash_dir, exist_ok=True)
    return trash_dir


@fb_bp.route('/create-folder', methods=['POST'])
@login_required
def create_folder():
    user_id = g.user_id
    data = request.get_json()
    filebase_type = (data.get('filebase_type') or 'local').strip()
    name = (data.get('name') or '').strip()
    local_path = (data.get('local_path') or '').strip()

    if filebase_type == 'net':
        network_path = (data.get('network_path') or data.get('local_path') or '').strip()
        if not network_path:
            return jsonify({'success': False, 'message': '网络路径不能为空'})

        if not name:
            return jsonify({'success': False, 'message': '网络文件库名称不能为空'})

        db = get_db()
        filebase_id = str(uuid.uuid4())
        now = time.time()
        db.execute(
            "INSERT INTO filebases (id, name, owner_id, filebase_type, local_path, created_at) VALUES (?, ?, ?, 'net', ?, ?)",
            (filebase_id, name, user_id, network_path, now)
        )
        db.execute(
            "INSERT INTO filebase_permissions (filebase_id, user_id, permission_level) VALUES (?, ?, ?)",
            (filebase_id, user_id, 'manage')
        )
        db.commit()

        return jsonify({
            'success': True,
            'fb': {'id': filebase_id, 'name': name, 'owner_id': user_id, 'created_at': now, 'filebase_type': 'net',
                   'local_path': network_path}
        })

    # 未提供路径时，自动在桌面创建目录
    if not local_path:
        if not name:
            return jsonify({'success': False, 'message': '请输入文件库名称'})
        ws = _get_user_workspace(user_id)
        local_path = os.path.join(ws, name)
        os.makedirs(local_path, exist_ok=True)
        folder_name = name
    else:
        if not os.path.isdir(local_path):
            return jsonify({'success': False, 'message': '目录不存在或无效'})
        folder_name = os.path.basename(local_path.rstrip('/\\'))
        if not folder_name:
            return jsonify({'success': False, 'message': '无效的目录路径'})

    db = get_db()
    filebase_id = str(uuid.uuid4())
    now = time.time()

    existing = db.execute(
        "SELECT id FROM filebases WHERE owner_id = ? AND local_path = ?",
        (user_id, local_path)
    ).fetchone()
    if existing:
        return jsonify({'success': False, 'message': '该目录已在文件库中'})

    db.execute(
        "INSERT INTO filebases (id, name, owner_id, filebase_type, local_path, created_at) VALUES (?, ?, ?, 'local', ?, ?)",
        (filebase_id, folder_name, user_id, local_path, now)
    )
    db.execute(
        "INSERT INTO filebase_permissions (filebase_id, user_id, permission_level) VALUES (?, ?, ?)",
        (filebase_id, user_id, 'manage')
    )
    db.commit()

    return jsonify({
        'success': True,
        'fb': {'id': filebase_id, 'name': folder_name, 'owner_id': user_id, 'created_at': now, 'filebase_type': 'local',
               'local_path': local_path}
    })


@fb_bp.route('/copy-folder', methods=['POST'])
@login_required
def copy_folder():
    user_id = g.user_id
    data = request.get_json()
    filebase_id = (data.get('fb_id') or '').strip()
    new_name = (data.get('new_name') or '').strip()

    if not filebase_id or not new_name:
        return jsonify({'success': False, 'message': '参数不完整'})

    db = get_db()
    kb_row = db.execute("SELECT * FROM filebases WHERE id = ?", (filebase_id,)).fetchone()
    if not kb_row:
        return jsonify({'success': False, 'message': '源文件库不存在'})

    src_path = kb_row['local_path']
    if not os.path.isdir(src_path):
        return jsonify({'success': False, 'message': '源目录不存在'})

    ws = _get_user_workspace(user_id)
    dst_path = os.path.join(ws, new_name)

    counter = 1
    orig_name = new_name
    while os.path.exists(dst_path) and counter < 100:
        new_name = orig_name + '_' + str(counter)
        dst_path = os.path.join(ws, new_name)
        counter += 1
    if os.path.exists(dst_path):
        return jsonify({'success': False, 'message': '无法生成唯一的名称'})

    try:
        shutil.copytree(src_path, dst_path)
    except Exception as e:
        return jsonify({'success': False, 'message': '复制目录失败: ' + str(e)})

    new_filebase_id = str(uuid.uuid4())
    now = time.time()
    db.execute(
        "INSERT INTO filebases (id, name, owner_id, filebase_type, local_path, created_at) VALUES (?, ?, ?, 'local', ?, ?)",
        (new_filebase_id, new_name, user_id, dst_path, now)
    )
    db.execute(
        "INSERT INTO filebase_permissions (filebase_id, user_id, permission_level) VALUES (?, ?, ?)",
        (new_filebase_id, user_id, 'manage')
    )
    db.commit()

    return jsonify({
        'success': True,
        'fb': {'id': new_filebase_id, 'name': new_name, 'owner_id': user_id, 'created_at': now, 'filebase_type': 'local', 'local_path': dst_path}
    })


@fb_bp.route('/list', methods=['GET'])
@login_required
def list_fb():
    user_id = g.user_id
    is_admin = (get_user_role(user_id) == 'admin')
    db = get_db()
    ws = _get_user_workspace(user_id)

    db_kbs = {}
    rows = db.execute("SELECT * FROM filebases").fetchall()
    for row in rows:
        db_kbs[row['local_path']] = row

    # 分别处理本地文件库和网络文件库
    local_existing_paths = set()
    for row in rows:
        if row['filebase_type'] != 'net':
            local_existing_paths.add(row['local_path'])

    fs_paths = set()

    try:
        ws_entries = sorted(os.listdir(ws))
    except Exception as e:
        import logging
        logging.getLogger(__name__).warning('无法扫描工作空间 %s: %s', ws, e)
        ws_entries = []

    for entry_name in ws_entries:
        if entry_name.startswith('.') or entry_name == 'trash':
            continue
        entry_path = os.path.join(ws, entry_name)
        if os.path.isdir(entry_path):
            fs_paths.add(entry_path)
            if entry_path not in db_kbs:
                filebase_id = str(uuid.uuid4())
                now = time.time()
                db.execute(
                    "INSERT INTO filebases (id, name, owner_id, filebase_type, local_path, created_at) VALUES (?, ?, ?, 'local', ?, ?)",
                    (filebase_id, entry_name, user_id, entry_path, now)
                )
                db.execute(
                    "INSERT INTO filebase_permissions (filebase_id, user_id, permission_level) VALUES (?, ?, ?)",
                    (filebase_id, user_id, 'manage')
                )
            else:
                row = db_kbs[entry_path]
                if row['filebase_type'] != 'net' and row['name'] != entry_name:
                    db.execute("UPDATE filebases SET name = ? WHERE id = ?", (entry_name, row['id']))

    # 删除不存在于文件系统的本地文件库（仅限 workspaces 目录下的）
    for path in local_existing_paths - fs_paths:
        row = db_kbs[path]
        # 只删除位于当前 workspaces 目录下的文件库，保留用户自定义路径的文件库
        if path.startswith(ws + os.sep) or path == ws:
            db.execute("DELETE FROM filebase_permissions WHERE filebase_id = ?", (row['id'],))
            db.execute("DELETE FROM filebases WHERE id = ?", (row['id'],))
    db.commit()

    if is_admin:
        visible_rows = db.execute("SELECT * FROM filebases").fetchall()
    else:
        visible_ids = get_visible_fb_ids(user_id, False)
        visible_rows = []
        for filebase_id in visible_ids:
            r = db.execute("SELECT * FROM filebases WHERE id = ?", (filebase_id,)).fetchone()
            if r:
                visible_rows.append(r)

    kbs = []
    for row in visible_rows:
        perm_row = db.execute(
            "SELECT permission_level FROM filebase_permissions WHERE filebase_id = ? AND user_id = ?",
            (row['id'], user_id)
        ).fetchone()
        permission = 'manage' if row['owner_id'] == user_id or is_admin else (
            perm_row['permission_level'] if perm_row else 'view'
        )

        owner_username = row['owner_id'][:8]  # 单用户桌面版，owner_id 即节点 ID，取前缀作显示名

        local_path = row['local_path']
        display_path = local_path
        if local_path:
            norm = os.path.normpath(local_path)
            parts = norm.split(os.sep)
            try:
                ws_idx = [p.lower() for p in parts].index('workspaces')
                if ws_idx + 2 < len(parts):
                    parts[ws_idx + 1] = owner_username
                path_parts = parts[ws_idx + 1:]
                display_path = '/'.join(path_parts)
            except ValueError:
                pass

        kbs.append({
            'id': row['id'],
            'name': row['name'],
            'owner_id': row['owner_id'],
            'owner_username': owner_username,
            'display_path': display_path,
            'created_at': row['created_at'],
            'permission': permission,
            'filebase_type': row['filebase_type'] or 'local',
            'local_path': local_path
        })

    # 追加远程文件库（共享自其他节点的文件库）
    from p2p.models import RemoteFilebaseStore, TrustStore
    remote_store = RemoteFilebaseStore()
    trust_store = TrustStore()
    for fb_id, fb_info in remote_store.get_all().items():
        owner_node = trust_store.get_node_info(fb_info['owner_node_id'])
        owner_name = owner_node['display_name'] if owner_node else fb_info['owner_node_id'][:8]
        kbs.append({
            'id': fb_id,
            'name': f'[远程] {owner_name}/{fb_info["name"]}',
            'owner_id': fb_info['owner_node_id'],
            'owner_username': owner_name,
            'display_path': f'远程节点: {fb_info["owner_addr"]}',
            'created_at': fb_info.get('created_at', 0),
            'permission': fb_info['permission'],
            'filebase_type': 'remote',
            'local_path': ''
        })

    return jsonify({'success': True, 'kbs': kbs})


@fb_bp.route('/<fb_id>', methods=['PUT'])
@login_required
@_require_fb_permission('manage')
def rename_fb(filebase_id):
    data = request.get_json()
    name = (data.get('name') or '').strip()
    if not name:
        return jsonify({'success': False, 'message': '名称不能为空'})

    db = get_db()
    kb_row = db.execute("SELECT local_path, filebase_type FROM filebases WHERE id = ?", (filebase_id,)).fetchone()
    if not kb_row:
        return jsonify({'success': False, 'message': '文件库不存在'})

    filebase_type = kb_row['filebase_type'] or 'local'
    old_path = kb_row['local_path']

    # 网络文件库：只更新数据库记录，不操作文件系统
    if filebase_type == 'net':
        db.execute("UPDATE filebases SET name = ? WHERE id = ?", (name, filebase_id))
        db.commit()
        return jsonify({'success': True, 'message': '重命名成功'})

    parent_dir = os.path.dirname(old_path)
    new_path = os.path.join(parent_dir, name)

    if os.path.exists(new_path):
        return jsonify({'success': False, 'message': '同名目录已存在'})

    try:
        os.rename(old_path, new_path)
    except Exception as e:
        return jsonify({'success': False, 'message': '重命名目录失败: ' + str(e)})

    db.execute("UPDATE filebases SET name = ?, local_path = ? WHERE id = ?", (name, new_path, filebase_id))
    db.commit()
    return jsonify({'success': True, 'message': '重命名成功'})


def _cleanup_synced_data(user_id, filebase_id):
    """删除文件库时清理 KB 中的同步数据"""
    try:
        from kb.sync_worker import get_sync_worker
        worker = get_sync_worker()
        worker.cleanup_filebase(user_id, filebase_id)
    except Exception as e:
        import logging
        logging.getLogger(__name__).warning(f"Failed to cleanup synced data: {e}")


@fb_bp.route('/<fb_id>', methods=['DELETE'])
@login_required
@_require_fb_permission('manage')
def delete_fb(filebase_id):
    db = get_db()
    row = db.execute("SELECT id, name, local_path, owner_id, filebase_type FROM filebases WHERE id = ?", (filebase_id,)).fetchone()
    if not row:
        return jsonify({'success': False, 'message': '文件库不存在'})

    local_path = row['local_path']
    filebase_type = row['filebase_type'] or 'local'
    trash_dir = _get_trash_dir()

    # 网络文件库：只删除数据库记录
    if filebase_type == 'net':
        db.execute("DELETE FROM filebase_permissions WHERE filebase_id = ?", (filebase_id,))
        db.execute("DELETE FROM filebases WHERE id = ?", (filebase_id,))
        db.commit()
        _cleanup_synced_data(row['owner_id'], filebase_id)
        return jsonify({'success': True, 'message': '网络文件库已移至回收站'})

    if local_path.startswith(trash_dir):
        if os.path.isdir(local_path):
            shutil.rmtree(local_path)
        db.execute("DELETE FROM filebase_permissions WHERE filebase_id = ?", (filebase_id,))
        db.execute("DELETE FROM filebases WHERE id = ?", (filebase_id,))
        db.commit()
        _cleanup_synced_data(row['owner_id'], filebase_id)
        return jsonify({'success': True, 'message': '文件库已彻底删除'})

    fb_name = row['name']
    timestamp = str(int(time.time()))
    target = os.path.join(trash_dir, fb_name + '_' + timestamp)
    os.makedirs(trash_dir, exist_ok=True)

    if os.path.isdir(local_path):
        shutil.move(local_path, target)

    db.execute("DELETE FROM filebase_permissions WHERE filebase_id = ?", (filebase_id,))
    db.execute("DELETE FROM filebases WHERE id = ?", (filebase_id,))
    db.commit()
    _cleanup_synced_data(row['owner_id'], filebase_id)
    return jsonify({'success': True, 'message': '文件库已移至trash目录'})


@fb_bp.route('/trash', methods=['DELETE'])
@login_required
def clear_trash():
    user_id = g.user_id
    trash_dir = _get_trash_dir()
    if not os.path.isdir(trash_dir):
        return jsonify({'success': True, 'deleted': 0})

    # 先清理 trash 目录下对应的数据库记录，避免残留 DB 记录
    db = get_db()
    trash_entries = list(os.listdir(trash_dir))
    for entry in trash_entries:
        entry_path = os.path.join(trash_dir, entry)
        try:
            db.execute(
                "DELETE FROM filebase_permissions WHERE filebase_id IN "
                "(SELECT id FROM filebases WHERE local_path = ?)",
                (entry_path,)
            )
            db.execute("DELETE FROM filebases WHERE local_path = ?", (entry_path,))
        except Exception:
            pass
    db.commit()

    count = 0
    for entry in trash_entries:
        entry_path = os.path.join(trash_dir, entry)
        try:
            if os.path.isdir(entry_path):
                shutil.rmtree(entry_path)
                count += 1
            elif os.path.isfile(entry_path):
                os.remove(entry_path)
                count += 1
        except Exception:
            pass

    return jsonify({'success': True, 'deleted': count, 'message': f'已清空 {count} 个项目'})


@fb_bp.route('/trash-list', methods=['GET'])
@login_required
def list_trash():
    user_id = g.user_id
    trash_dir = _get_trash_dir()
    if not os.path.isdir(trash_dir):
        return jsonify({'success': True, 'items': []})

    items = []
    for entry in os.listdir(trash_dir):
        entry_path = os.path.join(trash_dir, entry)
        if os.path.isdir(entry_path):
            stat = os.stat(entry_path)
            size = 0
            try:
                size = sum(os.path.getsize(os.path.join(r, f)) for r, ds, fs in os.walk(entry_path) for f in fs)
            except Exception:
                pass
            items.append({'name': entry, 'path': entry_path, 'mtime': stat.st_mtime, 'size': size})
    items.sort(key=lambda x: x['mtime'], reverse=True)
    return jsonify({'success': True, 'items': items})


@fb_bp.route('/trash-restore', methods=['POST'])
@login_required
def restore_from_trash():
    user_id = g.user_id

    data = request.get_json()
    item_name = (data.get('name') or '').strip()
    if not item_name:
        return jsonify({'success': False, 'message': '未指定项目'})

    trash_dir = _get_trash_dir()
    src = os.path.join(trash_dir, item_name)
    if not os.path.isdir(src):
        return jsonify({'success': False, 'message': '项目不存在'})

    dst_name = item_name.rsplit('_', 1)[0] if '_' in item_name else item_name
    ws = _get_user_workspace(user_id)
    dst = os.path.join(ws, dst_name)
    orig_dst = dst
    counter = 1
    while os.path.exists(dst) and counter < 100:
        dst = orig_dst + '_' + str(counter)
        counter += 1
    if os.path.exists(dst):
        return jsonify({'success': False, 'message': '目标路径已存在'})

    shutil.move(src, dst)

    db = get_db()
    filebase_id = str(uuid.uuid4())
    now = time.time()
    name = os.path.basename(dst)
    db.execute(
        "INSERT INTO filebases (id, name, owner_id, filebase_type, local_path, created_at) VALUES (?, ?, ?, 'local', ?, ?)",
        (filebase_id, name, user_id, dst, now)
    )
    db.execute(
        "INSERT INTO filebase_permissions (filebase_id, user_id, permission_level) VALUES (?, ?, ?)",
        (filebase_id, user_id, 'manage')
    )
    db.commit()

    return jsonify({'success': True, 'fb': {'id': filebase_id, 'name': name, 'owner_id': user_id, 'created_at': now, 'filebase_type': 'local', 'local_path': dst}})


@fb_bp.route('/trash-item', methods=['DELETE'])
@login_required
def delete_trash_item():
    user_id = g.user_id

    name = request.args.get('name', '').strip()
    if not name:
        return jsonify({'success': False, 'message': '未指定项目'})

    trash_dir = _get_trash_dir()
    target = os.path.join(trash_dir, name)
    if not os.path.exists(target):
        return jsonify({'success': False, 'message': '项目不存在'})

    if not target.startswith(os.path.normpath(trash_dir)):
        return jsonify({'success': False, 'message': '路径非法'})

    try:
        if os.path.isdir(target):
            shutil.rmtree(target)
        else:
            os.remove(target)
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})

    return jsonify({'success': True})


@fb_bp.route('/<fb_id>/transfer', methods=['POST'])
@login_required
@_require_fb_permission('manage')
def transfer_fb(filebase_id):
    data = request.get_json()
    new_owner_id = (data.get('new_owner_id') or '').strip()
    keep_role = (data.get('keep_role') or 'editor').strip()

    if not new_owner_id:
        return jsonify({'success': False, 'message': '请指定新所有者'})
    if keep_role not in ('view', 'edit', 'manage'):
        keep_role = 'edit'

    db = get_db()
    kb_row = db.execute("SELECT owner_id FROM filebases WHERE id = ?", (filebase_id,)).fetchone()
    if not kb_row:
        return jsonify({'success': False, 'message': '文件库不存在'})

    db.execute("UPDATE filebases SET owner_id = ? WHERE id = ?", (new_owner_id, filebase_id))

    db.execute("DELETE FROM filebase_permissions WHERE filebase_id = ? AND user_id = ?", (filebase_id, new_owner_id))
    db.execute("INSERT INTO filebase_permissions (filebase_id, user_id, permission_level) VALUES (?, ?, ?)",
               (filebase_id, new_owner_id, 'manage'))

    db.execute("DELETE FROM filebase_permissions WHERE filebase_id = ? AND user_id = ?", (filebase_id, g.user_id))
    db.execute("INSERT INTO filebase_permissions (filebase_id, user_id, permission_level) VALUES (?, ?, ?)",
               (filebase_id, g.user_id, keep_role))

    db.commit()
    return jsonify({'success': True, 'message': '所有权移交成功'})


# ==================== 成员权限管理 ====================

@fb_bp.route('/<fb_id>/members', methods=['GET'])
@login_required
@_require_fb_permission('view')
def list_members(filebase_id):
    db = get_db()
    kb_row = db.execute("SELECT owner_id FROM filebases WHERE id = ?", (filebase_id,)).fetchone()

    members = []
    owner_id = kb_row['owner_id'] if kb_row else None
    if owner_id:
        members.append({
            'user_id': owner_id,
            'username': 'admin',
            'permission': 'manage',
            'is_owner': True
        })

    perm_rows = db.execute(
        "SELECT user_id, permission_level FROM filebase_permissions WHERE filebase_id = ? AND user_id != ?",
        (filebase_id, owner_id or '')
    ).fetchall()

    for row in perm_rows:
        members.append({
            'user_id': row['user_id'],
            'username': row['user_id'][:8],
            'permission': row['permission_level'],
            'is_owner': False
        })

    return jsonify({'success': True, 'members': members})


@fb_bp.route('/<fb_id>/members', methods=['POST'])
@login_required
@_require_fb_permission('manage')
def add_member(filebase_id):
    data = request.get_json()
    target_username = (data.get('username') or '').strip()
    permission = (data.get('permission') or 'view').strip()

    if permission not in ('view', 'edit', 'manage'):
        return jsonify({'success': False, 'message': '无效的权限级别'})

    # 单用户桌面版，唯一用户为本机节点
    from server.auth import _get_node_id
    target_user_id = _get_node_id()
    if target_username != 'admin':
        return jsonify({'success': False, 'message': '用户不存在'})

    db = get_db()
    existing = db.execute(
        "SELECT * FROM filebase_permissions WHERE filebase_id = ? AND user_id = ?",
        (filebase_id, target_user_id)
    ).fetchone()
    if existing:
        db.execute(
            "UPDATE filebase_permissions SET permission_level = ? WHERE filebase_id = ? AND user_id = ?",
            (permission, filebase_id, target_user_id)
        )
    else:
        db.execute(
            "INSERT INTO filebase_permissions (filebase_id, user_id, permission_level) VALUES (?, ?, ?)",
            (filebase_id, target_user_id, permission)
        )
    db.commit()

    return jsonify({'success': True, 'message': '成员已添加/更新'})


@fb_bp.route('/<fb_id>/members/<member_id>', methods=['PUT'])
@login_required
@_require_fb_permission('manage')
def update_member(filebase_id, member_id):
    data = request.get_json()
    permission = (data.get('permission') or 'view').strip()
    if permission not in ('view', 'edit', 'manage'):
        return jsonify({'success': False, 'message': '无效的权限级别'})

    db = get_db()
    db.execute(
        "UPDATE filebase_permissions SET permission_level = ? WHERE filebase_id = ? AND user_id = ?",
        (permission, filebase_id, member_id)
    )
    db.commit()
    return jsonify({'success': True, 'message': '权限已更新'})


@fb_bp.route('/<fb_id>/members/<member_id>', methods=['DELETE'])
@login_required
@_require_fb_permission('manage')
def remove_member(filebase_id, member_id):
    db = get_db()
    db.execute(
        "DELETE FROM filebase_permissions WHERE filebase_id = ? AND user_id = ?",
        (filebase_id, member_id)
    )
    db.commit()
    return jsonify({'success': True, 'message': '成员已移除'})


def _is_admin(user_id):
    """检查用户是否为管理员"""
    return get_user_role(user_id) == 'admin'


# ==================== 全文搜索 ====================

@fb_bp.route('/search', methods=['GET'])
@login_required
def search_documents():
    user_id = g.user_id
    q = (request.args.get('q') or '').strip()
    if not q:
        return jsonify({'success': True, 'results': []})

    is_admin = _is_admin(user_id)
    visible_ids = get_visible_fb_ids(user_id, is_admin)
    if not visible_ids:
        return jsonify({'success': True, 'results': []})

    db = get_db()
    results = []
    keywords = q.lower().split()

    for filebase_id in visible_ids:
        kb_row = db.execute("SELECT name, filebase_type, local_path FROM filebases WHERE id = ?", (filebase_id,)).fetchone()
        if not kb_row:
            continue
        fb_name = kb_row['name'] or ''
        local_path = (kb_row['local_path'] if 'local_path' in kb_row.keys() else '') or ''

        if local_path and os.path.isdir(local_path):
            results.extend(_search_local_dir(local_path, filebase_id, fb_name, keywords))

    return jsonify({'success': True, 'results': results, 'query': q})


def _search_local_dir(base_path, filebase_id, fb_name, keywords):
    results = []
    try:
        for root, dirs, files in os.walk(base_path):
            dirs[:] = [d for d in dirs if not d.startswith('~$')]
            for fname in files:
                if fname.startswith('~$'):
                    continue
                matched = False
                match_type = ''
                fname_lower = fname.lower()

                for kw in keywords:
                    if kw in fname_lower:
                        matched = True
                        match_type = 'filename'
                        break

                if not matched:
                    ext = os.path.splitext(fname)[1].lower()
                    if ext in ('.md', '.txt', '.html', '.htm', '.xml', '.json', '.csv'):
                        file_path = os.path.join(root, fname)
                        try:
                            with open(file_path, 'r', encoding='utf-8') as f:
                                content = f.read().lower()
                            for kw in keywords:
                                if kw in content:
                                    matched = True
                                    match_type = 'content'
                                    break
                        except Exception:
                            pass

                if matched:
                    full_path = os.path.join(root, fname)
                    rel_path = os.path.relpath(full_path, base_path).replace('\\', '/')
                    stat = os.stat(full_path)
                    results.append({
                        'document_id': rel_path,
                        'fb_id': filebase_id,
                        'fb_name': fb_name,
                        'filename': fname,
                        'file_type': os.path.splitext(fname)[1],
                        'file_size': stat.st_size,
                        'updated_at': stat.st_mtime,
                        'match_type': match_type,
                        'rel_path': rel_path
                    })
    except PermissionError:
        pass
    return results


# ==================== 本地目录浏览 ====================

def _resolve_local_path(db, filebase_id, subdir=''):
    kb_row = db.execute("SELECT local_path FROM filebases WHERE id = ?", (filebase_id,)).fetchone()
    if not kb_row:
        return None, None
    local_path = kb_row['local_path']
    target = os.path.join(local_path, subdir) if subdir else local_path
    target = os.path.normpath(target)
    if not target.startswith(os.path.normpath(local_path)):
        return None, None
    return local_path, target


def _trigger_fb_sync(filebase_id):
    """文件变更后触发同步"""
    try:
        from kb.sync_worker import get_sync_worker
        from flask import g
        worker = get_sync_worker()
        worker._trigger_sync(g.user_id, filebase_id)
    except Exception:
        import logging
        logging.getLogger(__name__).warning(f"Failed to trigger sync for {filebase_id}")


def _toggle_fb_sync_visibility(filebase_id, visible):
    """切换同步数据在 KB 中的可见性，visible=True 显示（imported/），False 隐藏（_disabled/）"""
    try:
        from flask import g
        from kb.database import get_db as get_kb_db
        conn = get_kb_db(g.user_id)
        if visible:
            old_prefix = f'_disabled/{filebase_id}/'
            new_prefix = f'imported/{filebase_id}/'
        else:
            old_prefix = f'imported/{filebase_id}/'
            new_prefix = f'_disabled/{filebase_id}/'

        conn.execute(
            "UPDATE wiki_files SET path = REPLACE(path, ?, ?) WHERE usr_id = ? AND path LIKE ?",
            (old_prefix, new_prefix, g.user_id, old_prefix + '%')
        )
        conn.execute(
            "UPDATE wiki_fts SET path = REPLACE(path, ?, ?) WHERE usr_id = ? AND path LIKE ?",
            (old_prefix, new_prefix, g.user_id, old_prefix + '%')
        )
        conn.commit()
    except Exception:
        import logging
        logging.getLogger(__name__).warning(f"Failed to toggle visibility for {filebase_id}")


@fb_bp.route('/<fb_id>/local-files', methods=['POST'])
@login_required
@_require_fb_permission('edit')
@_ensure_local_fb_route
def upload_local_files(filebase_id):
    if getattr(g, 'is_remote_fb', False):
        from p2p import proxy as p2p_proxy
        node = _get_node_identity()
        info = g.remote_fb_info
        subdir = request.args.get('subdir', '').strip()
        uploaded = []
        for key in request.files:
            for f in request.files.getlist(key):
                if not f.filename:
                    continue
                result = p2p_proxy.remote_upload_file(
                    info['owner_addr'], node, filebase_id,
                    subdir, f.filename, f.stream, 0
                )
                if result:
                    uploaded.extend(result.get('uploaded', []))
        return jsonify({'success': True, 'uploaded': uploaded})

    db = get_db()
    subdir = request.args.get('subdir', '').strip()
    local_path, target_dir = _resolve_local_path(db, filebase_id, subdir)
    if local_path is None:
        return jsonify({'success': False, 'message': '文件库不存在或路径非法'})

    if not os.path.isdir(target_dir):
        os.makedirs(target_dir, exist_ok=True)

    uploaded = []
    for key in request.files:
        for f in request.files.getlist(key):
            if not f.filename:
                continue
            # 安全处理文件名，保留相对路径但防止 ../ 跳出
            safe_filename = f.filename
            safe_filename = safe_filename.replace('..', '')
            if safe_filename.startswith('/') or safe_filename.startswith('\\'):
                safe_filename = safe_filename[1:]
            file_path = os.path.join(target_dir, safe_filename)
            os.makedirs(os.path.dirname(file_path), exist_ok=True)
            f.save(file_path)
            stat = os.stat(file_path)
            uploaded.append({
                'name': f.filename,
                'size': stat.st_size,
                'mtime': stat.st_mtime
            })

    if uploaded:
        _trigger_fb_sync(filebase_id)

    return jsonify({'success': True, 'uploaded': uploaded})


@fb_bp.route('/<fb_id>/local-files/dir', methods=['POST'])
@login_required
@_require_fb_permission('edit')
@_ensure_local_fb_route
def create_local_dir(filebase_id):
    if getattr(g, 'is_remote_fb', False):
        from p2p import proxy as p2p_proxy
        node = _get_node_identity()
        info = g.remote_fb_info
        data = request.get_json() or {}
        result = p2p_proxy.remote_create_dir(info['owner_addr'], node, filebase_id, data.get('name', ''), data.get('parent', ''))
        return jsonify(result or {'success': False, 'message': '远程节点不可用'})

    db = get_db()
    data = request.get_json() or {}
    name = (data.get('name') or '').strip()
    parent = (data.get('parent') or '').strip()
    if not name:
        return jsonify({'success': False, 'message': '目录名不能为空'})
    if '/' in name or '\\' in name:
        return jsonify({'success': False, 'message': '目录名不能包含路径分隔符'})

    local_path, target_dir = _resolve_local_path(db, filebase_id, parent)
    if local_path is None:
        return jsonify({'success': False, 'message': '文件库不存在或路径非法'})

    new_dir = os.path.join(target_dir, name)
    counter = 1
    orig_name = name
    while os.path.exists(new_dir) and counter < 100:
        name = orig_name + '_' + str(counter)
        new_dir = os.path.join(target_dir, name)
        counter += 1
    if os.path.exists(new_dir):
        return jsonify({'success': False, 'message': '无法生成唯一的目录名称'})

    os.makedirs(new_dir, exist_ok=True)
    rel = os.path.relpath(new_dir, local_path).replace('\\', '/')
    return jsonify({'success': True, 'path': rel})


@fb_bp.route('/<fb_id>/local-files/create', methods=['POST'])
@login_required
@_require_fb_permission('edit')
@_ensure_local_fb_route
def create_local_file(filebase_id):
    if getattr(g, 'is_remote_fb', False):
        from p2p import proxy as p2p_proxy
        node = _get_node_identity()
        info = g.remote_fb_info
        data = request.get_json() or {}
        result = p2p_proxy.remote_create_file(info['owner_addr'], node, filebase_id, data.get('name', ''), data.get('parent', ''))
        return jsonify(result or {'success': False, 'message': '远程节点不可用'})

    db = get_db()
    data = request.get_json() or {}
    name = (data.get('name') or '').strip()
    parent = (data.get('parent') or '').strip()
    if not name:
        return jsonify({'success': False, 'message': '文件名不能为空'})
    if '/' in name or '\\' in name:
        return jsonify({'success': False, 'message': '文件名不能包含路径分隔符'})

    local_path, target_dir = _resolve_local_path(db, filebase_id, parent)
    if local_path is None:
        return jsonify({'success': False, 'message': '文件库不存在或路径非法'})

    base, ext = os.path.splitext(name)
    if ext:
        filename = name
        default_ext = ext
    else:
        filename = name + '.md'
        default_ext = '.md'
    file_path = os.path.join(target_dir, filename)

    counter = 1
    while os.path.exists(file_path) and counter < 100:
        filename = base + '_' + str(counter) + default_ext
        file_path = os.path.join(target_dir, filename)
        counter += 1
    if os.path.exists(file_path):
        return jsonify({'success': False, 'message': '无法生成唯一的文件名'})

    with open(file_path, 'w', encoding='utf-8') as f:
        f.write('')
    rel = os.path.relpath(file_path, local_path).replace('\\', '/')
    _trigger_fb_sync(filebase_id)
    return jsonify({'success': True, 'path': rel})


@fb_bp.route('/<fb_id>/local-files/create-office', methods=['POST'])
@login_required
@_require_fb_permission('edit')
@_ensure_local_fb_route
def create_office_file(filebase_id):
    if getattr(g, 'is_remote_fb', False):
        from p2p import proxy as p2p_proxy
        node = _get_node_identity()
        info = g.remote_fb_info
        data = request.get_json() or {}
        result = p2p_proxy.remote_create_file(info['owner_addr'], node, filebase_id, data.get('name', ''), data.get('parent', ''))
        return jsonify(result or {'success': False, 'message': '远程节点不可用'})

    db = get_db()
    data = request.get_json() or {}
    name = (data.get('name') or '').strip()
    parent = (data.get('parent') or '').strip()
    if not name:
        return jsonify({'success': False, 'message': '文件名不能为空'})
    if '/' in name or '\\' in name:
        return jsonify({'success': False, 'message': '文件名不能包含路径分隔符'})

    local_path, target_dir = _resolve_local_path(db, filebase_id, parent)
    if local_path is None:
        return jsonify({'success': False, 'message': '文件库不存在或路径非法'})

    ext = os.path.splitext(name)[1].lower()
    if ext not in ('.docx', '.xlsx', '.pptx'):
        return jsonify({'success': False, 'message': f'不支持的文件类型: {ext}'})

    base = os.path.splitext(name)[0]
    file_path = os.path.join(target_dir, name)
    counter = 1
    while os.path.exists(file_path) and counter < 100:
        filename = f'{base}_{counter}{ext}'
        file_path = os.path.join(target_dir, filename)
        counter += 1
    if os.path.exists(file_path):
        return jsonify({'success': False, 'message': '无法生成唯一的文件名'})

    try:
        if ext == '.docx':
            from docx import Document
            from tools.mystyle import MyStyle
            doc = Document()
            set_page(doc)
            clear_styles(doc)
            add_my_styles(doc)
            doc.save(file_path)
        elif ext == '.xlsx':
            import openpyxl
            wb = openpyxl.Workbook()
            wb.save(file_path)
        elif ext == '.pptx':
            from pptx import Presentation
            prs = Presentation()
            prs.save(file_path)
    except ImportError as e:
        return jsonify({'success': False, 'message': f'缺少创建 {ext} 文件所需的库: {str(e)}'})
    except Exception as e:
        return jsonify({'success': False, 'message': f'创建文件失败: {str(e)}'})

    rel = os.path.relpath(file_path, local_path).replace('\\', '/')
    _trigger_fb_sync(filebase_id)
    return jsonify({'success': True, 'path': rel})


@fb_bp.route('/<fb_id>/local-files/content', methods=['PUT'])
@login_required
@_require_fb_permission('edit')
@_ensure_local_fb_route
def save_local_file_content(filebase_id):
    if getattr(g, 'is_remote_fb', False):
        from p2p import proxy as p2p_proxy
        node = _get_node_identity()
        info = g.remote_fb_info
        data = request.get_json() or {}
        result = p2p_proxy.remote_save_file(
            info['owner_addr'], node, filebase_id,
            (data.get('path') or '').strip(),
            data.get('content', ''),
            data.get('client_mtime', 0)
        )
        return jsonify(result or {'success': False, 'message': '远程节点不可用'})

    db = get_db()
    data = request.get_json() or {}
    path = (data.get('path') or '').strip()
    content = data.get('content', '')
    client_mtime = data.get('client_mtime', 0)

    if not path:
        return jsonify({'success': False, 'message': '未指定文件路径'})

    local_path, target = _resolve_local_path(db, filebase_id, path)
    if local_path is None:
        return jsonify({'success': False, 'message': '文件库不存在或路径非法'})

    # mtime 乐观锁：检查文件是否被外部修改
    if os.path.isfile(target) and client_mtime:
        current_mtime = os.stat(target).st_mtime
        if current_mtime != client_mtime:
            return jsonify({
                'success': False,
                'conflict': True,
                'message': '文件已被他人修改，是否覆盖？',
                'server_mtime': current_mtime
            }), 409

    os.makedirs(os.path.dirname(target), exist_ok=True)
    with open(target, 'w', encoding='utf-8') as f:
        f.write(content)
    stat = os.stat(target)
    _trigger_fb_sync(filebase_id)
    return jsonify({'success': True, 'mtime': stat.st_mtime})


@fb_bp.route('/<fb_id>/local-files', methods=['GET'])
@login_required
@_require_fb_permission('view')
@_ensure_local_fb_route
def list_local_files(filebase_id):
    if getattr(g, 'is_remote_fb', False):
        from p2p import proxy as p2p_proxy
        node = _get_node_identity()
        info = g.remote_fb_info
        subdir = request.args.get('subdir', '')
        result = p2p_proxy.remote_list_files(info['owner_addr'], node, filebase_id, subdir)
        return jsonify(result or {'success': False, 'message': '远程节点不可用'})

    db = get_db()
    kb_row = db.execute("SELECT local_path FROM filebases WHERE id = ?", (filebase_id,)).fetchone()
    if not kb_row:
        return jsonify({'success': False, 'message': '文件库不存在'})

    local_path = kb_row['local_path']
    if not os.path.isdir(local_path):
        return jsonify({'success': False, 'message': '本地目录不存在', 'files': [], 'categories': []})

    subdir = request.args.get('subdir', '').strip()
    tool = request.args.get('tool', '').strip()
    target_path = os.path.join(local_path, subdir) if subdir else local_path
    target_path = os.path.normpath(target_path)

    if not target_path.startswith(os.path.normpath(local_path)):
        return jsonify({'success': False, 'message': '不允许访问上级目录'})

    if not os.path.isdir(target_path):
        return jsonify({'success': False, 'message': '目录不存在'})

    extensions = TOOL_EXTENSIONS.get(tool) if tool else None

    files = []
    categories = []
    try:
        for entry in os.scandir(target_path):
            if entry.name.startswith('~$'):
                continue
            stat = entry.stat()
            if entry.is_dir():
                categories.append({
                    'name': entry.name,
                    'path': os.path.relpath(entry.path, local_path).replace('\\', '/')
                })
            elif entry.is_file():
                _, ext = os.path.splitext(entry.name)
                if extensions and ext.lower() not in extensions:
                    continue
                files.append({
                    'name': entry.name,
                    'path': os.path.relpath(entry.path, local_path).replace('\\', '/'),
                    'size': stat.st_size,
                    'mtime': stat.st_mtime,
                    'ext': ext.lower()
                })
    except PermissionError:
        return jsonify({'success': False, 'message': '没有权限访问此目录'})

    categories.sort(key=lambda x: x['name'].lower())
    files.sort(key=lambda x: x['name'].lower())

    return jsonify({
        'success': True,
        'files': files,
        'categories': categories,
        'current_path': subdir.replace('\\', '/') if subdir else ''
    })


@fb_bp.route('/<fb_id>/local-categories', methods=['GET'])
@login_required
@_require_fb_permission('view')
def list_local_categories(filebase_id):
    db = get_db()
    kb_row = db.execute("SELECT local_path FROM filebases WHERE id = ?", (filebase_id,)).fetchone()
    if not kb_row:
        return jsonify({'success': False, 'message': '文件库不存在'})

    local_path = kb_row['local_path']
    if not os.path.isdir(local_path):
        return jsonify({'success': False, 'categories': []})

    recursive = request.args.get('recursive', '0') == '1'

    if recursive:
        categories = _scan_categories_recursive(local_path, local_path)
        return jsonify({'success': True, 'categories': categories})

    subdir = request.args.get('subdir', '').strip()
    target_path = os.path.join(local_path, subdir) if subdir else local_path
    target_path = os.path.normpath(target_path)

    if not target_path.startswith(os.path.normpath(local_path)):
        return jsonify({'success': False, 'categories': []})

    categories = _scan_categories(target_path, local_path)
    return jsonify({'success': True, 'categories': categories})


def _scan_categories(target_path, base_path):
    categories = []
    try:
        for entry in os.scandir(target_path):
            if entry.name.startswith('~$') or not entry.is_dir():
                continue
            rel_path = os.path.relpath(entry.path, base_path).replace('\\', '/')
            child_categories = []
            try:
                for child in os.scandir(entry.path):
                    if not child.name.startswith('~$') and child.is_dir():
                        child_categories.append({
                            'name': child.name,
                            'path': os.path.relpath(child.path, base_path).replace('\\', '/')
                        })
            except PermissionError:
                pass
            child_categories.sort(key=lambda x: x['name'].lower())
            categories.append({
                'name': entry.name,
                'path': rel_path,
                'children': child_categories
            })
    except PermissionError:
        pass
    categories.sort(key=lambda x: x['name'].lower())
    return categories


def _scan_categories_recursive(target_path, base_path):
    categories = []
    try:
        for entry in os.scandir(target_path):
            if entry.name.startswith('~$') or not entry.is_dir():
                continue
            rel_path = os.path.relpath(entry.path, base_path).replace('\\', '/')
            children = _scan_categories_recursive(entry.path, base_path)
            categories.append({
                'name': entry.name,
                'path': rel_path,
                'children': children
            })
    except PermissionError:
        pass
    categories.sort(key=lambda x: x['name'].lower())
    return categories


@fb_bp.route('/<fb_id>/local-files/download', methods=['GET'])
@login_required
@_require_fb_permission('view')
@_ensure_local_fb_route
def download_local_file(filebase_id):
    if getattr(g, 'is_remote_fb', False):
        from p2p import proxy as p2p_proxy
        node = _get_node_identity()
        info = g.remote_fb_info
        resp = p2p_proxy.remote_download_file(info['owner_addr'], node, filebase_id, request.args.get('path', ''))
        if resp:
            return Response(resp.iter_content(chunk_size=8192), content_type=resp.headers.get('Content-Type', 'application/octet-stream'))
        return jsonify({'success': False, 'message': '远程节点不可用'})

    db = get_db()
    kb_row = db.execute("SELECT local_path FROM filebases WHERE id = ?", (filebase_id,)).fetchone()
    if not kb_row:
        return jsonify({'success': False, 'message': '文件库不存在'})

    local_path = kb_row['local_path']
    rel_path = request.args.get('path', '').strip()
    if not rel_path:
        return jsonify({'success': False, 'message': '未指定文件路径'})

    file_path = os.path.normpath(os.path.join(local_path, rel_path))
    if not file_path.startswith(os.path.normpath(local_path)):
        return jsonify({'success': False, 'message': '不允许访问上级目录'})

    if not os.path.isfile(file_path):
        return jsonify({'success': False, 'message': '文件不存在'})

    return send_file(file_path, as_attachment=True, download_name=os.path.basename(file_path))


@fb_bp.route('/<fb_id>/local-files/save-as', methods=['POST'])
@login_required
@_require_fb_permission('view')
def save_local_file_as(filebase_id):
    if getattr(g, 'is_remote_fb', False):
        from p2p import proxy as p2p_proxy
        node = _get_node_identity()
        info = g.remote_fb_info
        data = request.get_json() or {}
        rel_path = data.get('path', '').strip()
        save_path = data.get('save_path', '').strip()
        if not rel_path:
            return jsonify({'success': False, 'message': '未指定源文件路径'})
        if not save_path:
            return jsonify({'success': False, 'message': '未指定保存路径'})
        try:
            resp = p2p_proxy.remote_download_file(info['owner_addr'], node, filebase_id, rel_path)
            if not resp:
                return jsonify({'success': False, 'message': '远程节点不可用'})
            os.makedirs(os.path.dirname(save_path), exist_ok=True)
            with open(save_path, 'wb') as f:
                for chunk in resp.iter_content(chunk_size=8192):
                    if chunk:
                        f.write(chunk)
            return jsonify({'success': True, 'message': '文件已保存'})
        except Exception as e:
            return jsonify({'success': False, 'message': f'保存失败: {str(e)}'}), 500

    db = get_db()
    kb_row = db.execute("SELECT local_path FROM filebases WHERE id = ?", (filebase_id,)).fetchone()
    if not kb_row:
        return jsonify({'success': False, 'message': '文件库不存在'})

    local_path = kb_row['local_path']
    data = request.get_json() or {}
    rel_path = data.get('path', '').strip()
    save_path = data.get('save_path', '').strip()

    if not rel_path:
        return jsonify({'success': False, 'message': '未指定源文件路径'})
    if not save_path:
        return jsonify({'success': False, 'message': '未指定保存路径'})

    src_file = os.path.normpath(os.path.join(local_path, rel_path))
    if not src_file.startswith(os.path.normpath(local_path)):
        return jsonify({'success': False, 'message': '路径非法'})
    if not os.path.isfile(src_file):
        return jsonify({'success': False, 'message': '源文件不存在'})

    try:
        os.makedirs(os.path.dirname(save_path), exist_ok=True)
        shutil.copy2(src_file, save_path)
        return jsonify({'success': True, 'message': '文件已保存'})
    except Exception as e:
        return jsonify({'success': False, 'message': f'保存失败: {str(e)}'}), 500


# ==================== KB 同步管理 ====================

@fb_bp.route('/<fb_id>/sync', methods=['POST'])
@login_required
@_require_fb_permission('manage')
def toggle_sync(filebase_id):
    """切换文件库同步状态"""
    data = request.get_json() or {}
    enabled = bool(data.get('enabled', False))

    db = get_db()
    kb_row = db.execute("SELECT owner_id FROM filebases WHERE id = ?", (filebase_id,)).fetchone()

    if not kb_row:
        return jsonify({'success': False, 'message': '文件库不存在'}), 404

    if kb_row['owner_id'] != g.user_id:
        return jsonify({'success': False, 'message': '只有文件库所有者可以管理同步'}), 403

    db.execute("UPDATE filebases SET is_synced_to_kb = ? WHERE id = ?", (1 if enabled else 0, filebase_id))
    db.commit()

    _toggle_fb_sync_visibility(filebase_id, enabled)

    if enabled:
        _trigger_fb_sync(filebase_id)

    return jsonify({'success': True, 'enabled': enabled})


@fb_bp.route('/<fb_id>/sync-now', methods=['POST'])
@login_required
@_require_fb_permission('manage')
def sync_now(filebase_id):
    """手动触发立即同步"""
    db = get_db()
    kb_row = db.execute("SELECT owner_id, is_synced_to_kb FROM filebases WHERE id = ?", (filebase_id,)).fetchone()

    if not kb_row:
        return jsonify({'success': False, 'message': '文件库不存在'}), 404

    if kb_row['owner_id'] != g.user_id:
        return jsonify({'success': False, 'message': '只有文件库所有者可以触发同步'}), 403

    if not kb_row['is_synced_to_kb']:
        return jsonify({'success': False, 'message': '请先启用同步功能'}), 400

    try:
        from kb.sync_worker import get_sync_worker
        worker = get_sync_worker()
        worker.trigger_sync_now(g.user_id, filebase_id)
    except Exception as e:
        import logging
        logging.getLogger(__name__).warning(f"Failed to trigger sync: {e}")

    return jsonify({'success': True, 'message': '同步已触发'})


@fb_bp.route('/<fb_id>/sync-status', methods=['GET'])
@login_required
@_require_fb_permission('view')
def get_sync_status(filebase_id):
    """获取同步状态"""
    db = get_db()
    kb_row = db.execute("SELECT owner_id, is_synced_to_kb, local_path FROM filebases WHERE id = ?", (filebase_id,)).fetchone()

    if not kb_row:
        return jsonify({'success': False, 'message': '文件库不存在'}), 404

    try:
        from kb.sync_state import get_sync_state_manager
        state_manager = get_sync_state_manager()
        state = state_manager.load_state(kb_row['owner_id'], filebase_id)

        total_files = 0
        if os.path.exists(kb_row['local_path']):
            for root, dirs, files in os.walk(kb_row['local_path']):
                dirs[:] = [d for d in dirs if not d.startswith('.')]
                files = [f for f in files if not f.startswith('.') and not f.startswith('~')]
                total_files += len(files)

        from kb.sync_converters import can_convert
        syncable_count = 0
        if os.path.exists(kb_row['local_path']):
            for root, dirs, files in os.walk(kb_row['local_path']):
                for f in files:
                    if can_convert(os.path.join(root, f)):
                        syncable_count += 1

        from kb.sync_worker import get_sync_worker
        worker = get_sync_worker()
        is_syncing = filebase_id in worker._processing_filebases

        return jsonify({
            'success': True,
            'enabled': bool(kb_row['is_synced_to_kb']),
            'is_owner': kb_row['owner_id'] == g.user_id,
            'is_syncing': is_syncing,
            'status': {
                'total_files': total_files,
                'syncable_files': syncable_count,
                'synced_files': state.synced_files,
                'last_sync': state.last_sync
            }
        })
    except Exception as e:
        import logging
        logging.getLogger(__name__).error(f"Failed to get sync status: {e}")
        return jsonify({
            'success': True,
            'enabled': bool(kb_row['is_synced_to_kb']),
            'is_owner': kb_row['owner_id'] == g.user_id,
            'is_syncing': False,
            'status': {
                'total_files': 0,
                'syncable_files': 0,
                'synced_files': 0,
                'last_sync': None
            }
        })




@fb_bp.route('/<fb_id>/run-tool', methods=['POST'])
@login_required
@_require_fb_permission('edit')
@_ensure_local_fb_route
def run_tool_on_fb(filebase_id):
    if getattr(g, 'is_remote_fb', False):
        from p2p import proxy as p2p_proxy
        node = _get_node_identity()
        info = g.remote_fb_info
        data = request.get_json() or {}
        resp = p2p_proxy.remote_run_tool(info['owner_addr'], node, filebase_id, data.get('tool'), data.get('files', []), data.get('subdir', ''))
        if resp:
            return Response(resp.iter_content(chunk_size=4096), mimetype='text/event-stream', content_type='text/event-stream')
        return jsonify({'success': False, 'message': '远程节点不可用'})

    data = request.get_json()
    tool = data.get('tool')
    subdir = data.get('subdir', '').strip()
    files = data.get('files')

    if not tool:
        return jsonify({'success': False, 'message': '未指定工具'})

    if tool not in TOOL_SCRIPTS:
        return jsonify({'success': False, 'message': f'未知的工具: {tool}'})

    script_path = TOOL_SCRIPTS[tool]
    if not os.path.exists(script_path):
        return jsonify({'success': False, 'message': f'脚本不存在: {tool}'})

    db = get_db()
    kb_row = db.execute("SELECT local_path FROM filebases WHERE id = ?", (filebase_id,)).fetchone()
    if not kb_row:
        return jsonify({'success': False, 'message': '文件库不存在'})

    local_path = kb_row['local_path']
    target_path = os.path.normpath(os.path.join(local_path, subdir)) if subdir else local_path

    if not target_path.startswith(os.path.normpath(local_path)):
        return jsonify({'success': False, 'message': '不允许访问的路径'})

    if not os.path.isdir(target_path):
        return jsonify({'success': False, 'message': f'目录不存在: {subdir or "根目录"}'})

    if not files:
        extensions = TOOL_EXTENSIONS.get(tool, ('.docx',))
        files = []
        for f in os.listdir(target_path):
            full = os.path.join(target_path, f)
            if os.path.isfile(full) and f.lower().endswith(extensions):
                files.append(f)

    _user_id = g.user_id

    def generate():
        try:
            env = os.environ.copy()
            env['PYTHONPATH'] = os.path.dirname(os.path.abspath(__file__))
            env['USER_ID'] = _user_id

            cmd_args = [sys.executable, "-u", script_path]
            for f in files:
                full_path = os.path.join(target_path, f)
                cmd_args.append(full_path)

            process = subprocess.Popen(
                cmd_args,
                stdout=subprocess.PIPE,
                stderr=subprocess.STDOUT,
                text=True,
                bufsize=1,
                cwd=target_path,
                env=env
            )

            output_lines = []
            for line in iter(process.stdout.readline, ''):
                if line:
                    content = line.rstrip()
                    output_lines.append(content)
                    yield f'data: {json.dumps({"type": "output", "content": content})}\n\n'

            process.stdout.close()
            process.wait()

            success = process.returncode == 0
            if not success:
                error_msg = '\n'.join(output_lines) if output_lines else "执行失败"
                yield f'data: {json.dumps({"type": "end", "success": False, "error": error_msg})}\n\n'
            else:
                yield f'data: {json.dumps({"type": "end", "success": True})}\n\n'

        except Exception as e:
            yield f'data: {json.dumps({"type": "end", "success": False, "error": str(e)})}\n\n'

    return Response(stream_with_context(generate()), mimetype='text/event-stream')


@fb_bp.route('/<fb_id>/local-files/content', methods=['GET'])
@login_required
@_require_fb_permission('view')
@_ensure_local_fb_route
def get_local_file_content(filebase_id):
    if getattr(g, 'is_remote_fb', False):
        from p2p import proxy as p2p_proxy
        node = _get_node_identity()
        info = g.remote_fb_info
        result = p2p_proxy.remote_get_file_content(info['owner_addr'], node, filebase_id, request.args.get('path', ''))
        return jsonify(result or {'success': False, 'message': '远程节点不可用'})

    db = get_db()
    kb_row = db.execute("SELECT local_path FROM filebases WHERE id = ?", (filebase_id,)).fetchone()
    if not kb_row:
        return jsonify({'success': False, 'message': '文件库不存在'})

    local_path = kb_row['local_path']
    rel_path = request.args.get('path', '').strip()
    if not rel_path:
        return jsonify({'success': False, 'message': '未指定文件路径'})

    file_path = os.path.normpath(os.path.join(local_path, rel_path))
    if not file_path.startswith(os.path.normpath(local_path)):
        return jsonify({'success': False, 'message': '不允许访问上级目录'})

    if not os.path.isfile(file_path):
        return jsonify({'success': False, 'message': '文件不存在'})

    ext = os.path.splitext(file_path)[1].lower()

    if ext in ('.md', '.txt', '.html', '.htm', '.xml', '.json', '.csv', '.yaml', '.yml', '.py', '.js', '.css'):
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            return jsonify({'success': True, 'content': content, 'file_type': ext})
        except UnicodeDecodeError:
            return jsonify({'success': False, 'message': '无法以文本方式读取此文件'})

    if ext == '.docx':
        try:
            from docx import Document
            doc = Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs]
            return jsonify({'success': True, 'content': '\n'.join(paragraphs), 'file_type': ext})
        except Exception:
            return jsonify({'success': False, 'message': '无法读取 docx 内容'})

    if ext in ('.xlsx', '.xls'):
        try:
            import openpyxl
            wb = openpyxl.load_workbook(file_path, data_only=True)
            lines = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                lines.append(f'# {sheet_name}')
                for row in ws.iter_rows(values_only=True):
                    lines.append('\t'.join(str(c) if c is not None else '' for c in row))
            return jsonify({'success': True, 'content': '\n'.join(lines), 'file_type': ext})
        except Exception:
            return jsonify({'success': False, 'message': '无法读取 xlsx 内容'})

    return jsonify({'success': False, 'message': '不支持在线查看此文件类型的内容'})


SUPPORTED_PREVIEW_EXTS = {'.docx', '.pptx', '.ppt', '.xlsx', '.xls'}


@fb_bp.route('/<fb_id>/local-files/preview', methods=['GET'])
@login_required
@_require_fb_permission('view')
@_ensure_local_fb_route
def file_preview(filebase_id):
    if getattr(g, 'is_remote_fb', False):
        from p2p import proxy as p2p_proxy
        node = _get_node_identity()
        info = g.remote_fb_info
        result = p2p_proxy.remote_get_file_content(info['owner_addr'], node, filebase_id, request.args.get('path', ''))
        return jsonify(result or {'success': False, 'message': '远程节点不可用'})

    db = get_db()
    kb_row = db.execute("SELECT local_path FROM filebases WHERE id = ?", (filebase_id,)).fetchone()
    if not kb_row:
        return jsonify({'success': False, 'message': '文件库不存在'})

    local_path = kb_row['local_path']
    rel_path = request.args.get('path', '').strip()
    if not rel_path:
        return jsonify({'success': False, 'message': '未指定文件路径'})

    file_path = os.path.normpath(os.path.join(local_path, rel_path))
    if not file_path.startswith(os.path.normpath(local_path)):
        return jsonify({'success': False, 'message': '不允许访问上级目录'})

    if not os.path.isfile(file_path):
        return jsonify({'success': False, 'message': '文件不存在'})

    ext = os.path.splitext(file_path)[1].lower()
    if ext not in SUPPORTED_PREVIEW_EXTS:
        return jsonify({'success': False, 'message': f'不支持的预览格式: {ext}'})

    try:
        from kb.sync_converters import MarkItDownConverter
        converter = MarkItDownConverter()
        markdown = converter.convert(file_path)
        if markdown is None:
            return jsonify({'success': False, 'message': '文件转换失败'})
        return jsonify({'success': True, 'markdown': markdown, 'file_type': ext})
    except Exception as e:
        return jsonify({'success': False, 'message': f'预览失败: {str(e)}'})


@fb_bp.route('/<fb_id>/local-files/batch-download', methods=['POST'])
@login_required
@_require_fb_permission('view')
def batch_download_local(filebase_id):
    db = get_db()
    kb_row = db.execute("SELECT local_path FROM filebases WHERE id = ?", (filebase_id,)).fetchone()
    if not kb_row:
        return jsonify({'success': False, 'message': '文件库不存在'})

    local_path = kb_row['local_path']
    data = request.get_json() or {}
    paths = data.get('paths', [])

    if not paths:
        return jsonify({'success': False, 'message': '请选择文件'})

    memory_file = io.BytesIO()
    with zipfile.ZipFile(memory_file, 'w', zipfile.ZIP_DEFLATED) as zf:
        for rel_path in paths:
            abs_path = os.path.normpath(os.path.join(local_path, rel_path))
            if not abs_path.startswith(os.path.normpath(local_path)):
                continue
            if os.path.isfile(abs_path):
                zf.write(abs_path, os.path.basename(abs_path))
            elif os.path.isdir(abs_path):
                for root, dirs, files in os.walk(abs_path):
                    for f in files:
                        fp = os.path.join(root, f)
                        arc = os.path.relpath(fp, os.path.dirname(abs_path)).replace('\\', '/')
                        zf.write(fp, arc)

    memory_file.seek(0)
    return send_file(memory_file, mimetype='application/zip',
                     as_attachment=True, download_name='files.zip')


@fb_bp.route('/<fb_id>/local-files/batch-save-as', methods=['POST'])
@login_required
@_require_fb_permission('view')
def batch_save_local_files(filebase_id):
    data = request.get_json() or {}
    paths = data.get('paths', [])
    dest_dir = data.get('dest_dir', '').strip()

    if not paths:
        return jsonify({'success': False, 'message': '请选择文件'})
    if not dest_dir:
        return jsonify({'success': False, 'message': '未指定目标目录'})

    if getattr(g, 'is_remote_fb', False):
        from p2p import proxy as p2p_proxy
        node = _get_node_identity()
        info = g.remote_fb_info
        try:
            os.makedirs(dest_dir, exist_ok=True)
            for rel_path in paths:
                if rel_path.endswith('/') or rel_path.endswith('\\'):
                    continue
                fname = os.path.basename(rel_path.replace('\\', '/'))
                save_path = os.path.join(dest_dir, fname)
                resp = p2p_proxy.remote_download_file(info['owner_addr'], node, filebase_id, rel_path)
                if not resp:
                    continue
                with open(save_path, 'wb') as f:
                    for chunk in resp.iter_content(chunk_size=8192):
                        if chunk:
                            f.write(chunk)
            return jsonify({'success': True, 'message': '文件已保存'})
        except Exception as e:
            return jsonify({'success': False, 'message': f'保存失败: {str(e)}'}), 500

    db = get_db()
    kb_row = db.execute("SELECT local_path FROM filebases WHERE id = ?", (filebase_id,)).fetchone()
    if not kb_row:
        return jsonify({'success': False, 'message': '文件库不存在'})

    local_path = kb_row['local_path']
    try:
        os.makedirs(dest_dir, exist_ok=True)
        for rel_path in paths:
            abs_path = os.path.normpath(os.path.join(local_path, rel_path))
            if not abs_path.startswith(os.path.normpath(local_path)):
                continue
            fname = os.path.basename(rel_path.replace('\\', '/'))
            target = os.path.join(dest_dir, fname)
            if os.path.isfile(abs_path):
                shutil.copy2(abs_path, target)
            elif os.path.isdir(abs_path):
                shutil.copytree(abs_path, target, dirs_exist_ok=True)
        return jsonify({'success': True, 'message': '文件已保存'})
    except Exception as e:
        return jsonify({'success': False, 'message': f'保存失败: {str(e)}'}), 500


@fb_bp.route('/<fb_id>/local-files/replace', methods=['PUT'])
@login_required
@_require_fb_permission('edit')
@_ensure_local_fb_route
def replace_local_file(filebase_id):
    if getattr(g, 'is_remote_fb', False):
        from p2p import proxy as p2p_proxy
        node = _get_node_identity()
        info = g.remote_fb_info
        path = request.args.get('path', '').strip()
        if 'file' not in request.files:
            return jsonify({'success': False, 'message': '请选择文件'})
        upload_file = request.files['file']
        result = p2p_proxy.remote_replace_file(info['owner_addr'], node, filebase_id, path, upload_file.filename, upload_file.stream)
        return jsonify(result or {'success': False, 'message': '远程节点不可用'})

    db = get_db()
    kb_row = db.execute("SELECT local_path FROM filebases WHERE id = ?", (filebase_id,)).fetchone()
    if not kb_row:
        return jsonify({'success': False, 'message': '文件库不存在'})

    local_path = kb_row['local_path']
    rel_path = request.args.get('path', '').strip()
    if not rel_path:
        return jsonify({'success': False, 'message': '未指定文件路径'})

    file_path = os.path.normpath(os.path.join(local_path, rel_path))
    if not file_path.startswith(os.path.normpath(local_path)):
        return jsonify({'success': False, 'message': '不允许访问上级目录'})

    if not os.path.isfile(file_path):
        return jsonify({'success': False, 'message': '原文件不存在'})

    if 'file' not in request.files:
        return jsonify({'success': False, 'message': '请选择文件'})

    upload_file = request.files['file']
    if not upload_file.filename:
        return jsonify({'success': False, 'message': '文件名不能为空'})

    try:
        upload_file.save(file_path)
        new_stat = os.stat(file_path)
        _trigger_fb_sync(filebase_id)
        return jsonify({
            'success': True,
            'message': '文件已替换',
            'size': new_stat.st_size,
            'mtime': new_stat.st_mtime
        })
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})


@fb_bp.route('/<fb_id>/local-files/move', methods=['PUT'])
@login_required
@_require_fb_permission('edit')
@_ensure_local_fb_route
def move_local_items(filebase_id):
    if getattr(g, 'is_remote_fb', False):
        from p2p import proxy as p2p_proxy
        node = _get_node_identity()
        info = g.remote_fb_info
        data = request.get_json() or {}
        result = p2p_proxy.remote_move_items(info['owner_addr'], node, filebase_id, data.get('sources', []), data.get('dest', ''))
        return jsonify(result or {'success': False, 'message': '远程节点不可用'})

    db = get_db()
    kb_row = db.execute("SELECT local_path FROM filebases WHERE id = ?", (filebase_id,)).fetchone()
    if not kb_row:
        return jsonify({'success': False, 'message': '文件库不存在'})

    local_path = kb_row['local_path']
    data = request.get_json() or {}
    sources = data.get('sources', [])
    dest = (data.get('dest') or '').strip()

    if not sources:
        return jsonify({'success': False, 'message': '请选择要移动的项目'})
    if not dest:
        return jsonify({'success': False, 'message': '请指定目标目录'})

    dest_path = os.path.normpath(os.path.join(local_path, dest))
    if not dest_path.startswith(os.path.normpath(local_path)):
        return jsonify({'success': False, 'message': '目标目录非法'})

    if not os.path.isdir(dest_path):
        os.makedirs(dest_path, exist_ok=True)

    moved = 0
    errors = []
    for src in sources:
        src_path = os.path.normpath(os.path.join(local_path, src))
        if not src_path.startswith(os.path.normpath(local_path)):
            errors.append(f'{src}: 路径非法')
            continue
        if not os.path.exists(src_path):
            errors.append(f'{src}: 不存在')
            continue
        target = os.path.join(dest_path, os.path.basename(src_path))
        if os.path.exists(target):
            errors.append(f'{src}: 目标位置已存在同名项目')
            continue
        try:
            shutil.move(src_path, target)
            moved += 1
        except Exception as e:
            errors.append(f'{src}: {str(e)}')

    _trigger_fb_sync(filebase_id)
    return jsonify({
        'success': True,
        'moved': moved,
        'errors': errors
    })


@fb_bp.route('/<fb_id>/local-files', methods=['DELETE'])
@login_required
@_require_fb_permission('edit')
@_ensure_local_fb_route
def delete_local_items(filebase_id):
    if getattr(g, 'is_remote_fb', False):
        from p2p import proxy as p2p_proxy
        node = _get_node_identity()
        info = g.remote_fb_info
        result = p2p_proxy.remote_delete_items(info['owner_addr'], node, filebase_id, (request.get_json() or {}).get('paths', []))
        return jsonify(result or {'success': False, 'message': '远程节点不可用'})

    db = get_db()
    kb_row = db.execute("SELECT local_path FROM filebases WHERE id = ?", (filebase_id,)).fetchone()
    if not kb_row:
        return jsonify({'success': False, 'message': '文件库不存在'})

    local_path = kb_row['local_path']
    data = request.get_json() or {}
    paths = data.get('paths', [])
    if not paths:
        return jsonify({'success': False, 'message': '请选择要删除的项目'})

    deleted = 0
    errors = []
    for rel in paths:
        if not rel:
            continue
        target = os.path.normpath(os.path.join(local_path, rel))
        if not target.startswith(os.path.normpath(local_path)):
            errors.append(f'{rel}: 路径非法')
            continue
        try:
            if os.path.isdir(target):
                shutil.rmtree(target)
            elif os.path.isfile(target):
                os.remove(target)
            else:
                errors.append(f'{rel}: 文件不存在')
                continue
            deleted += 1
        except Exception as e:
            errors.append(f'{rel}: {str(e)}')

    _trigger_fb_sync(filebase_id)
    return jsonify({'success': True, 'deleted': deleted, 'errors': errors})


@fb_bp.route('/<fb_id>/local-files/rename', methods=['PUT'])
@login_required
@_require_fb_permission('edit')
@_ensure_local_fb_route
def rename_local_item(filebase_id):
    if getattr(g, 'is_remote_fb', False):
        from p2p import proxy as p2p_proxy
        node = _get_node_identity()
        info = g.remote_fb_info
        data = request.get_json() or {}
        result = p2p_proxy.remote_rename_item(info['owner_addr'], node, filebase_id, data.get('path', ''), data.get('new_name', ''))
        return jsonify(result or {'success': False, 'message': '远程节点不可用'})

    db = get_db()
    kb_row = db.execute("SELECT local_path FROM filebases WHERE id = ?", (filebase_id,)).fetchone()
    if not kb_row:
        return jsonify({'success': False, 'message': '文件库不存在'})

    local_path = kb_row['local_path']
    data = request.get_json() or {}
    rel_path = (data.get('path') or '').strip()
    new_name = (data.get('new_name') or '').strip()

    if not rel_path:
        return jsonify({'success': False, 'message': '未指定文件路径'})
    if not new_name:
        return jsonify({'success': False, 'message': '新名称不能为空'})
    if '/' in new_name or '\\' in new_name:
        return jsonify({'success': False, 'message': '新名称不能包含路径分隔符'})

    old = os.path.normpath(os.path.join(local_path, rel_path))
    if not old.startswith(os.path.normpath(local_path)):
        return jsonify({'success': False, 'message': '路径非法'})

    if not os.path.exists(old):
        return jsonify({'success': False, 'message': '文件或目录不存在'})

    parent_dir = os.path.dirname(old)
    new_path = os.path.join(parent_dir, new_name)
    if os.path.exists(new_path):
        return jsonify({'success': False, 'message': '同名文件或目录已存在'})

    try:
        os.rename(old, new_path)
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})

    new_rel = os.path.relpath(new_path, local_path).replace('\\', '/')
    _trigger_fb_sync(filebase_id)
    return jsonify({'success': True, 'new_path': new_rel})


@fb_bp.route('/<fb_id>/local-files/copy', methods=['POST'])
@login_required
@_require_fb_permission('edit')
@_ensure_local_fb_route
def copy_local_items(filebase_id):
    if getattr(g, 'is_remote_fb', False):
        from p2p import proxy as p2p_proxy
        node = _get_node_identity()
        info = g.remote_fb_info
        data = request.get_json() or {}
        result = p2p_proxy.remote_copy_items(info['owner_addr'], node, filebase_id, data.get('sources', []), data.get('dest', ''))
        return jsonify(result or {'success': False, 'message': '远程节点不可用'})

    db = get_db()
    kb_row = db.execute("SELECT local_path FROM filebases WHERE id = ?", (filebase_id,)).fetchone()
    if not kb_row:
        return jsonify({'success': False, 'message': '文件库不存在'})

    local_path = kb_row['local_path']
    data = request.get_json() or {}
    sources = data.get('sources', [])
    dest = (data.get('dest') or '').strip()

    if not sources:
        return jsonify({'success': False, 'message': '请选择要复制的项目'})

    dest_dir = os.path.normpath(os.path.join(local_path, dest)) if dest else local_path
    if not dest_dir.startswith(os.path.normpath(local_path)):
        return jsonify({'success': False, 'message': '目标路径非法'})

    os.makedirs(dest_dir, exist_ok=True)

    copied = 0
    errors = []
    for rel in sources:
        if not rel:
            continue
        src = os.path.normpath(os.path.join(local_path, rel))
        if not src.startswith(os.path.normpath(local_path)):
            errors.append(f'{rel}: 路径非法')
            continue
        try:
            basename = os.path.basename(src)
            dst = os.path.join(dest_dir, basename)

            counter = 1
            orig_dst = dst
            while os.path.exists(dst):
                name, ext = os.path.splitext(basename)
                dst = os.path.join(dest_dir, f'{name}_{counter}{ext}')
                counter += 1

            if os.path.isdir(src):
                shutil.copytree(src, dst)
            elif os.path.isfile(src):
                shutil.copy2(src, dst)
            else:
                errors.append(f'{rel}: 文件不存在')
                continue
            copied += 1
        except Exception as e:
            errors.append(f'{rel}: {str(e)}')

    _trigger_fb_sync(filebase_id)
    return jsonify({'success': True, 'copied': copied, 'errors': errors})


@fb_bp.route('/<fb_id>/local-files/open', methods=['GET'])
@login_required
@_require_fb_permission('view')
def open_local_file(filebase_id):
    db = get_db()
    kb_row = db.execute("SELECT local_path FROM filebases WHERE id = ?", (filebase_id,)).fetchone()
    if not kb_row:
        return jsonify({'success': False, 'message': '文件库不存在'})

    local_path = kb_row['local_path']
    rel_path = request.args.get('path', '').strip()
    if not rel_path:
        return jsonify({'success': False, 'message': '未指定文件路径'})

    file_path = os.path.normpath(os.path.join(local_path, rel_path))
    if not file_path.startswith(os.path.normpath(local_path)):
        return jsonify({'success': False, 'message': '不允许访问上级目录'})

    if not os.path.isfile(file_path):
        return jsonify({'success': False, 'message': '文件不存在'})

    # PDF、图片等浏览器可直接渲染的类型用 inline 预览，其他类型才强制下载
    ext = os.path.splitext(file_path)[1].lower()
    previewable_exts = {'.pdf', '.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp', '.svg', '.txt', '.csv'}
    as_attachment = ext not in previewable_exts

    return send_file(file_path, as_attachment=as_attachment, download_name=os.path.basename(file_path))


@fb_bp.route('/<fb_id>/local-files/open-with-app', methods=['GET'])
@login_required
@_require_fb_permission('view')
def open_with_app(filebase_id):
    """用系统默认软件打开文件（仅限本地文件库）"""
    import platform
    import subprocess

    db = get_db()
    kb_row = db.execute("SELECT local_path FROM filebases WHERE id = ?", (filebase_id,)).fetchone()
    if not kb_row:
        return jsonify({'success': False, 'message': '文件库不存在'})

    local_path = kb_row['local_path']
    rel_path = request.args.get('path', '').strip()
    if not rel_path:
        return jsonify({'success': False, 'message': '未指定文件路径'})

    file_path = os.path.normpath(os.path.join(local_path, rel_path))
    if not file_path.startswith(os.path.normpath(local_path)):
        return jsonify({'success': False, 'message': '路径非法'})

    if not os.path.isfile(file_path):
        return jsonify({'success': False, 'message': '文件不存在'})

    try:
        system = platform.system()
        if system == 'Windows':
            os.startfile(file_path)
        elif system == 'Darwin':  # macOS
            subprocess.run(['open', file_path], check=True)
        else:  # Linux
            subprocess.run(['xdg-open', file_path], check=True)
        return jsonify({'success': True, 'message': '已用本地软件打开'})
    except Exception as e:
        return jsonify({'success': False, 'message': f'打开失败: {str(e)}'}), 500


# ==================== P2P 共享管理 ====================

@fb_bp.route('/<fb_id>/share', methods=['POST'])
@login_required
@_require_fb_permission('manage')
def share_filebase(filebase_id):
    """将文件库共享给其他 P2P 节点"""
    data = request.get_json() or {}
    target_nodes = data.get('nodes', [])
    permission = (data.get('permission') or 'view').strip()

    if not target_nodes:
        return jsonify({'success': False, 'message': '请选择目标节点'})
    if permission not in ('view', 'edit', 'manage'):
        return jsonify({'success': False, 'message': '无效的权限级别'})

    db = get_db()
    row = db.execute("SELECT name, owner_id FROM filebases WHERE id = ?", (filebase_id,)).fetchone()
    if not row:
        return jsonify({'success': False, 'message': '文件库不存在'})

    fb_name = row['name']
    identity = get_node_identity()
    if not identity:
        return jsonify({'success': False, 'message': '节点身份未初始化'})

    owner_addr = f'{identity.node_id}:{identity.port}'

    now = time.time()
    success_count = 0
    for node in target_nodes:
        node_id = node.get('node_id', '')
        node_name = node.get('display_name', '')
        node_addr = node.get('addr', '')

        if not node_id or not node_addr:
            continue

        host = node_addr.split(':')[0] if ':' in node_addr else node_addr
        owner_full_addr = f'{host}:{identity.port}'

        try:
            import requests
            notify_url = f'http://{node_addr}/p2p/share/notify'
            resp = requests.post(notify_url, json={
                'fb_id': filebase_id,
                'fb_name': fb_name,
                'owner_addr': owner_full_addr,
                'permission': permission,
                'node_id': identity.node_id,
                'node_name': identity.display_name,
                'node_public_key': identity.get_public_key_b64()
            }, timeout=10)
            if resp.ok:
                success_count += 1
        except Exception as e:
            logger.warning("Failed to notify node %s: %s", node_addr, e)

        db.execute(
            "INSERT OR REPLACE INTO shared_nodes (filebase_id, node_id, node_name, node_addr, permission_level, created_at) VALUES (?, ?, ?, ?, ?, ?)",
            (filebase_id, node_id, node_name, node_addr, permission, now)
        )

    db.commit()

    return jsonify({
        'success': True,
        'message': f'已共享给 {success_count}/{len(target_nodes)} 个节点',
        'shared_count': success_count,
        'total': len(target_nodes)
    })


@fb_bp.route('/<fb_id>/shared-nodes', methods=['GET'])
@login_required
@_require_fb_permission('manage')
def list_shared_nodes(filebase_id):
    """获取文件库已共享的节点列表"""
    db = get_db()
    rows = db.execute(
        "SELECT node_id, node_name, node_addr, permission_level, created_at FROM shared_nodes WHERE filebase_id = ? ORDER BY created_at DESC",
        (filebase_id,)
    ).fetchall()

    nodes = []
    for r in rows:
        nodes.append({
            'node_id': r['node_id'],
            'node_name': r['node_name'],
            'node_addr': r['node_addr'],
            'permission': r['permission_level'],
            'created_at': r['created_at']
        })

    return jsonify({'success': True, 'nodes': nodes})


@fb_bp.route('/<fb_id>/shared-nodes/<node_id>', methods=['DELETE'])
@login_required
@_require_fb_permission('manage')
def revoke_share(filebase_id, node_id):
    """撤销对某个节点的共享"""
    db = get_db()
    db.execute("DELETE FROM shared_nodes WHERE filebase_id = ? AND node_id = ?", (filebase_id, node_id))
    db.commit()

    # 通知远程节点移除文件库
    row = db.execute("SELECT node_addr FROM shared_nodes WHERE filebase_id = ? AND node_id = ?",
                     (filebase_id, node_id)).fetchone()
    if row:
        node_addr = row['node_addr']
        try:
            import requests
            identity = get_node_identity()
            from p2p.proxy import _sign_request
            if identity:
                headers = _sign_request(identity, 'DELETE', f'/p2p/fb/{filebase_id}/revoke')
                requests.delete(f'http://{node_addr}/p2p/fb/{filebase_id}/revoke',
                                headers=headers, timeout=10)
        except Exception:
            pass

    return jsonify({'success': True, 'message': '共享已撤销'})


@fb_bp.route('/share-batch', methods=['POST'])
@login_required
def batch_share():
    """一键共享给所有在线节点（由前端指定节点列表）"""
    data = request.get_json() or {}
    fb_id = data.get('fb_id', '')
    permission = (data.get('permission') or 'view').strip()
    all_nodes = data.get('all_nodes', [])

    if not fb_id or not all_nodes:
        return jsonify({'success': False, 'message': '参数不完整'})

    db = get_db()
    row = db.execute("SELECT owner_id FROM filebases WHERE id = ?", (fb_id,)).fetchone()
    if not row:
        return jsonify({'success': False, 'message': '文件库不存在'})

    identity = get_node_identity()
    if not identity:
        return jsonify({'success': False, 'message': '节点身份未初始化'})

    fb_row = db.execute("SELECT name FROM filebases WHERE id = ?", (fb_id,)).fetchone()
    fb_name = fb_row['name'] if fb_row else ''
    owner_addr = f'{identity.node_id}:{identity.port}'
    now = time.time()

    success_count = 0
    for node in all_nodes:
        node_id = node.get('node_id', '')
        node_name = node.get('display_name', '')
        node_addr = node.get('addr', '')
        if not node_id or not node_addr:
            continue
        try:
            import requests
            host = node_addr.split(':')[0] if ':' in node_addr else node_addr
            notify_url = f'http://{node_addr}/p2p/share/notify'
            resp = requests.post(notify_url, json={
                'fb_id': fb_id,
                'fb_name': fb_name,
                'owner_addr': f'{host}:{identity.port}',
                'permission': permission,
                'node_id': identity.node_id,
                'node_name': identity.display_name,
                'node_public_key': identity.get_public_key_b64()
            }, timeout=10)
            if resp.ok:
                success_count += 1
        except Exception as e:
            logger.warning("batch share failed for %s: %s", node_addr, e)

        db.execute(
            "INSERT OR REPLACE INTO shared_nodes (filebase_id, node_id, node_name, node_addr, permission_level, created_at) VALUES (?, ?, ?, ?, ?, ?)",
            (fb_id, node_id, node_name, node_addr, permission, now)
        )
    db.commit()

    return jsonify({
        'success': True,
        'message': f'已共享给 {success_count}/{len(all_nodes)} 个在线节点',
        'shared_count': success_count,
        'total': len(all_nodes)
    })


# ==================== P2P 节点信息 ====================

@fb_bp.route('/p2p/node', methods=['GET'])
@login_required
def get_p2p_node_info():
    """获取本机 P2P 节点身份信息"""
    identity = get_node_identity()
    if not identity:
        return jsonify({'success': False, 'message': '节点身份未初始化'})

    return jsonify({
        'success': True,
        'node_id': identity.node_id,
        'display_name': identity.display_name,
        'port': identity.port,
    })


@fb_bp.route('/p2p/node', methods=['PUT'])
@login_required
def update_p2p_node_info():
    """更新本机 P2P 节点身份配置"""
    data = request.get_json() or {}
    display_name = (data.get('display_name') or '').strip()
    port = data.get('port')

    identity = get_node_identity()
    if not identity:
        return jsonify({'success': False, 'message': '节点身份未初始化'})

    changed = False
    if display_name and display_name != identity.display_name:
        identity.display_name = display_name
        changed = True
    if port is not None:
        try:
            new_port = int(port)
            if new_port < 1024 or new_port > 65535:
                return jsonify({'success': False, 'message': '端口号范围 1024-65535'})
            if new_port != identity.port:
                identity.port = new_port
                changed = True
        except (ValueError, TypeError):
            return jsonify({'success': False, 'message': '端口号无效'})

    if changed:
        if not identity.save_config():
            return jsonify({'success': False, 'message': '配置保存失败'})
        # 重启 P2P 发现服务以应用新配置
        try:
            discovery = get_p2p_discovery()
            if discovery:
                discovery.stop()
                discovery.display_name = identity.display_name
                discovery.port = identity.port
                discovery.start()
        except Exception as e:
            logger.warning("Failed to restart P2P discovery: %s", e)

    return jsonify({'success': True, 'message': '配置已更新'})


@fb_bp.route('/p2p/discovered-nodes', methods=['GET'])
@login_required
def get_discovered_nodes():
    """获取局域网发现的其他 P2P 节点"""
    discovery = get_p2p_discovery()
    if not discovery:
        return jsonify({'success': True, 'nodes': []})

    nodes = discovery.get_discovered_nodes()
    return jsonify({'success': True, 'nodes': nodes})


@fb_bp.route('/<fb_id>/convert-doc', methods=['POST'])
@login_required
@_require_fb_permission('edit')
def convert_doc_files(filebase_id):
    """扫描文件库中的 .doc 文件并转换为 .docx（对 KB 同步无感）"""
    db = get_db()
    row = db.execute("SELECT local_path FROM filebases WHERE id = ?", (filebase_id,)).fetchone()
    if not row:
        return jsonify({'success': False, 'message': '文件库不存在'}), 404

    local_path = row['local_path']
    if not local_path or not os.path.exists(local_path):
        return jsonify({'success': False, 'message': '文件库路径不存在'}), 400

    from tools.doc_process import doc_to_docx
    import logging
    logger = logging.getLogger(__name__)

    doc_dirs = set()
    for root, dirs, files in os.walk(local_path):
        dirs[:] = [d for d in dirs if not d.startswith('.')]
        for f in files:
            if f.lower().endswith('.doc') and not f.startswith('~$'):
                doc_dirs.add(root)

    if not doc_dirs:
        return jsonify({'success': True, 'message': '没有需要转换的 .doc 文件', 'converted': 0, 'failed': 0})

    total_ok = 0
    total_err = 0
    errors = []

    for workdir in sorted(doc_dirs):
        err_msg = doc_to_docx(workdir)
        if err_msg:
            total_err += 1
            errors.append(err_msg)
            logger.warning(f"doc 转换失败 [{workdir}]: {err_msg}")
        else:
            count = sum(1 for f in os.listdir(workdir)
                        if f.lower().endswith('.doc') and not f.startswith('~$'))
            total_ok += count

    return jsonify({
        'success': True,
        'message': f'转换完成: {total_ok} 成功, {total_err} 个目录有失败',
        'converted': total_ok,
        'failed_dirs': total_err,
        'errors': errors if errors else None
    })
