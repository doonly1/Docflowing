import os
import re
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from mystyle import set_page, add_my_styles
from doc_process import doc_to_docx, save_docx


def _make_page_field():
    """构造 "第 X 页" PAGE 域的 XML 片段（python-docx 无高层 API，直接拼 XML）。"""
    # w:fldChar begin / instrText PAGE / fldChar separate / placeholder / fldChar end
    w = 'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
    return parse_xml(
        '<w:r %s>'
        '<w:fldChar w:fldCharType="begin"/>'
        '<w:instrText xml:space="preserve"> PAGE </w:instrText>'
        '<w:fldChar w:fldCharType="separate"/>'
        '<w:t xml:space="preserve"> 1 </w:t>'
        '<w:fldChar w:fldCharType="end"/>'
        '</w:r>' % w
    )


def _make_text_run(text):
    """构造纯文本 run 的 XML 片段。"""
    return parse_xml('<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:t xml:space="preserve">%s</w:t></w:r>' % text)


def _clear_footer(footer):
    """清除页脚中所有段落内容。"""
    for p_elem in footer._element.findall(qn('w:p')):
        footer._element.remove(p_elem)


def add_page_number_single(file_path):
    """为单个文件添加页码（设置页面→添加样式→插入页码→保存）

    若是 .doc 文件会先转换为 .docx；
    若文件名非数字前缀，会先另存为数字前缀文件再添加页码，不覆盖原文件。
    """
    workdir = os.path.dirname(file_path)

    # .doc 文件先转换为 .docx
    if file_path.lower().endswith('.doc') and not file_path.lower().endswith('.docx'):
        doc_to_docx(workdir)
        file_path = os.path.splitext(file_path)[0] + '.docx'
        if not os.path.exists(file_path):
            print(f'转换失败：{os.path.basename(file_path)}')
            return

    # 非数字前缀文件需另存，避免覆盖原文件
    basename = os.path.basename(file_path)
    if not basename[:4].isdigit():
        doc = Document(file_path)
        save_path = save_docx(doc, basename, workdir)
        if not save_path:
            print(f'另存失败：{basename}')
            return
        file_path = save_path

    # 添加页码
    doc = Document(file_path)
    add_my_styles(doc)

    # 清除首页页脚、禁用奇偶页不同
    _clear_footer(doc.sections[0].footer)
    doc.settings.odd_and_even_pages_header_footer = False
    for section in doc.sections:
        section.footer.is_linked_to_previous = True

    footer = doc.sections[0].footer
    if not footer.paragraphs:
        footer.add_paragraph()
    para = footer.paragraphs[0]

    # 应用 foter 样式（字体、字号、右对齐等）
    try:
        para.style = doc.styles['Foter']
    except Exception:
        pass

    # "- 1 -" 格式
    para._p.append(_make_text_run('- '))
    para._p.append(_make_page_field())
    para._p.append(_make_text_run(' -'))

    doc.save(file_path)
    print(f'添加页码：{os.path.basename(file_path)} 成功。')


def add_page_numbers(workdir):
    """为目录中所有 docx 文件添加页码（含 .doc 转换、另存、页码插入均由 add_page_number_single 完成）"""
    doc_to_docx(workdir)
    files = [f for f in os.listdir(workdir)
             if f.lower().endswith('.docx') and not f.startswith("~$")]

    for file in files:
        add_page_number_single(os.path.join(workdir, file))

if __name__ == '__main__':
    import sys
    paths = sys.argv[1:] if len(sys.argv) > 1 else [os.path.dirname(__file__)]
    for path in paths:
        if os.path.isfile(path):
            add_page_number_single(path)
        elif os.path.isdir(path):
            add_page_numbers(path)
