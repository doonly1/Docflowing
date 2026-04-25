import os,time
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn

from mystyle import add_my_styles, para_fm, run_fm
from float_picture import parse_xml, nsdecls, add_float_picture
from load_config import load_user_config


def apply_font_scaling(run, scaling):
    """对 run 内的所有文字应用字体横向缩放

    :param run: python-docx 的 Run 对象
    :param scaling: 缩放比例
    """
    try:
        if not run.text:
            return
        rPr = run._r.get_or_add_rPr()
        # 先一次性移除所有旧的 w:w 元素，避免循环中反复 append
        for existing in rPr.findall(qn('w:w')):
            rPr.remove(existing)
        # 再添加新的 w:w，namespace 由 parse_xml 自动处理，无需手动写 xmlns
        w_elem = parse_xml('<w:w {} w:val="%d"/>'.format(nsdecls('w')) % scaling)
        rPr.append(w_elem)
    except Exception as e:
        print(f'字体缩放失败: {e}')


def find_signature(paras):
    """在段落列表中查找署名文本和日期段落。

    Returns:
        tuple: (sign_text, date_para)
            - sign_text: 署名文本，未找到时为 '未找到署名'
            - date_para: 日期段落对象，未找到时为 None
    """
    for para in paras:
        para_text = ''.join(para.text.split())
        if '年' in para_text and '月' in para_text \
                and '日' in para_text and len(para_text) < 12:
            sign_para = paras[paras.index(para) - 1]  # 日期上一段
            sign_text = ''.join(sign_para.text.split())
            if len(sign_text) > 3 and len(sign_text) < 25 and sign_text.isalpha():
                return sign_text, para
            return '未找到署名', para
    return '未找到署名', None


def add_seal(workdir):
    """在文档中添加印章"""

    # 检查是否有可处理的文件
    files_to_process = [f for f in os.listdir(workdir) if f.lower().endswith('.docx') and f[:4].isdigit()]
    if not files_to_process:
        print('没有可处理的文件，请先添加页码')
        return

    for file in files_to_process:
        doc = Document(os.path.join(workdir, file))
        paras = doc.paragraphs
        add_my_styles(doc)

        # 添加印章
        print(f'▼添加印章：{file}')
        signature, date_para = find_signature(paras)
        picture_name = get_stamp_path(signature)
        if date_para and picture_name:
            try:
                # 读取页面实际尺寸（带默认值：A4纸，左右边距2.8cm/2.6cm）
                section = doc.sections[0]
                left_margin = section.left_margin if section.left_margin is not None else Cm(2.8)
                right_margin = section.right_margin if section.right_margin is not None else Cm(2.6)
                page_width = section.page_width if section.page_width is not None else Cm(21)
                text_area_width_pt = (page_width - left_margin - right_margin) / 12700  # EMU转pt
    
                n = len(''.join(date_para.text.split()))
                seal_cm = 3.8
                seal_pt = seal_cm * 28.35

                date_text_width_pt = n * 16
                seal_x = text_area_width_pt - date_text_width_pt / 2 - seal_pt / 2 - 40
                seal_y = -seal_pt / 1.41

                add_float_picture(date_para, picture_name,
                                    width=Pt(seal_pt), height=Pt(seal_pt),
                                    pos_x=Pt(seal_x), pos_y=Pt(seal_y),
                                    pos_h_relative='margin', pos_v_relative='paragraph')
            except Exception as e:
                print(f'失败：{e}')
                
        #套红并缩放
        para_fm(doc.styles['Normal'], 0, 0, 28.95, 0, 0, 0, 'J')
        para = paras[0].insert_paragraph_before()   #最前段插入发文单位
        para.style = doc.styles['H0']
        para_fm(para,0,0,1,0,0,0,'C')
        run0 = para.add_run(signature+'文件')
        run_fm(run0,'方正小标宋简体',72,0,255,0,0)
        apply_font_scaling(run0, int(560/(len(signature)+2)))
        
        para = paras[0].insert_paragraph_before()   #插入空行
        para.style = doc.styles['Fangsong']
        para_fm(para,0,0,28.95,0,0,0,'C')
        run_fm(para.add_run(''),'仿宋', 16,0,0,0,0)

        #插入文号及分割线
        para = paras[0].insert_paragraph_before()
        para.style = doc.styles['Fangsong']
        para_fm(para,0,0,28.95,0,0,0,'C')
        wenhao = get_fawenzihao(signature)
        run_fm(para.add_run(wenhao),'仿宋', 16,0,0,0,0)
        pPr = para._p.get_or_add_pPr()
        # 底边框：红色 FF0000、单实线、粗细 12（=1.5pt）、间距 3
        pPr.insert(1, parse_xml(
            '<w:pBdr {}>'.format(nsdecls('w')) +
            '<w:bottom w:val="single" w:sz="12" w:space="3" w:color="FF0000"/>' +
            '</w:pBdr>'
        ))
        print(f'文号：{wenhao}')

        para = paras[0].insert_paragraph_before()   #插入空行
        para_fm(para,0,0,28.95,0,0,0,'C')
        run_fm(para.add_run(''),'仿宋', 16,0,0,0,0)

        #文档保存docx
        para_fm(doc.styles['Normal'], 0, 0, 28.95, 0, 0, 32, 'J')
        save_path = os.path.join(workdir, str(wenhao)+file[4:])
        doc.save(save_path)
        print('文档已保存：', os.path.normpath(save_path))

    
def get_stamp_path(sign_text):
    """获取印章图片路径"""
    config = load_user_config()
    if not config:
        return None
    
    # 获取项目根目录（脚本所在目录）
    script_dir = os.path.dirname(os.path.abspath(__file__))
    
    # 遍历配置，查找匹配的公司名称或简称
    for company, info in config.items():
        if not isinstance(info, dict):
            continue  # 跳过非公司配置（compare、last_workdir 等）
        if company == sign_text or sign_text in info.get('简称', []):
            # 优先使用配置中的印章位置
            stamp_path = info.get('印章位置')
            if stamp_path:
                # 如果是相对路径，基于项目根目录
                if not os.path.isabs(stamp_path):
                    stamp_path = os.path.join(script_dir, stamp_path)
                return stamp_path
            # 如果没有配置印章位置，使用默认方式
            return os.path.join(script_dir, 'config', sign_text + '.png')
    
    return None


def get_fawenzihao(sign_text):
    """获取发文字号"""
    config = load_user_config()
    daizi = '未找到'
    if config:
        for company, info in config.items():
            if not isinstance(info, dict):
                continue  # 跳过非公司配置
            if company == sign_text or sign_text in info.get('简称', []):
                daizi = info.get('代字')
                break  
    year = time.strftime("%Y", time.localtime())
    wenhao = daizi + '〔' + year + '〕' + '1号'
    return wenhao


if __name__ == "__main__":
    import sys
    if len(sys.argv) > 1:
        workdir = sys.argv[1]
    else:
        workdir = os.path.dirname(__file__)
    add_seal(workdir)
