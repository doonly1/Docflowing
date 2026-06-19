# -*- coding: utf-8 -*-
"""
Word文档比较工具
比较 原稿.docx 和 终稿.docx 文档，基于原稿生成差异比较结果
"""

import os
import sys
import glob
import difflib
from copy import deepcopy
import yaml
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from doc_process import doc_to_docx
from load_config import load_user_config
from logging_config import setup_logging, get_logger

setup_logging()
logger = get_logger(__name__)

# 开关：Diff 驱动的重分割（方案 B）
# True  = 先做全文 char_diff，只在 equal 区域的标点处分句（方案 B）
# False = 直接按标点分句（方案 A-1），保留 mismatch 回退（方案 A-2）
ENABLE_DIFF_DRIVEN_SPLIT = False


# 标点集按层级固定（中英文混合，从强到弱）
PUNCTUATION_LEVELS = [
    "。！？!?",     # L1：强终结符
    "；：;",     # L2：中强分隔符（不含半角:，避免数字/比分/时间误拆分）
    "，,",        # L3（末级）：弱分隔符，终端拆分
]
# 顿号"、"不参与拆分


def _build_thresholds(values):
    """将阈值数字列表与固定标点集合并为内部格式 [(punct, max_len), ...]"""
    result = []
    for i, punct in enumerate(PUNCTUATION_LEVELS):
        result.append((punct, values[i] if i < len(values) else None))
    return result


def load_compare_config():
    """从yaml配置文件加载比较参数（支持用户自定义配置）"""
    config = load_user_config()
    
    sentence_threshold = 0.40
    para_threshold = 0.40
    short_para_char_threshold = 60
    thresholds = _build_thresholds([30, 30])
    
    if config:
        compare_config = config.get('compare', {})
        sentence_threshold = compare_config.get('sentence_similarity_threshold', sentence_threshold)
        para_threshold = compare_config.get('para_similarity_threshold', para_threshold)
        short_para_char_threshold = compare_config.get('short_para_char_threshold', short_para_char_threshold)
        
        raw = compare_config.get('semantic_unit_thresholds')
        if raw:
            if isinstance(raw, list) and raw and isinstance(raw[0], (int, float)):
                thresholds = _build_thresholds(raw)
            else:
                thresholds = [(item['punctuation'], item.get('threshold')) for item in raw]
    
    return sentence_threshold, para_threshold, short_para_char_threshold, thresholds


def find_docx_files(workdir):
    """从目录中选择原版和终版文档
    
    Args:
        workdir: 工作目录
    
    Returns:
        (original, final) 文档路径元组
    """
    # 查找目录中的所有 docx 文件
    doc_to_docx(workdir)
    all_docs = glob.glob(os.path.join(workdir, "*.docx"))
    
    if len(all_docs) < 2:
        logger.warning("目录中只有 %s 个 docx 文档，至少需要 2 个进行比较", len(all_docs))
        return None, None
    
    # 目录中有 2 个及以上文档，让用户选择
    def select(prompt, choices):
        logger.info("\n%s:", prompt)
        for i, doc in enumerate(choices, 1):
            logger.info("  %s. %s", i, os.path.basename(doc))
        if not sys.stdin.isatty():
            logger.warning("非交互模式，自动选择第一个文件")
            return choices[0]
        while True:
            try:
                choice = input(f"(回车默认1):").strip()
                idx = int(choice) if choice else 1
                if idx == 0:
                    return None
                if 1 <= idx <= len(choices):
                    return choices[idx - 1]
                logger.warning("请输入1-%s", len(choices))
            except ValueError:
                logger.warning("请输入数字")
            except EOFError:
                logger.warning("stdin 不可用，自动选择第一个文件")
                return choices[0]
    
    # 选择原稿
    original = select("选择原稿", all_docs)
    if not original:
        return None, None
    
    # 选择终稿
    final = select("选择终稿", all_docs)
    if not final:
        return None, None

    return original, final


def get_paragraphs_with_style(doc):
    """获取文档所有段落及样式信息"""
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text
        paragraphs.append({
            'text': text,
            'style': para.style.name if para.style else 'Normal'
        })
    return paragraphs


def char_level_diff(original_text, final_text):
    """
    字符级别差异比较
    返回: list of (标记, 字符)
    标记: 'equal', 'delete', 'insert'
    """
    # 使用SequenceMatcher进行字符级比较
    matcher = difflib.SequenceMatcher(None, original_text, final_text)
    
    result = []
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == 'equal':
            result.append(('equal', original_text[i1:i2]))
        elif tag == 'delete':
            result.append(('delete', original_text[i1:i2]))
        elif tag == 'insert':
            result.append(('insert', final_text[j1:j2]))
        elif tag == 'replace':
            # 替换操作：新增在前，删除在后
            result.append(('insert', final_text[j1:j2]))
            result.append(('delete', original_text[i1:i2]))
    
    return result


def _split_at_reliable_boundaries(text, reliable_flags, thresholds, level=0, base_offset=0):
    """
    在可靠位置按层级递归拆分文本。

    只有落在 reliable_flags 对应位置为 True 的区域内的标点，才作为有效边界。
    这确保了分割边界在原稿和终稿中都是存在的、未被编辑的。

    拆分结果包含标点：每个句子片段以边界标点结尾（如果标点可靠）。

    Args:
        text: 待分句文本
        reliable_flags: 与 text 等长的 bool 列表，标记每个位置是否可靠
        thresholds: [(punct_chars, max_len), ...] 按层级递归
        level: 当前递归层级
        base_offset: 当前 text 在原始 reliable_flags 中的起始偏移量
    """
    import re
    if level >= len(thresholds):
        return [text]  # 超出定义层级，不再拆分

    punct_chars, max_len = thresholds[level]
    is_last_level = (level == len(thresholds) - 1)

    # 用捕获组保留标点，同时收集每个标点的位置索引
    parts = re.split(f'([{re.escape(punct_chars)}])', text)

    # 收集可靠标点的位置
    # parts 格式：[内容, 标点, 内容, 标点, ..., 内容(可能无)]
    # 需要映射回 reliable_flags
    reliable_boundaries = []  # [(char_pos, punct), ...]
    char_pos = 0
    for i in range(0, len(parts) - 1, 2):
        content = parts[i]
        punct = parts[i + 1]
        # 内容区域
        char_pos += len(content)
        # 标点：检查其位置是否可靠
        if char_pos < len(reliable_flags) and reliable_flags[char_pos]:
            reliable_boundaries.append((char_pos, punct))
        char_pos += len(punct)
    # 最后一段（无标点跟随的内容）
    # 不需要额外处理 char_pos

    # 按可靠边界拆分，结果包含标点
    segments = []
    if reliable_boundaries:
        prev = 0
        for bound_pos, punct in reliable_boundaries:
            # 片段：prev 到标点位置（含标点）
            if bound_pos >= prev:
                segments.append(text[prev:bound_pos + 1])
            prev = bound_pos + 1
        # 最后一段（标点之后的剩余内容）
        if prev < len(text):
            segments.append(text[prev:])
    else:
        # 没有可靠边界，整段保留
        segments = [text]

    # 每个子段独立判断是否继续递归
    result = []
    seg_start_offset = 0  # 当前 segment 在 text 中的起始偏移
    for seg in segments:
        if not seg:
            continue
        seg_len = len(seg)
        if is_last_level:
            # 末级：直接保留，不再递归
            result.append(seg)
        elif seg_len <= max_len:
            # 未超阈值 → 语义单元，保留
            result.append(seg)
        else:
            # 超过阈值 → 递归下钻
            # 使用偏移量计算子段在 reliable_flags 中的对应位置，避免 text.index() 的重复匹配问题
            seg_start = base_offset + seg_start_offset
            seg_end = seg_start + seg_len
            seg_reliable = reliable_flags[seg_start_offset:seg_start_offset + seg_len]
            sub_segs = _split_at_reliable_boundaries(seg, seg_reliable, thresholds, level + 1, seg_start)
            result.extend(sub_segs)
        seg_start_offset += seg_len

    return result


def split_by_reliable_boundaries(orig_text, final_text, semantic_unit_thresholds=None):
    """
    Diff 驱动的可靠边界分句（方案 B）。
    
    核心思路：先做全文 char_diff，再从 opcodes 的 equal 区域中提取可靠的标点边界。
    - equal 区域的标点：在原稿和终稿中都存在，未被编辑，是可靠的分割边界
    - replace/insert/delete 区域的标点：可能是编辑产物，不作为边界
    
    这样从根本上避免了"标点本身被修改"导致的分句不一致问题。
    
    Args:
        orig_text: 原稿文本
        final_text: 终稿文本
        semantic_unit_thresholds: 多级语义单元拆分配置
    
    Returns:
        (orig_sentences, final_sentences): 原稿和终稿的分句结果
    """
    # 向后兼容：无配置时使用传统分句
    if semantic_unit_thresholds is None:
        orig_sentences = split_into_sentences(orig_text, None)
        final_sentences = split_into_sentences(final_text, None)
        return orig_sentences, final_sentences
    
    # 1. 获取 opcodes
    matcher = difflib.SequenceMatcher(None, orig_text, final_text)
    opcodes = matcher.get_opcodes()
    
    # 2. 标记原稿和终稿中哪些字符位置是 equal（可靠的）
    orig_reliable = [False] * len(orig_text)
    final_reliable = [False] * len(final_text)
    
    for tag, i1, i2, j1, j2 in opcodes:
        if tag == 'equal':
            for i in range(i1, i2):
                orig_reliable[i] = True
            for j in range(j1, j2):
                final_reliable[j] = True
    
    # 3. 按层级在可靠边界处拆分
    orig_sentences = _split_at_reliable_boundaries(
        orig_text, orig_reliable, semantic_unit_thresholds
    )
    final_sentences = _split_at_reliable_boundaries(
        final_text, final_reliable, semantic_unit_thresholds
    )
    
    return orig_sentences, final_sentences


def split_into_sentences(text, semantic_unit_thresholds=None):
    """按语义单元层次递归分割句子
    
    多级拆分策略（默认）：
    L1：先用强标点（。！？）切分，子句 ≤ 30 字则停止
    L2：超长子句再按次强标点（；：）切分，子句 ≤ 30 字则停止
    L3（末级）：按逗号切分，不再判断阈值。顿号永远不参与切分
    
    Args:
        text: 待分句文本
        semantic_unit_thresholds: [(punctuation_chars, max_len), ...] 按此顺序递归
                                  末级的 max_len 被忽略（永远拆分到该级）
                                  None 表示向后兼容的旧行为（逗号+句号一起切）
    """
    import re
    if semantic_unit_thresholds is None:
        # 向后兼容：所有标点一起切
        sentences = re.split(r'([，,。！!？?；;：:])', text)
        result = []
        for i in range(0, len(sentences) - 1, 2):
            result.append(sentences[i] + sentences[i + 1])
        if len(sentences) % 2 == 1 and sentences[-1]:
            result.append(sentences[-1])
        return result
    
    return _recursive_split_with_thresholds(text, semantic_unit_thresholds)


def _recursive_split_with_thresholds(text, thresholds, level=0):
    """递归按级切分
    
    非末级：用对应标点切分，子段 ≤ 阈值则停止下钻
    末级：直接切分到该级，不再判断阈值
    """
    import re
    if level >= len(thresholds):
        return [text]  # 超出定义层级，不再拆分
    
    punct_chars, max_len = thresholds[level]
    is_last_level = (level == len(thresholds) - 1)
    
    # 用捕获组保留标点
    parts = re.split(f'([{re.escape(punct_chars)}])', text)
    
    # 合并内容+标点
    segments = []
    for i in range(0, len(parts) - 1, 2):
        segment = parts[i] + parts[i + 1]
        segments.append(segment)
    if len(parts) % 2 == 1 and parts[-1]:
        segments.append(parts[-1])
    
    # 每个子段独立判断
    result = []
    for seg in segments:
        if is_last_level:
            # 末级：直接拆分到这级，不再判断阈值
            result.append(seg)
        elif len(seg) <= max_len:
            # 未超阈值 → 语义单元，保留
            result.append(seg)
        else:
            # 超过阈值 → 递归下钻
            result.extend(_recursive_split_with_thresholds(seg, thresholds, level + 1))
    
    return result


def _longest_increasing_subsequence(indices):
    """
    求 indices 的最长递增子序列，返回 LIS 中元素在 indices 中的位置集合。
    用于句子匹配后检测顺序变化：LIS 内的配对顺序一致，LIS 外的为"移动"。
    """
    if not indices:
        return set()
    
    n = len(indices)
    # dp[i]: 以 indices[i] 结尾的 LIS 长度
    dp = [1] * n
    # prev[i]: 前驱索引，用于回溯路径
    prev = [-1] * n
    
    for i in range(1, n):
        for j in range(i):
            if indices[j] < indices[i] and dp[j] + 1 > dp[i]:
                dp[i] = dp[j] + 1
                prev[i] = j
    
    # 找 LIS 末端
    max_len = max(dp)
    end = dp.index(max_len)
    
    # 回溯 LIS 路径
    lis_positions = set()
    pos = end
    while pos != -1:
        lis_positions.add(pos)
        pos = prev[pos]
    
    return lis_positions


def _find_boundary(boundaries, pos):
    """根据位置找到对应的终稿段落索引"""
    for b_start, b_end, sf in boundaries:
        if b_start <= pos < b_end:
            return sf
    # 如果 pos 恰好落在边界末端，归入前一个段落
    # 如果 pos == 0 或超出范围，归入最后一个段落或第一个段落
    if boundaries:
        if pos >= boundaries[-1][1]:
            return boundaries[-1][2]
        return boundaries[0][2]
    return None


def _split_diff_to_boundaries(para_diffs, boundaries, tag, text, start_pos):
    """
    将一个 diff 片段按终稿段落边界拆分，分配到对应的终稿段落。
    
    Args:
        para_diffs: {f_idx: [(tag, text), ...]} 分配结果
        boundaries: [(start, end, f_idx), ...] 终稿段落边界表
        tag: diff 标记 ('equal' 或 'insert')
        text: diff 文本
        start_pos: 该片段在 combined_final_text 中的起始位置
    """
    end_pos = start_pos + len(text)
    text_offset = 0
    
    for b_start, b_end, sf in boundaries:
        # 跳过不重叠的段落
        if end_pos <= b_start:
            break
        if start_pos >= b_end:
            continue
        
        # 计算交集在 text 中的偏移
        chunk_start = max(start_pos, b_start) - start_pos
        chunk_end = min(end_pos, b_end) - start_pos
        
        chunk = text[chunk_start:chunk_end]
        if chunk:
            para_diffs[sf].append((tag, chunk))


def _build_boundaries(items, key_fn):
    """
    构建边界表：[(start, end, idx), ...]
    
    Args:
        items: 有序列表
        key_fn: 从 item 提取文本的函数
    
    Returns:
        boundaries: [(start, end, idx), ...] 按 items 顺序排列
    """
    boundaries = []
    cursor = 0
    for item in items:
        text = key_fn(item)
        text_len = len(text)
        boundaries.append((cursor, cursor + text_len, item))
        cursor += text_len
    return boundaries


def _distribute_diff_to_boundaries(diffs, boundaries):
    """
    将字符级 diff 结果按边界表分配到对应的段落/句子。
    
    公共逻辑：遍历 diffs，对 delete 按当前 f_pos 归入对应边界，
    对 equal/insert 按边界切割后分配。
    
    Args:
        diffs: char_level_diff 返回的 [(tag, text), ...]
        boundaries: [(start, end, idx), ...] 边界表
    
    Returns:
        {idx: [(tag, text), ...]} 分配结果
    """
    idxs = [b[2] for b in boundaries]
    para_diffs = {idx: [] for idx in idxs}
    
    f_pos = 0
    for tag, text in diffs:
        if not text:
            continue
        text_len = len(text)
        
        if tag == 'delete':
            target = _find_boundary(boundaries, f_pos)
            if target is not None:
                para_diffs[target].append((tag, text))
        elif tag in ('equal', 'insert'):
            _split_diff_to_boundaries(para_diffs, boundaries, tag, text, f_pos)
            f_pos += text_len
    
    return para_diffs


def _apply_diff_run(result_para, diffs):
    """
    将 char_level_diff 的结果写入结果段落。
    连续相同 tag 的文本会合并为一个 run，减少 Word 中的 run 碎片。
    """
    for tag, text in diffs:
        if not text:
            continue
        if tag == 'equal':
            result_para.add_run(text)
        elif tag == 'delete':
            run = result_para.add_run(text)
            run.font.color.rgb = RGBColor(0, 0, 255)
            run.font.strike = True
        elif tag == 'insert':
            run = result_para.add_run(text)
            run.font.color.rgb = RGBColor(255, 0, 0)


def _merge_unmatched_sentences(orig_sentences, final_sentences,
                                used_orig, used_final, sentence_match,
                                SENTENCE_SIM_THRESHOLD,
                                merge_groups, group_used_orig, group_used_final):
    """
    兜底合并：贪心匹配 + 现有合并/拆分检测后仍有未匹配句时，
    尝试与相邻已匹配句合并，兜底处理拆分碎片。

    例如 "2" 和 "0！最终..." 被标点拆分后各自独立，但合并后
    "2:0！最终..." 能与 "2:0。" 形成更好的匹配。

    Args:
        orig_sentences, final_sentences: 分句结果
        used_orig, used_final: 当前已匹配/已分组的索引集合
        sentence_match: 1:1 匹配结果 [(o_idx, f_idx, ratio), ...]
        merge_groups: 已有的合并组，兜底结果追加至此（in-place）
        group_used_orig, group_used_final: 已分组的索引集（in-place）
    """
    IMPROVEMENT_THRESHOLD = 0.15  # 合并后 ratio 提升超过此值才生效

    unmatched_orig = sorted(set(range(len(orig_sentences))) - used_orig)
    if not unmatched_orig:
        return

    # 收集已被合并组占用的 final 索引（避免冲突）
    occupied_final = set()
    for _, f_indices, _ in merge_groups:
        occupied_final.update(f_indices)

    for oi in unmatched_orig:
        # 找 oi 前面最近的已匹配 orig 句
        prev_candidates = [(m[0], m[1]) for m in sentence_match if m[0] < oi]
        if not prev_candidates:
            continue
        prev_oi, prev_fi = max(prev_candidates, key=lambda x: x[0])

        # 如果 prev_fi 已被合并组占用，跳过
        if prev_fi in occupied_final:
            continue

        # 合并后 vs 同一 final 句的 ratio
        merged_text = orig_sentences[prev_oi] + orig_sentences[oi]
        merged_ratio = difflib.SequenceMatcher(
            None, merged_text, final_sentences[prev_fi]
        ).ratio()

        # 当前最佳 solo ratio：prev 的匹配率 + oi 单独匹配的最佳值
        prev_ratio = next((r for o, f, r in sentence_match if o == prev_oi), 0)
        oi_best = max(
            difflib.SequenceMatcher(
                None, orig_sentences[oi], final_sentences[fi]
            ).ratio()
            for fi in range(len(final_sentences))
        )
        best_solo = max(prev_ratio, oi_best)

        if merged_ratio - best_solo > IMPROVEMENT_THRESHOLD:
            merge_groups.append(([prev_oi, oi], [prev_fi], merged_ratio))
            group_used_orig.add(oi)


def _detect_sentence_merge_split(orig_sentences, final_sentences, used_orig, used_final, SENTENCE_SIM_THRESHOLD, sentence_match=None):
    """
    句子级拆分/合并检测。
    
    在贪心1:1匹配之后，对仍未匹配的句子尝试分组匹配：
    - 合并：N个连续原稿句子 → 1个终稿句子
    - 拆分：1个原稿句子 → N个连续终稿句子
    - 合并升级：已1:1匹配的连续句子对，拼接后ratio更高时升级为合并组
    
    Args:
        orig_sentences: 原稿句子列表
        final_sentences: 终稿句子列表
        used_orig: 已在1:1匹配中占用的原稿句子索引集合
        used_final: 已在1:1匹配中占用的终稿句子索引集合
        SENTENCE_SIM_THRESHOLD: 句子相似度阈值
        sentence_match: 1:1匹配列表 [(orig_idx, final_idx, ratio), ...]，用于合并升级检测
    
    Returns:
        merge_groups: [([o_idx1, o_idx2, ...], f_idx, ratio), ...]  合并组
        split_groups: [(o_idx, [f_idx1, f_idx2, ...], ratio), ...]  拆分组
        new_used_orig: 被分组匹配占用的原稿句子索引集合
        new_used_final: 被分组匹配占用的终稿句子索引集合
        replaced_match_indices: 被合并组替代的1:1匹配在sentence_match中的索引集合
    """
    MAX_GROUP_SIZE = 3  # 最多合并/拆分3句
    SEARCH_WINDOW = 5    # 合并/拆分检测时仅搜索相邻±5句范围内的候选，避免O(n²)遍历
    
    unmatched_orig = sorted(i for i in range(len(orig_sentences)) if i not in used_orig)
    unmatched_final = sorted(i for i in range(len(final_sentences)) if i not in used_final)
    
    all_candidates = []  # (ratio, 'merge'/'split'/'merge_upgrade', group_info)
    
    # ---- A. 合并检测（纯未匹配）：N原稿句 → 1终稿句 ----
    if unmatched_orig and unmatched_final:
        for f_idx in unmatched_final:
            f_text = final_sentences[f_idx]
            if not f_text.strip():
                continue
            
            # 窗口限制：只搜索原稿中 f_idx±SEARCH_WINDOW 范围内的未匹配句
            window_start = max(0, f_idx - SEARCH_WINDOW)
            window_end = min(len(orig_sentences), f_idx + SEARCH_WINDOW + 1)
            window_orig = [oi for oi in unmatched_orig if window_start <= oi < window_end]
            
            for seg_start in range(len(window_orig)):
                for group_size in range(2, MAX_GROUP_SIZE + 1):
                    seg_end = seg_start + group_size
                    if seg_end > len(window_orig):
                        break
                    
                    o_indices = window_orig[seg_start:seg_end]
                    
                    # 连续性检查：索引必须相邻
                    if o_indices != list(range(o_indices[0], o_indices[0] + group_size)):
                        break
                    
                    combined_orig = "".join(orig_sentences[oi] for oi in o_indices)
                    if not combined_orig.strip():
                        continue
                    
                    ratio = difflib.SequenceMatcher(None, combined_orig, f_text).ratio()
                    
                    if ratio >= SENTENCE_SIM_THRESHOLD:
                        best_single_ratio = 0
                        for oi in o_indices:
                            single_ratio = difflib.SequenceMatcher(
                                None, orig_sentences[oi], f_text
                            ).ratio()
                            best_single_ratio = max(best_single_ratio, single_ratio)
                        
                        if ratio > best_single_ratio:
                            all_candidates.append((ratio, 'merge', (o_indices, f_idx)))
    
    # ---- B. 拆分检测（纯未匹配）：1原稿句 → N终稿句 ----
    if unmatched_orig and unmatched_final:
        for o_idx in unmatched_orig:
            o_text = orig_sentences[o_idx]
            if not o_text.strip():
                continue
            
            # 窗口限制：只搜索终稿中 o_idx±SEARCH_WINDOW 范围内的未匹配句
            window_start = max(0, o_idx - SEARCH_WINDOW)
            window_end = min(len(final_sentences), o_idx + SEARCH_WINDOW + 1)
            window_final = [fi for fi in unmatched_final if window_start <= fi < window_end]
            
            for seg_start in range(len(window_final)):
                for group_size in range(2, MAX_GROUP_SIZE + 1):
                    seg_end = seg_start + group_size
                    if seg_end > len(window_final):
                        break
                    
                    f_indices = window_final[seg_start:seg_end]
                    
                    # 连续性检查
                    if f_indices != list(range(f_indices[0], f_indices[0] + group_size)):
                        break
                    
                    combined_final = "".join(final_sentences[fi] for fi in f_indices)
                    if not combined_final.strip():
                        continue
                    
                    ratio = difflib.SequenceMatcher(None, o_text, combined_final).ratio()
                    
                    if ratio >= SENTENCE_SIM_THRESHOLD:
                        best_single_ratio = 0
                        for fi in f_indices:
                            single_ratio = difflib.SequenceMatcher(
                                None, o_text, final_sentences[fi]
                            ).ratio()
                            best_single_ratio = max(best_single_ratio, single_ratio)
                        
                        if ratio > best_single_ratio:
                            all_candidates.append((ratio, 'split', (o_idx, f_indices)))
    
    # ---- C. 合并升级检测：已1:1匹配的连续句子对，拼接后更优时升级 ----
    # 场景：原稿 "第一条 XX。第二条 YY。" → 终稿 "第一条 XX，YY。"
    # 1:1匹配: orig[0]↔final[0](0.91), orig[1]↔final[1](0.75)
    # 合并升级: orig[0]+orig[1] ↔ final[0]+final[1] 拼接后ratio更高
    if sentence_match and len(sentence_match) >= 2:
        # 按orig索引排序
        sorted_matches = sorted(sentence_match, key=lambda m: m[0])
        
        # 检查连续的1:1匹配对
        for i in range(len(sorted_matches)):
            for group_size in range(2, MAX_GROUP_SIZE + 1):
                end_i = i + group_size
                if end_i > len(sorted_matches):
                    break
                
                consecutive_group = sorted_matches[i:end_i]
                
                # 检查orig索引连续
                orig_indices = [m[0] for m in consecutive_group]
                if orig_indices != list(range(orig_indices[0], orig_indices[0] + group_size)):
                    break
                
                # 检查final索引连续
                final_indices = [m[1] for m in consecutive_group]
                if final_indices != list(range(final_indices[0], final_indices[0] + group_size)):
                    continue
                
                # 拼接原稿和终稿
                combined_orig = "".join(orig_sentences[oi] for oi in orig_indices)
                combined_final = "".join(final_sentences[fi] for fi in final_indices)
                
                if not combined_orig.strip() or not combined_final.strip():
                    continue
                
                combined_ratio = difflib.SequenceMatcher(None, combined_orig, combined_final).ratio()
                
                # 1:1匹配的平均ratio
                avg_single_ratio = sum(m[2] for m in consecutive_group) / len(consecutive_group)
                
                # 合并后ratio必须高于1:1匹配的平均ratio，且高于阈值
                if combined_ratio > avg_single_ratio and combined_ratio >= SENTENCE_SIM_THRESHOLD:
                    replaced_indices = set()
                    for m_idx, m in enumerate(sentence_match):
                        if m[0] in orig_indices:
                            replaced_indices.add(m_idx)
                    
                    all_candidates.append((
                        combined_ratio, 'merge_upgrade',
                        (orig_indices, final_indices, replaced_indices)
                    ))
    
    # ---- D. 匹配扩展检测：已1:1匹配 + 相邻未匹配原稿句子 → 合并 ----
    # 场景：原稿 "加强领导。明确责任。落实到位。" → 终稿 "加强领导、明确责任、落实到位。"
    # 1:1匹配: orig[0]↔final[0](0.5)，orig[1]和orig[2]未匹配
    # 扩展：orig[0]+orig[1]+orig[2] ↔ final[0] 拼接后ratio更高
    if sentence_match and unmatched_orig:
        # 构建 1:1匹配查找表
        match_by_orig = {m[0]: (m[1], m[2]) for m in sentence_match}
        match_by_final = {m[1]: (m[0], m[2]) for m in sentence_match}
        
        for o_idx, f_idx, single_ratio in sentence_match:
            # 向前/向后扫描相邻的未匹配原稿句子
            for expand_size in range(2, MAX_GROUP_SIZE + 1):
                # 尝试向两个方向扩展
                for direction in ['forward', 'backward', 'both']:
                    if direction == 'forward':
                        # 向后扩展：o_idx, o_idx+1, ..., o_idx+expand_size-1
                        expanded_orig = list(range(o_idx, o_idx + expand_size))
                    elif direction == 'backward':
                        # 向前扩展：o_idx-expand_size+1, ..., o_idx
                        expanded_orig = list(range(o_idx - expand_size + 1, o_idx + 1))
                    else:
                        # 双向扩展
                        half = (expand_size - 1) // 2
                        expanded_orig = list(range(o_idx - half, o_idx + expand_size - half))
                    
                    # 边界检查
                    if expanded_orig[0] < 0 or expanded_orig[-1] >= len(orig_sentences):
                        continue
                    
                    # 检查扩展部分是否都是未匹配的（除了原始o_idx）
                    expand_part = [oi for oi in expanded_orig if oi != o_idx]
                    if any(oi in used_orig for oi in expand_part):
                        continue
                    
                    # 连续性检查（已经是连续的 by construction）
                    
                    # 终稿句子：用已匹配的f_idx（扩展后整个组对应同一个终稿句子）
                    combined_orig = "".join(orig_sentences[oi] for oi in expanded_orig)
                    f_text = final_sentences[f_idx]
                    
                    if not combined_orig.strip():
                        continue
                    
                    combined_ratio = difflib.SequenceMatcher(None, combined_orig, f_text).ratio()
                    
                    if combined_ratio > single_ratio and combined_ratio >= SENTENCE_SIM_THRESHOLD:
                        # 记录被替代的1:1匹配
                        replaced_indices = set()
                        for m_idx, m in enumerate(sentence_match):
                            if m[0] in expanded_orig:
                                replaced_indices.add(m_idx)
                        
                        all_candidates.append((
                            combined_ratio, 'merge_upgrade',
                            (expanded_orig, [f_idx], replaced_indices)
                        ))
                        break  # 一个方向找到一个就够
    
    # ---- E. 拆分升级检测：已1:1匹配 + 相邻未匹配终稿句子 → 拆分 ----
    # 1:1匹配: orig[0]↔final[0](0.89)，final[1]未匹配
    # 拆分升级: orig[0] ↔ final[0]+final[1] 拼接后ratio更高
    if sentence_match and unmatched_final:
        for o_idx, f_idx, single_ratio in sentence_match:
            # 向前/向后扫描相邻的未匹配终稿句子
            for expand_size in range(2, MAX_GROUP_SIZE + 1):
                found = False
                for direction in ['forward', 'backward', 'both']:
                    if direction == 'forward':
                        # 向后扩展：f_idx, f_idx+1, ..., f_idx+expand_size-1
                        expanded_final = list(range(f_idx, f_idx + expand_size))
                    elif direction == 'backward':
                        # 向前扩展：f_idx-expand_size+1, ..., f_idx
                        expanded_final = list(range(f_idx - expand_size + 1, f_idx + 1))
                    else:
                        # 双向扩展
                        half = (expand_size - 1) // 2
                        expanded_final = list(range(f_idx - half, f_idx + expand_size - half))
                    
                    # 边界检查
                    if expanded_final[0] < 0 or expanded_final[-1] >= len(final_sentences):
                        continue
                    
                    # 检查扩展部分是否都是未匹配的（除了原始f_idx）
                    expand_part = [fi for fi in expanded_final if fi != f_idx]
                    if any(fi in used_final for fi in expand_part):
                        continue
                    
                    # 拼接终稿句子
                    combined_final = "".join(final_sentences[fi] for fi in expanded_final)
                    o_text = orig_sentences[o_idx]
                    
                    if not combined_final.strip():
                        continue
                    
                    combined_ratio = difflib.SequenceMatcher(None, o_text, combined_final).ratio()
                    
                    if combined_ratio > single_ratio and combined_ratio >= SENTENCE_SIM_THRESHOLD:
                        # 记录被替代的1:1匹配
                        replaced_indices = set()
                        for m_idx, m in enumerate(sentence_match):
                            if m[0] == o_idx:
                                replaced_indices.add(m_idx)
                        
                        all_candidates.append((
                            combined_ratio, 'split_upgrade',
                            (o_idx, expanded_final, replaced_indices)
                        ))
                        found = True
                        break  # 一个方向找到一个就够
                if found:
                    break  # 一个expand_size找到一个就够
    
    # ---- 贪心分配 ----
    all_candidates.sort(key=lambda x: x[0], reverse=True)
    
    allocated_orig = set()
    allocated_final = set()
    merge_groups = []
    split_groups = []
    replaced_match_indices = set()
    
    for ratio, cand_type, group_info in all_candidates:
        if cand_type == 'merge':
            o_indices, f_idx = group_info
            if any(oi in allocated_orig for oi in o_indices):
                continue
            if f_idx in allocated_final:
                continue
            for oi in o_indices:
                allocated_orig.add(oi)
            allocated_final.add(f_idx)
            merge_groups.append((o_indices, f_idx, ratio))
        
        elif cand_type == 'split':
            o_idx, f_indices = group_info
            if o_idx in allocated_orig:
                continue
            if any(fi in allocated_final for fi in f_indices):
                continue
            allocated_orig.add(o_idx)
            for fi in f_indices:
                allocated_final.add(fi)
            split_groups.append((o_idx, f_indices, ratio))
        
        elif cand_type == 'merge_upgrade':
            orig_indices, final_indices, replaced_indices = group_info
            if any(oi in allocated_orig for oi in orig_indices):
                continue
            if any(fi in allocated_final for fi in final_indices):
                continue
            # 确认分配
            for oi in orig_indices:
                allocated_orig.add(oi)
            for fi in final_indices:
                allocated_final.add(fi)
            # merge_upgrade 记录为合并组（所有终稿句子作为一组）
            merge_groups.append((orig_indices, final_indices, ratio))
            replaced_match_indices.update(replaced_indices)
        
        elif cand_type == 'split_upgrade':
            o_idx, f_indices, replaced_indices = group_info
            if o_idx in allocated_orig:
                continue
            if any(fi in allocated_final for fi in f_indices):
                continue
            allocated_orig.add(o_idx)
            for fi in f_indices:
                allocated_final.add(fi)
            # split_upgrade 记录为拆分组
            split_groups.append((o_idx, f_indices, ratio))
            replaced_match_indices.update(replaced_indices)
    
    return merge_groups, split_groups, allocated_orig, allocated_final, replaced_match_indices


def _output_sentence_split_diff(result_para, orig_text, final_sentences_map, f_indices):
    """
    句子级拆分组输出：1原稿句子 → N终稿句子（同一段落内）。
    
    将原稿句子文本与合并后的终稿句子文本做字符级 diff，
    然后按终稿句子边界切割 diff 结果，分别输出到 result_para 的不同 run 区域。
    所有输出在同一 result_para 内，不创建新段落。
    
    Args:
        result_para: 结果段落对象
        orig_text: 原稿句子文本
        final_sentences_map: {f_idx: sentence_text} 终稿句子映射
        f_indices: 终稿句子索引列表（按终稿顺序）
    """
    combined_final_text = "".join(final_sentences_map[fi] for fi in f_indices)
    diffs = char_level_diff(orig_text, combined_final_text)
    
    boundaries = _build_boundaries(
        f_indices, lambda fi: final_sentences_map[fi]
    )
    para_diffs = _distribute_diff_to_boundaries(diffs, boundaries)
    
    for fi in f_indices:
        sf_diffs = para_diffs[fi]
        
        if not sf_diffs:
            result_para.add_run(final_sentences_map[fi])
        else:
            all_equal = all(t == 'equal' for t, _ in sf_diffs)
            if all_equal:
                for t, txt in sf_diffs:
                    result_para.add_run(txt)
            else:
                _apply_diff_run(result_para, sf_diffs)


def sentence_level_diff(orig_text, final_text, result_para, SENTENCE_SIM_THRESHOLD,
                         short_para_char_threshold=0, semantic_unit_thresholds=None):
    """
    句子级别差异比较（方案 B：Diff 驱动的可靠边界分句）

    核心改进（方案 B）：
    - 使用 split_by_reliable_boundaries 替代 split_into_sentences
    - 先做全文 char_diff，从 equal 区域提取可靠的标点边界
    - 避免了"标点本身被编辑修改"导致的分句不一致问题

    其他逻辑保持不变：
    - 贪心1:1匹配 + 句子级拆分/合并检测
    - 使用逆序对检测句子互换
    - 低于相似度阈值的配对直接标记新增+删除

    Args:
        short_para_char_threshold: >0 时，若原文+终稿总字符数不超过此值，
            直接走字符级diff，跳过句子拆分
        semantic_unit_thresholds: 多级语义单元拆分配置，传给 split_by_reliable_boundaries
    """
    # 短段落直接字符级diff，避免句子拆分产生不合理的碎片
    if short_para_char_threshold > 0 and len(orig_text) + len(final_text) <= short_para_char_threshold:
        para_diff = char_level_diff(orig_text, final_text)
        _apply_diff_run(result_para, para_diff)
        return

    # ---- 分句：方案 B（Diff 驱动） 或 方案 A（直接标点） ----
    if ENABLE_DIFF_DRIVEN_SPLIT:
        # 方案 B：先做全文 char_diff，从 equal 区域提取可靠的标点边界
        orig_sentences, final_sentences = split_by_reliable_boundaries(
            orig_text, final_text, semantic_unit_thresholds
        )
    else:
        # 方案 A-1：直接按标点层级分句
        orig_sentences = split_into_sentences(orig_text, semantic_unit_thresholds)
        final_sentences = split_into_sentences(final_text, semantic_unit_thresholds)

    # 只有一个句子时，直接字符级比对
    if len(orig_sentences) == 1 and len(final_sentences) == 1:
        para_diff = char_level_diff(orig_text, final_text)
        _apply_diff_run(result_para, para_diff)
        return
    
    # ---- 步骤1: 贪心1:1匹配（按置信度降序分配）----
    # 收集所有候选对，按 ratio 降序排序后分配（高置信度优先）
    all_pairs = []  # (ratio, o_idx, f_idx)
    for f_idx, f_sent in enumerate(final_sentences):
        for o_idx, o_sent in enumerate(orig_sentences):
            ratio = difflib.SequenceMatcher(None, o_sent, f_sent).ratio()
            if ratio >= SENTENCE_SIM_THRESHOLD:
                all_pairs.append((ratio, o_idx, f_idx))
    
    # 按 ratio 降序，高置信度优先分配；ratio 相同时保持原稿顺序稳定
    all_pairs.sort(key=lambda x: (-x[0], x[1]))
    
    used_orig = set()
    used_final = set()
    sentence_match = []
    
    for ratio, o_idx, f_idx in all_pairs:
        if o_idx in used_orig or f_idx in used_final:
            continue
        sentence_match.append((o_idx, f_idx, ratio))
        used_orig.add(o_idx)
        used_final.add(f_idx)
    
    # ---- 步骤2: 句子级拆分/合并检测 ----
    merge_groups, split_groups, group_used_orig, group_used_final, replaced_match_indices = \
        _detect_sentence_merge_split(
            orig_sentences, final_sentences, used_orig, used_final, 
            SENTENCE_SIM_THRESHOLD, sentence_match
        )
    
    # 移除被合并升级替代的1:1匹配
    if replaced_match_indices:
        sentence_match = [m for i, m in enumerate(sentence_match) if i not in replaced_match_indices]
        # 重建 used_orig / used_final
        used_orig = set(m[0] for m in sentence_match) | group_used_orig
        used_final = set(m[1] for m in sentence_match) | group_used_final
    
    # 将分组占用的索引加入已用集合
    used_orig.update(group_used_orig)
    used_final.update(group_used_final)

    # ---- 步骤2.5: 未匹配句相邻合并兜底 ----
    # 贪心匹配 + 现有合并/拆分检测后仍有未匹配句时，
    # 尝试与相邻已匹配句合并，兜底处理拆分碎片
    _merge_unmatched_sentences(
        orig_sentences, final_sentences,
        used_orig, used_final, sentence_match,
        SENTENCE_SIM_THRESHOLD,
        merge_groups, group_used_orig, group_used_final
    )
    # 将兜底合并加入已用集合
    used_orig.update(group_used_orig)
    used_final.update(group_used_final)

    # 构建查找表
    # merge_map: frozenset(f_indices) -> (o_indices, f_indices, ratio)  合并组（支持多终稿句）
    merge_map = {}  # f_idx -> (o_indices, f_indices, ratio)
    merge_final_set = set()  # 被合并组占用的终稿句子索引
    for o_indices, f_indices, ratio in merge_groups:
        # 将合并组信息关联到每个参与的终稿句子
        group_key = (tuple(o_indices), tuple(f_indices), ratio)
        for fi in f_indices:
            merge_map[fi] = group_key
        merge_final_set.update(f_indices)
    
    # split_map: o_idx -> (f_indices, ratio)  拆分组
    split_map = {}
    for o_idx, f_indices, ratio in split_groups:
        split_map[o_idx] = (f_indices, ratio)
    
    # split_final_set: 被拆分组占用的终稿句子索引
    split_final_set = set()
    for o_idx, f_indices, ratio in split_groups:
        split_final_set.update(f_indices)
    
    # ---- 步骤3: 逆序对检测句子顺序变化 ----
    # 扩展：1:1匹配 + 合并组 + 拆分组都参与逆序对检测
    # 统一表示为 (orig_key, final_key) 用于排序比较
    # 合并组: orig_key = max(o_indices), final_key = max(f_indices)
    # 拆分组: orig_key = o_idx, final_key = max(f_indices)
    
    all_pairs = []  # [(orig_key, final_key), ...] 用于逆序对检测
    
    for o_idx, f_idx, ratio in sentence_match:
        all_pairs.append((o_idx, f_idx))
    
    for o_indices, f_indices, ratio in merge_groups:
        all_pairs.append((max(o_indices), max(f_indices)))
    
    for o_idx, f_indices, ratio in split_groups:
        all_pairs.append((o_idx, max(f_indices)))
    
    moved_orig_set = set()       # 向后移的句子索引（原位蓝删 + 新位红增）
    forward_moved_set = set()    # 向前移的句子索引（终稿位直接显示，不做diff）
    
    all_pairs_sorted = sorted(all_pairs, key=lambda p: p[0])
    for i in range(len(all_pairs_sorted)):
        for j in range(i + 1, len(all_pairs_sorted)):
            if all_pairs_sorted[i][1] > all_pairs_sorted[j][1]:
                # 交叉：orig_i 在前但 final_i 在后
                moved_orig_set.add(all_pairs_sorted[i][0])
                forward_moved_set.add(all_pairs_sorted[j][0])
    
    # ---- 步骤4: 按终稿句子顺序输出 ----
    
    # 构建: final_idx -> (orig_idx, ratio)  1:1匹配
    final_to_match = {}
    for o_idx, f_idx, ratio in sentence_match:
        final_to_match[f_idx] = (o_idx, ratio)
    
    # 构建: orig_idx -> final_idx（用于判断移动句子在终稿中的位置）
    orig_to_final = {}
    for o_idx, f_idx, ratio in sentence_match:
        orig_to_final[o_idx] = f_idx
    
    # 已处理的原稿句子索引
    processed_orig = set()
    # 原稿位置追踪指针
    next_orig = 0
    
    # 预计算：每个原稿句子索引之后（含自身）最近的已配对原稿句子索引
    # 包括1:1匹配 + 分组匹配
    all_matched_orig_indices = set(orig_to_final.keys()) | group_used_orig
    next_matched_orig_after = [len(orig_sentences)] * len(orig_sentences)
    latest_matched = len(orig_sentences)
    for i in range(len(orig_sentences) - 1, -1, -1):
        if i in all_matched_orig_indices:
            latest_matched = i
        next_matched_orig_after[i] = latest_matched
    
    for f_idx in range(len(final_sentences)):
        f_sent = final_sentences[f_idx]
        
        # 确定当前终稿句子匹配的原稿位置（用于推进指针）
        if f_idx in final_to_match:
            current_orig_idx = final_to_match[f_idx][0]
        elif f_idx in merge_map:
            current_orig_idx = min(merge_map[f_idx][0])
        else:
            # 未匹配终稿句子：只推进到下一个已配对原稿句子之前
            if next_orig < len(orig_sentences):
                current_orig_idx = next_matched_orig_after[next_orig]
            else:
                current_orig_idx = next_orig
        
        # 输出原稿中"本应在此位置但已移走或已删除"的句子（蓝色删除）
        while next_orig < current_orig_idx:
            if next_orig not in processed_orig:
                if next_orig in moved_orig_set:
                    run = result_para.add_run(orig_sentences[next_orig])
                    run.font.color.rgb = RGBColor(0, 0, 255)
                    run.font.strike = True
                elif next_orig not in orig_to_final and next_orig not in split_map:
                    # 未匹配且非拆分组原稿句子（纯删除）→ 蓝色删除
                    run = result_para.add_run(orig_sentences[next_orig])
                    run.font.color.rgb = RGBColor(0, 0, 255)
                    run.font.strike = True
            processed_orig.add(next_orig)
            next_orig += 1
        
        # ---- 输出当前终稿句子 ----
        
        # 跳过被合并组占用的非首句终稿句子（会在首句时一起处理）
        if f_idx in merge_final_set:
            group_key = merge_map[f_idx]
            o_indices_tuple, f_indices_tuple, ratio = group_key
            # 只在首句时触发合并组输出
            if f_indices_tuple[0] == f_idx:
                # ---- 合并组输出：N原稿句 → M终稿句 ----
                o_indices = list(o_indices_tuple)
                f_indices = list(f_indices_tuple)
                combined_orig = "".join(orig_sentences[oi] for oi in o_indices)
                
                if max(o_indices) in moved_orig_set:
                    # 合并组向后移 → 终稿新位各句红色新增
                    for fi in f_indices:
                        run = result_para.add_run(final_sentences[fi])
                        run.font.color.rgb = RGBColor(255, 0, 0)
                elif min(o_indices) in forward_moved_set:
                    # 合并组向前移 → 直接显示终稿文本
                    for fi in f_indices:
                        result_para.add_run(final_sentences[fi])
                else:
                    # 顺序一致：做字符级diff，按终稿句子边界切割
                    final_sentences_map = {fi: final_sentences[fi] for fi in f_indices}
                    _output_sentence_split_diff(
                        result_para, combined_orig, final_sentences_map, f_indices
                    )
                
                for oi in o_indices:
                    processed_orig.add(oi)
                    if oi >= next_orig:
                        next_orig = oi + 1
            # 非首句，跳过（已在首句时输出）
            continue
        
        # 跳过被拆分组占用的终稿句子（会在拆分组首句时一起处理）
        if f_idx in split_final_set:
            is_split_head = False
            for o_idx, (f_indices, ratio) in split_map.items():
                if f_indices[0] == f_idx and o_idx not in processed_orig:
                    is_split_head = True
                    combined_orig = orig_sentences[o_idx]
                    final_sentences_map = {fi: final_sentences[fi] for fi in f_indices}
                    _output_sentence_split_diff(
                        result_para, combined_orig, final_sentences_map, f_indices
                    )
                    processed_orig.add(o_idx)
                    if o_idx >= next_orig:
                        next_orig = o_idx + 1
                    break
            
            if is_split_head:
                continue
            else:
                continue
        
        if f_idx not in final_to_match:
            # 未匹配的终稿句子 → 红色新增
            run = result_para.add_run(f_sent)
            run.font.color.rgb = RGBColor(255, 0, 0)
        else:
            # ---- 1:1匹配输出 ----
            o_idx, ratio = final_to_match[f_idx]
            o_sent = orig_sentences[o_idx]
            processed_orig.add(o_idx)
            if o_idx >= next_orig:
                next_orig = o_idx + 1
            
            if ratio < SENTENCE_SIM_THRESHOLD:
                # 相似度低 → 直接标记新增+删除
                run1 = result_para.add_run(f_sent)
                run1.font.color.rgb = RGBColor(255, 0, 0)
                run2 = result_para.add_run(o_sent)
                run2.font.color.rgb = RGBColor(0, 0, 255)
                run2.font.strike = True
            elif o_idx in moved_orig_set:
                # 向后移动的句子 → 在终稿新位输出红色新增
                run = result_para.add_run(f_sent)
                run.font.color.rgb = RGBColor(255, 0, 0)
            elif o_idx in forward_moved_set:
                # 向前移动的句子 → 直接显示终稿文本，不做字符级diff
                result_para.add_run(f_sent)
            else:
                # 顺序一致，走字符级 diff
                para_diff = char_level_diff(o_sent, f_sent)
                _apply_diff_run(result_para, para_diff)
    
    # 输出剩余未处理的原稿句子（移动到更后位置 或 纯删除）
    while next_orig < len(orig_sentences):
        if next_orig not in processed_orig:
            run = result_para.add_run(orig_sentences[next_orig])
            run.font.color.rgb = RGBColor(0, 0, 255)
            run.font.strike = True
        processed_orig.add(next_orig)
        next_orig += 1


# ======================== 原地标记辅助函数 ========================

def _get_para_font_info(paragraph):
    """从段落提取字体信息（字体名、东亚字体、字号、加粗）
    作为该段落新增run的默认字体
    优先取最后一个run的有效字体名，若没有则向前回溯
    """
    font_info = {'name': None, 'eastAsia': None, 'size': None, 'bold': None}

    if paragraph.runs:
        # 从最后一个run向前回溯，找到第一个有字体名信息的run
        for run in reversed(paragraph.runs):
            has_name = False
            try:
                rPr = run.font.element.rPr
                if rPr is not None:
                    rFonts = rPr.find(qn('w:rFonts'))
                    if rFonts is not None:
                        ea = rFonts.get(qn('w:eastAsia'))
                        if ea:
                            font_info['eastAsia'] = ea
                            font_info['name'] = ea
                            has_name = True
                        ascii_f = rFonts.get(qn('w:ascii'))
                        if ascii_f and not font_info.get('name'):
                            font_info['name'] = ascii_f
                            has_name = True
            except Exception:
                pass

            fn = run.font.name
            if fn:
                font_info['name'] = fn
                has_name = True

            # 只有找到字体名才回溯继续搜索
            if has_name:
                if not font_info.get('size'):
                    font_info['size'] = run.font.size
                if font_info.get('bold') is None:
                    font_info['bold'] = run.font.bold
                break

        # 如果整个段落都没有字体名，至少收集字号
        if not font_info.get('name'):
            for run in paragraph.runs:
                if not font_info.get('size'):
                    font_info['size'] = run.font.size
                if font_info.get('bold') is None:
                    font_info['bold'] = run.font.bold
                if font_info.get('size') is not None:
                    break

    # 回退到段落样式的XML
    if not font_info.get('name'):
        try:
            style = paragraph.style
            if style:
                s_elem = style._element
                s_pPr = s_elem.find(qn('w:pPr'))
                if s_pPr is not None:
                    s_rPr = s_pPr.find(qn('w:rPr'))
                else:
                    s_rPr = s_elem.find(qn('w:rPr'))
                if s_rPr is not None:
                    s_rFonts = s_rPr.find(qn('w:rFonts'))
                    if s_rFonts is not None:
                        ea = s_rFonts.get(qn('w:eastAsia'))
                        if ea:
                            font_info['eastAsia'] = ea
                            font_info['name'] = ea
                        ascii_f = s_rFonts.get(qn('w:ascii'))
                        if ascii_f and not font_info.get('name'):
                            font_info['name'] = ascii_f
                    s_sz = s_rPr.find(qn('w:sz'))
                    if s_sz is not None and not font_info.get('size'):
                        sz_val = s_sz.get(qn('w:val'))
                        if sz_val:
                            font_info['size'] = Pt(int(sz_val) // 2)

                if not font_info.get('name') and style.font:
                    font_info['name'] = style.font.name
                    if not font_info.get('size'):
                        font_info['size'] = style.font.size
                    if not font_info.get('eastAsia') and style.font.name:
                        font_info['eastAsia'] = style.font.name
        except Exception:
            pass

    # 最终 fallback
    if not font_info.get('name') and not font_info.get('eastAsia'):
        font_info['name'] = '仿宋'
        font_info['eastAsia'] = '仿宋'
    if not font_info.get('size'):
        font_info['size'] = Pt(16)

    if font_info.get('name') and not font_info.get('eastAsia'):
        font_info['eastAsia'] = font_info['name']
    elif font_info.get('eastAsia') and not font_info.get('name'):
        font_info['name'] = font_info['eastAsia']

    return font_info


def _apply_run_font(run, font_info):
    """将字体信息应用到run对象（同时设置ascii和eastAsia字体）"""
    font_name = font_info.get('name')
    east_asia = font_info.get('eastAsia')

    try:
        if font_info.get('size'):
            run.font.size = font_info['size']
        if font_info.get('bold') is not None:
            run.font.bold = font_info['bold']

        if font_name or east_asia:
            rPr = run.element.rPr
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                run.element.append(rPr)

            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.append(rFonts)

            if font_name:
                rFonts.set(qn('w:ascii'), font_name)
                rFonts.set(qn('w:hAnsi'), font_name)
                run.font.name = font_name
            if east_asia:
                rFonts.set(qn('w:eastAsia'), east_asia)
            elif font_name:
                rFonts.set(qn('w:eastAsia'), font_name)
    except Exception:
        pass


def _is_pure_image_para(paragraph):
    """检查段落是否只包含图片（无文字内容），这类段落跳过标记规则"""
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    if paragraph._element.find('.//w:drawing', ns) is None:
        return False
    # 纯图片段落的特征：文本全空
    return not paragraph.text.strip()


def _clear_para_runs(paragraph):
    """清空段落的所有run和hyperlink元素，保留段落级格式
    
    注意：保留包含图片（<w:drawing>）的 run，避免误删图片。
    """
    p_elem = paragraph._element
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    for child in list(p_elem):
        tag = child.tag
        if tag.endswith('}r') or tag.endswith('}hyperlink'):
            # 保留包含图片的 run/hyperlink（不删除）
            if child.find('.//w:drawing', ns) is not None:
                continue
            p_elem.remove(child)


def _remove_element_safe(element):
    """安全地从 DOM 中移除段落元素。"""
    try:
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)
    except Exception:
        pass


def _set_para_red_with_deletion(result_para, orig_text, font_info):
    """将一个段落标记为已删除（蓝色+删除线），保留段落中的图片/表格等元素"""
    has_content = False
    for run in result_para.runs:
        has_content = True
        run.font.color.rgb = RGBColor(0, 0, 255)
        run.font.strike = True
        _apply_run_font(run, font_info)

    if not has_content and orig_text:
        run = result_para.add_run(orig_text)
        run.font.color.rgb = RGBColor(0, 0, 255)
        run.font.strike = True
        _apply_run_font(run, font_info)


def _create_new_para_elem(para_elem, text, font_info, color=None, is_after=True):
    """在para_elem之后或之前创建新段落XML元素，返回新元素"""
    if is_after:
        ref_elem = para_elem
    else:
        prev = para_elem.getprevious()
        ref_elem = prev if prev is not None else para_elem

    ref_pPr = ref_elem.find(qn('w:pPr'))

    temp_doc = Document()
    temp_p = temp_doc.add_paragraph(text)
    for run in temp_p.runs:
        _apply_run_font(run, font_info)
        if color:
            run.font.color.rgb = color

    new_elem = temp_p._element

    if ref_pPr is not None:
        existing_pPr = new_elem.find(qn('w:pPr'))
        if existing_pPr is not None:
            new_elem.remove(existing_pPr)
        copied_pPr = deepcopy(ref_pPr)
        new_elem.insert(0, copied_pPr)

    parent = para_elem.getparent()
    if is_after:
        next_sibling = para_elem.getnext()
        if next_sibling is not None:
            parent.insert(list(parent).index(next_sibling), new_elem)
        else:
            parent.append(new_elem)
    else:
        parent.insert(list(parent).index(para_elem), new_elem)

    return new_elem


def _apply_para_format_from_last(doc, new_para):
    """将文档中最后一个段落的段落级格式应用到新段落"""
    if len(doc.paragraphs) < 2:
        return
    last_para_elem = doc.paragraphs[-2]._element
    last_pPr = last_para_elem.find(qn('w:pPr'))
    if last_pPr is not None:
        new_elem = new_para._element
        existing_pPr = new_elem.find(qn('w:pPr'))
        if existing_pPr is not None:
            new_elem.remove(existing_pPr)
        new_elem.insert(0, deepcopy(last_pPr))


def _apply_diff_run_inplace(result_para, diffs, font_info):
    """将diff结果写入段落，保留字体格式"""
    for tag, text in diffs:
        if not text:
            continue
        run = result_para.add_run(text)
        _apply_run_font(run, font_info)
        if tag == 'delete':
            run.font.color.rgb = RGBColor(0, 0, 255)
            run.font.strike = True
        elif tag == 'insert':
            run.font.color.rgb = RGBColor(255, 0, 0)


def _find_para_before_for_insert(result_doc, f_idx, f_match, final_paras, orig_paras_len):
    """查找插入位置的参考段落索引"""
    search_f = f_idx - 1
    while search_f >= 0:
        if not final_paras[search_f]['text']:
            search_f -= 1
            continue
        o = f_match[search_f]
        if o != -1:
            if isinstance(o, list):
                return min(o), search_f
            return o, search_f
        search_f -= 1
    return 0, -1


def _output_split_diff_inplace(result_doc, base_para, orig_text, all_f_indices,
                                final_paras, SENTENCE_SIM_THRESHOLD, font_info):
    """原地版拆分输出：1原稿段落 → N终稿段落"""
    combined_final_text = "".join(
        final_paras[sf]['text'] for sf in all_f_indices
    )
    diffs = char_level_diff(orig_text, combined_final_text)
    
    boundaries = _build_boundaries(
        all_f_indices, lambda sf: final_paras[sf]['text']
    )
    para_diffs_map = _distribute_diff_to_boundaries(diffs, boundaries)

    # 输出：第一个拆分到 base_para，后续插入新段落
    first = True
    insert_after_elem = base_para._element
    for sf in all_f_indices:
        sf_diffs = para_diffs_map[sf]

        if first:
            out_para = base_para
            _clear_para_runs(out_para)
            first = False
        else:
            piece_text = ""
            for tag, txt in sf_diffs:
                piece_text += txt
            if not piece_text:
                piece_text = final_paras[sf]['text']

            _create_new_para_elem(
                insert_after_elem, piece_text, font_info,
                is_after=True
            )
            insert_after_elem = insert_after_elem.getnext()
            continue

        if not sf_diffs:
            run = out_para.add_run(final_paras[sf]['text'])
            _apply_run_font(run, font_info)
        else:
            all_equal = all(t == 'equal' for t, _ in sf_diffs)
            if all_equal:
                for t, txt in sf_diffs:
                    run = out_para.add_run(txt)
                    _apply_run_font(run, font_info)
            else:
                for tag, text in sf_diffs:
                    if not text:
                        continue
                    run = out_para.add_run(text)
                    _apply_run_font(run, font_info)
                    if tag == 'delete':
                        run.font.color.rgb = RGBColor(0, 0, 255)
                        run.font.strike = True
                    elif tag == 'insert':
                        run.font.color.rgb = RGBColor(255, 0, 0)


def _extract_table_texts(doc):
    """
    提取文档中所有表格的单元格文本，返回扁平列表。
    每个元素包含表格索引、行列位置和文本内容。
    
    Args:
        doc: python-docx Document 对象
    
    Returns:
        [(table_idx, row_idx, col_idx, text), ...]
    """
    cells = []
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                text = cell.text
                cells.append((t_idx, r_idx, c_idx, text))
    return cells


def _compare_table_cells(orig_cells, final_cells, result_doc):
    """
    比较表格单元格并在结果文档中标记差异。
    
    匹配策略：先按 (table_idx, row_idx, col_idx) 精确匹配，
    然后对文本不同的单元格做字符级 diff 标记。
    
    Args:
        orig_cells: 原稿单元格列表 [(t_idx, r_idx, c_idx, text), ...]
        final_cells: 终稿单元格列表 [(t_idx, r_idx, c_idx, text), ...]
        result_doc: 结果文档（已加载，基于原稿）
    """
    # 构建索引查找表
    final_cell_map = {}
    for t_idx, r_idx, c_idx, text in final_cells:
        key = (t_idx, r_idx, c_idx)
        final_cell_map[key] = text
    
    # 遍历结果文档的表格，标记差异
    for t_idx, table in enumerate(result_doc.tables):
        if t_idx >= len(result_doc.tables):
            break
        for r_idx, row in enumerate(table.rows):
            if r_idx >= len(table.rows):
                break
            for c_idx, cell in enumerate(row.cells):
                if c_idx >= len(row.cells):
                    break
                key = (t_idx, r_idx, c_idx)
                orig_text = cell.text
                final_text = final_cell_map.get(key)
                
                if final_text is None:
                    # 原稿有此单元格，终稿没有 → 蓝色删除
                    if orig_text.strip():
                        _clear_cell_paragraphs(cell)
                        para = cell.paragraphs[0]
                        run = para.add_run(orig_text)
                        run.font.color.rgb = RGBColor(0, 0, 255)
                        run.font.strike = True
                elif orig_text != final_text:
                    # 文本不同 → 字符级 diff
                    _clear_cell_paragraphs(cell)
                    para = cell.paragraphs[0]
                    diffs = char_level_diff(orig_text, final_text)
                    _apply_diff_run(para, diffs)
    
    # 终稿新增的表格/单元格（在原稿中不存在的）
    # 用标记记录新增的表格和单元格索引
    orig_cell_keys = set((t_idx, r_idx, c_idx) for t_idx, r_idx, c_idx, _ in orig_cells)
    new_cells = [(t_idx, r_idx, c_idx, text) for t_idx, r_idx, c_idx, text in final_cells
                 if (t_idx, r_idx, c_idx) not in orig_cell_keys and text.strip()]
    
    if new_cells:
        logger.info("检测到 %d 个新增表格单元格（终稿独有），无法在原稿副本中自动插入，请手动检查", len(new_cells))


def _clear_cell_paragraphs(cell):
    """清空单元格所有段落的文本内容，保留段落结构"""
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ''


def compare_with_python_inplace(original_path, final_path, output_path):
    """在原稿复制版上开展比对和标记，保留图片、表格、页眉页脚等元素

    流程：
    1. 文件级拷贝原稿 → output_path
    2. 段落匹配（贪婪匹配 + 合并/拆分检测 + 逆序对移动检测）
    3. 在副本上原地标记差异：修改段落文本样式、插入新增段落
    """
    import shutil

    try:
        # 预加载配置
        SENTENCE_SIM_THRESHOLD, PARA_SIM_THRESHOLD, SHORT_PARA_CHAR_THRESHOLD, SEMANTIC_UNIT_THRESHOLDS = load_compare_config()

        # 打开文档（匹配用）
        orig_doc = Document(original_path)
        final_doc = Document(final_path)

        orig_paras = get_paragraphs_with_style(orig_doc)
        final_paras = get_paragraphs_with_style(final_doc)

        # 拷贝原稿为结果文档
        shutil.copy2(original_path, output_path)
        result_doc = Document(output_path)

        # ── 段落匹配 ──
        orig_matched = set()
        final_matched = set()
        orig_map = {}
        final_map = {}

        for i, p in enumerate(orig_paras):
            if p['text'] not in orig_map:
                orig_map[p['text']] = []
            orig_map[p['text']].append(i)

        for i, p in enumerate(final_paras):
            if p['text'] not in final_map:
                final_map[p['text']] = []
            final_map[p['text']].append(i)

        for text, indices in final_map.items():
            if not text:  # 空段落不参与精确匹配，避免跨索引跳跃
                continue
            if text in orig_map:
                for o_idx in orig_map[text]:
                    for f_idx in indices:
                        if o_idx not in orig_matched and f_idx not in final_matched:
                            orig_matched.add(o_idx)
                            final_matched.add(f_idx)
                            break

        f_match = [-1] * len(final_paras)
        matched_orig = set()

        for text, f_indices in final_map.items():
            if not text:  # 空段落不参与精确匹配
                continue
            if text in orig_map:
                for o_idx in orig_map[text]:
                    for f_idx in f_indices:
                        if o_idx not in matched_orig and f_match[f_idx] == -1:
                            f_match[f_idx] = o_idx
                            matched_orig.add(o_idx)
                            break

        # 预计算"有效（非空）索引"，用于 position_score 时空段落不计入位置偏移
        orig_nonempty_idx = []  # orig_nonempty_idx[o_idx] = 该段在"仅非空段"中的序号
        orig_nonempty_count = 0
        for i, p in enumerate(orig_paras):
            if p['text'].strip():
                orig_nonempty_idx.append(orig_nonempty_count)
                orig_nonempty_count += 1
            else:
                orig_nonempty_idx.append(-1)  # 空段标记为 -1

        final_nonempty_idx = []
        final_nonempty_count = 0
        for i, p in enumerate(final_paras):
            if p['text'].strip():
                final_nonempty_idx.append(final_nonempty_count)
                final_nonempty_count += 1
            else:
                final_nonempty_idx.append(-1)

        match_candidates = []
        for f_idx in range(len(final_paras)):
            if f_match[f_idx] != -1:
                continue
            # 终稿空段落不参与相似度匹配，直接作为新增处理
            if not final_paras[f_idx]['text'].strip():
                continue
            for o_idx in range(len(orig_paras)):
                if o_idx in matched_orig:
                    continue
                # 原稿空段落不参与相似度匹配
                if not orig_paras[o_idx]['text'].strip():
                    continue
                ratio = difflib.SequenceMatcher(
                    None, orig_paras[o_idx]['text'], final_paras[f_idx]['text']
                ).ratio()
                if ratio >= PARA_SIM_THRESHOLD:
                    # 用非空段有效索引计算 position_score，空段落不干扰位置偏移
                    adj_o = orig_nonempty_idx[o_idx]
                    adj_f = final_nonempty_idx[f_idx]
                    if adj_o >= 0 and adj_f >= 0:
                        position_score = 1.0 / (abs(adj_o - adj_f) + 1)
                    else:
                        position_score = 0.5  #  fallback（理论上不会走到这里）
                    total_score = ratio * 0.7 + position_score * 0.3
                    match_candidates.append((total_score, ratio, o_idx, f_idx))

        match_candidates.sort(reverse=True, key=lambda x: x[0])

        for total_score, ratio, o_idx, f_idx in match_candidates:
            if o_idx in matched_orig:
                continue
            f_text = final_paras[f_idx]['text']
            o_text = orig_paras[o_idx]['text']
            if f_text == o_text:
                f_text_count = sum(1 for p in final_paras if p['text'] == f_text)
                if f_text_count > 1:
                    first_occurrence = True
                    for i in range(f_idx):
                        if final_paras[i]['text'] == f_text:
                            first_occurrence = False
                            break
                    if not first_occurrence:
                        continue
            if f_match[f_idx] != -1:
                existing_o_idx = f_match[f_idx]
                should_create_one_to_many = False
                if isinstance(existing_o_idx, list):
                    min_idx = min(existing_o_idx)
                    max_idx = max(existing_o_idx)
                    if (min_idx - 1 <= o_idx <= max_idx + 1):
                        should_create_one_to_many = True
                else:
                    if abs(o_idx - existing_o_idx) == 1 and ratio >= 0.6:
                        should_create_one_to_many = True
                if should_create_one_to_many:
                    if isinstance(existing_o_idx, list):
                        f_match[f_idx].append(o_idx)
                    else:
                        f_match[f_idx] = [existing_o_idx, o_idx]
                    matched_orig.add(o_idx)
                else:
                    continue
            else:
                f_match[f_idx] = o_idx
                matched_orig.add(o_idx)

        # 拆分组反向映射
        o_match = {}
        split_matched_final = set()

        # 合并检测
        if True:
            unmatched_orig = set()
            for o_i in range(len(orig_paras)):
                if o_i not in matched_orig:
                    unmatched_orig.add(o_i)

            merge_candidates = []
            for f_i in range(len(final_paras)):
                o_i = f_match[f_i]
                if o_i == -1 or isinstance(o_i, list):
                    continue
                for adjacent_o in [o_i - 1, o_i + 1]:
                    if adjacent_o < 0 or adjacent_o >= len(orig_paras):
                        continue
                    if adjacent_o in matched_orig:
                        continue
                    if not orig_paras[adjacent_o]['text']:
                        continue
                    if adjacent_o < o_i:
                        combined = orig_paras[adjacent_o]['text'] + orig_paras[o_i]['text']
                    else:
                        combined = orig_paras[o_i]['text'] + orig_paras[adjacent_o]['text']
                    combined_ratio = difflib.SequenceMatcher(
                        None, combined, final_paras[f_i]['text']
                    ).ratio()
                    current_ratio = difflib.SequenceMatcher(
                        None, orig_paras[o_i]['text'], final_paras[f_i]['text']
                    ).ratio()
                    if combined_ratio > current_ratio and combined_ratio >= PARA_SIM_THRESHOLD:
                        improvement = combined_ratio - current_ratio
                        merge_candidates.append((improvement, combined_ratio, o_i, adjacent_o, f_i))

            merge_candidates.sort(reverse=True, key=lambda x: x[0])
            for improvement, combined_ratio, o_i, adjacent_o, f_i in merge_candidates:
                if adjacent_o in matched_orig:
                    continue
                if isinstance(f_match[f_i], list) and adjacent_o in f_match[f_i]:
                    continue
                if isinstance(f_match[f_i], list):
                    f_match[f_i] = sorted(f_match[f_i] + [adjacent_o])
                else:
                    f_match[f_i] = sorted([o_i, adjacent_o])
                matched_orig.add(adjacent_o)

        # 拆分检测
        final_unmatched_after_step2 = [i for i in range(len(final_paras)) if f_match[i] == -1]
        if final_unmatched_after_step2:
            split_candidates = []
            for f_idx in final_unmatched_after_step2:
                f_text = final_paras[f_idx]['text']
                if not f_text:
                    continue
                for o_idx in range(len(orig_paras)):
                    o_text = orig_paras[o_idx]['text']
                    if not o_text:
                        continue
                    if o_idx in o_match:
                        continue
                    ratio = difflib.SequenceMatcher(None, o_text, f_text).ratio()
                    if ratio >= PARA_SIM_THRESHOLD:
                        position_score = 1.0 / (abs(o_idx - f_idx) + 1)
                        total_score = ratio * 0.7 + position_score * 0.3
                        split_candidates.append((total_score, ratio, o_idx, f_idx))

            # 拆分升级检测
            split_upgrade_candidates = []
            for f_idx in final_unmatched_after_step2:
                if f_idx in split_matched_final:
                    continue
                f_text = final_paras[f_idx]['text']
                if not f_text:
                    continue
                for adjacent_f in [f_idx - 1, f_idx + 1]:
                    if adjacent_f < 0 or adjacent_f >= len(final_paras):
                        continue
                    o_i = f_match[adjacent_f]
                    if o_i == -1 or isinstance(o_i, list):
                        continue
                    o_text = orig_paras[o_i]['text']
                    if not o_text:
                        continue
                    if adjacent_f < f_idx:
                        combined_final = final_paras[adjacent_f]['text'] + f_text
                    else:
                        combined_final = f_text + final_paras[adjacent_f]['text']
                    combined_ratio = difflib.SequenceMatcher(None, o_text, combined_final).ratio()
                    current_ratio = difflib.SequenceMatcher(
                        None, o_text, final_paras[adjacent_f]['text']
                    ).ratio()
                    if combined_ratio > current_ratio and combined_ratio >= PARA_SIM_THRESHOLD:
                        improvement = combined_ratio - current_ratio
                        split_upgrade_candidates.append(
                            (improvement, combined_ratio, o_i, adjacent_f, f_idx)
                        )

            split_upgrade_candidates.sort(reverse=True, key=lambda x: x[0])
            for improvement, combined_ratio, o_i, adjacent_f, f_idx in split_upgrade_candidates:
                if o_i in o_match:
                    continue
                if f_idx in split_matched_final:
                    continue
                if isinstance(f_match[adjacent_f], list) and f_idx in f_match[adjacent_f]:
                    continue
                o_match[o_i] = [adjacent_f, f_idx]
                split_matched_final.add(f_idx)

            split_candidates.sort(reverse=True, key=lambda x: x[0])
            for total_score, ratio, o_idx, f_idx in split_candidates:
                if f_idx in split_matched_final:
                    continue
                if o_idx in o_match and len(o_match[o_idx]) >= 3:
                    continue
                if o_idx in o_match:
                    existing_f_indices = o_match[o_idx]
                    if not any(abs(f_idx - ef) == 1 for ef in existing_f_indices):
                        continue
                if o_idx not in o_match:
                    o_match[o_idx] = []
                o_match[o_idx].append(f_idx)
                split_matched_final.add(f_idx)

        # 段落级逆序对检测（段落互换/重排）
        moved_para_orig_set = set()
        forward_moved_para_orig_set = set()

        para_pairs = []
        for f_i in range(len(final_paras)):
            o_i = f_match[f_i]
            if o_i != -1 and not isinstance(o_i, list):
                # 空段落不参与移动检测 —— 空段落是格式元素，其位置变化不代表内容移动
                if not final_paras[f_i]['text'].strip() and not orig_paras[o_i]['text'].strip():
                    continue
                para_pairs.append((o_i, f_i))

        para_pairs_sorted = sorted(para_pairs, key=lambda x: x[0])
        for i in range(len(para_pairs_sorted)):
            for j in range(i + 1, len(para_pairs_sorted)):
                if para_pairs_sorted[i][1] > para_pairs_sorted[j][1]:
                    moved_para_orig_set.add(para_pairs_sorted[i][0])
                    forward_moved_para_orig_set.add(para_pairs_sorted[j][0])

        # ── 至此，匹配逻辑完成，开始原地标记 ──

        # 预收集所有原稿段落的字体信息
        orig_font_info = [_get_para_font_info(p) for p in orig_doc.paragraphs]

        # 快照结果文档中的段落对象（用于原地修改，不受后续插入影响）
        result_paras = list(result_doc.paragraphs)

        processed_orig = set()
        next_orig_idx = 0
        outputted_split_final = set()

        for f_idx in range(len(final_paras)):
            o_idx = f_match[f_idx]
            final_text = final_paras[f_idx]['text']

            if f_idx in outputted_split_final:
                continue

            # 确定目标原稿索引
            if o_idx == -1:
                # 查找下一个有匹配的最终稿段落的 o_idx，作为蓝后的截止位置
                # 避免将后续正常出现的原稿段落误标为删除
                next_target = None
                for nf in range(f_idx + 1, len(final_paras)):
                    no = f_match[nf]
                    if no != -1:
                        if isinstance(no, list):
                            next_target = min(no)
                        else:
                            next_target = no
                        break
                if next_target is None:
                    next_target = len(orig_paras)
                target_o = next_target
            elif isinstance(o_idx, list):
                target_o = min(o_idx)
            else:
                target_o = o_idx

            # ── 新增段落（终稿独有）：红前蓝后 ──
            if o_idx == -1:
                ref_o_idx, _ = _find_para_before_for_insert(
                    result_doc, f_idx, f_match, final_paras, len(orig_paras)
                )
                if ref_o_idx is not None and ref_o_idx < len(orig_font_info):
                    insert_font = orig_font_info[ref_o_idx]
                else:
                    insert_font = orig_font_info[0] if orig_font_info else {
                        'name': '仿宋', 'eastAsia': '仿宋', 'size': Pt(16), 'bold': None
                    }

                if not final_text:
                    if next_orig_idx < len(result_paras):
                        _create_new_para_elem(
                            result_paras[next_orig_idx]._element,
                            '', insert_font, None,
                            is_after=False
                        )
                    else:
                        new_p = result_doc.add_paragraph('')
                        _apply_para_format_from_last(result_doc, new_p)
                else:
                    if next_orig_idx < len(result_paras):
                        _create_new_para_elem(
                            result_paras[next_orig_idx]._element,
                            final_text, insert_font, RGBColor(255, 0, 0),
                            is_after=False
                        )
                    else:
                        new_p = result_doc.add_paragraph()
                        _apply_para_format_from_last(result_doc, new_p)
                        run = new_p.add_run(final_text)
                        run.font.color.rgb = RGBColor(255, 0, 0)
                        _apply_run_font(run, insert_font)

                # 蓝后：处理删除段落
                while next_orig_idx < target_o:
                    if next_orig_idx not in processed_orig:
                        if next_orig_idx in moved_para_orig_set:
                            _set_para_red_with_deletion(
                                result_paras[next_orig_idx],
                                orig_paras[next_orig_idx]['text'],
                                orig_font_info[next_orig_idx]
                            )
                            processed_orig.add(next_orig_idx)
                            next_orig_idx += 1
                        elif next_orig_idx not in matched_orig:
                            if not orig_paras[next_orig_idx]['text']:
                                # 未匹配的空段落 → 从 DOM 删除，不留空行
                                _remove_element_safe(result_paras[next_orig_idx]._element)
                            else:
                                _set_para_red_with_deletion(
                                    result_paras[next_orig_idx],
                                    orig_paras[next_orig_idx]['text'],
                                    orig_font_info[next_orig_idx]
                                )
                            processed_orig.add(next_orig_idx)
                            next_orig_idx += 1
                        else:
                            break
                    else:
                        next_orig_idx += 1

            # 非新增段落
            else:
                # 步骤A: 处理当前匹配前的删除段落
                while next_orig_idx < target_o:
                    if next_orig_idx not in processed_orig:
                        if next_orig_idx in moved_para_orig_set:
                            _set_para_red_with_deletion(
                                result_paras[next_orig_idx],
                                orig_paras[next_orig_idx]['text'],
                                orig_font_info[next_orig_idx]
                            )
                            processed_orig.add(next_orig_idx)
                            next_orig_idx += 1
                        elif next_orig_idx not in matched_orig:
                            if not orig_paras[next_orig_idx]['text']:
                                # 未匹配的空段落 → 从 DOM 删除，不留空行
                                _remove_element_safe(result_paras[next_orig_idx]._element)
                            else:
                                _set_para_red_with_deletion(
                                    result_paras[next_orig_idx],
                                    orig_paras[next_orig_idx]['text'],
                                    orig_font_info[next_orig_idx]
                                )
                            processed_orig.add(next_orig_idx)
                            next_orig_idx += 1
                        else:
                            break
                    else:
                        next_orig_idx += 1

                # B2: N→1 合并段落
                if isinstance(o_idx, list):
                    combined_orig_text = ""
                    for idx in o_idx:
                        if idx not in processed_orig:
                            combined_orig_text += orig_paras[idx]['text']
                            processed_orig.add(idx)
                            if idx >= next_orig_idx:
                                next_orig_idx = idx + 1

                    first_o = min(o_idx)
                    p = result_paras[first_o]

                    # 纯图片段落跳过合并标记，保持原样
                    if _is_pure_image_para(p):
                        for idx in o_idx:
                            processed_orig.add(idx)
                        continue

                    _clear_para_runs(p)
                    font_info = orig_font_info[first_o]

                    sentence_level_diff(combined_orig_text, final_text, p, SENTENCE_SIM_THRESHOLD,
                                         SHORT_PARA_CHAR_THRESHOLD, SEMANTIC_UNIT_THRESHOLDS)
                    for r in p.runs:
                        _apply_run_font(r, font_info)

                    for extra_o in sorted(o_idx):
                        if extra_o != first_o and extra_o < len(result_paras):
                            # 纯图片段落跳过合并标记
                            if _is_pure_image_para(result_paras[extra_o]):
                                # 确保 extra_o 被标记为已处理，避免后续再被打标
                                processed_orig.add(extra_o)
                                continue
                            _clear_para_runs(result_paras[extra_o])
                            if orig_paras[extra_o]['text']:
                                run = result_paras[extra_o].add_run(
                                    orig_paras[extra_o]['text']
                                )
                                run.font.color.rgb = RGBColor(0, 0, 255)
                                run.font.strike = True
                                _apply_run_font(run, orig_font_info[extra_o])

                # B3: 一对一匹配
                else:
                    if o_idx in moved_para_orig_set:
                        orig_text = orig_paras[o_idx]['text']
                        font_info = orig_font_info[o_idx]

                        p_new = result_doc.add_paragraph()
                        _apply_para_format_from_last(result_doc, p_new)
                        sentence_level_diff(orig_text, final_text, p_new, SENTENCE_SIM_THRESHOLD,
                                             SHORT_PARA_CHAR_THRESHOLD, SEMANTIC_UNIT_THRESHOLDS)
                        for r in p_new.runs:
                            _apply_run_font(r, font_info)

                        if next_orig_idx < len(result_paras):
                            target_elem = result_paras[next_orig_idx]._element
                            parent = target_elem.getparent()
                            parent.insert(list(parent).index(target_elem), p_new._element)
                    else:
                        processed_orig.add(o_idx)
                        next_orig_idx = max(next_orig_idx, o_idx + 1)

                        orig_text = orig_paras[o_idx]['text']
                        p = result_paras[o_idx]
                        font_info = orig_font_info[o_idx]

                        # 纯图片段落不参与标记，保持原样
                        if _is_pure_image_para(p):
                            continue

                        if o_idx in o_match:
                            split_f_indices = sorted(o_match[o_idx])
                            all_f_indices = [f_idx] + [sf for sf in split_f_indices if sf != f_idx]

                            _clear_para_runs(p)
                            _output_split_diff_inplace(
                                result_doc, p, orig_text, all_f_indices,
                                final_paras, SENTENCE_SIM_THRESHOLD, font_info
                            )

                            for sf in all_f_indices:
                                if sf != f_idx:
                                    outputted_split_final.add(sf)

                        elif orig_text == final_text and o_idx not in moved_para_orig_set \
                                and o_idx not in forward_moved_para_orig_set:
                            pass
                        else:
                            _clear_para_runs(p)
                            sentence_level_diff(orig_text, final_text, p, SENTENCE_SIM_THRESHOLD,
                                                 SHORT_PARA_CHAR_THRESHOLD, SEMANTIC_UNIT_THRESHOLDS)
                            for r in p.runs:
                                _apply_run_font(r, font_info)

        # 步骤C: 处理剩余的删除段落和未处理的移动段落
        while next_orig_idx < len(orig_paras):
            if next_orig_idx not in processed_orig:
                if next_orig_idx not in matched_orig:
                    if not orig_paras[next_orig_idx]['text']:
                        # 未匹配的空段落 → 从 DOM 删除，不留空行
                        _remove_element_safe(result_paras[next_orig_idx]._element)
                    else:
                        _set_para_red_with_deletion(
                            result_paras[next_orig_idx],
                            orig_paras[next_orig_idx]['text'],
                            orig_font_info[next_orig_idx]
                        )
                elif next_orig_idx in moved_para_orig_set:
                    # 移动段落的新位置已在前方输出，
                    # 原位置标记为蓝色删除
                    _set_para_red_with_deletion(
                        result_paras[next_orig_idx],
                        orig_paras[next_orig_idx]['text'],
                        orig_font_info[next_orig_idx]
                    )
            next_orig_idx += 1

        # ── 表格内容对比 ──
        orig_cells = _extract_table_texts(orig_doc)
        final_cells = _extract_table_texts(final_doc)
        if orig_cells or final_cells:
            _compare_table_cells(orig_cells, final_cells, result_doc)

        result_doc.save(output_path)
        return True, "短句级比对（原地）"

    except Exception as e:
        import traceback
        traceback.print_exc()
        return False, str(e)


def check_and_convert_file(file_path):
    """检查文件类型，doc 转换为 docx"""

    ext = os.path.splitext(file_path)[1].lower()
    if ext == '.doc':
        doc_dir = os.path.dirname(file_path)
        doc_to_docx(doc_dir)
        docx_path = os.path.splitext(file_path)[0] + '.docx'
        if os.path.exists(docx_path):
            return docx_path
        logger.warning("警告：doc 文件转换失败: %s", file_path)
        return None
    elif ext == '.docx':
        return file_path
    else:
        logger.error("错误：不支持的文件类型 '%s'，仅支持 .doc 和 .docx", ext)
        return None


def main(workdir, original_path=None, final_path=None):
    """主函数
    Args:
        workdir: 工作目录
        original_path: 直接传入的原稿路径（可选）
        final_path: 直接传入的终稿路径（可选）
    """ 
    
    # 判断传入的是文件还是目录
    if original_path and final_path:
        if not os.path.exists(original_path):
            logger.error("文件不存在: %s", original_path)
            return
        if not os.path.exists(final_path):
            logger.error("文件不存在: %s", final_path)
            return
        original = check_and_convert_file(original_path)
        final = check_and_convert_file(final_path)
    else:
        # 传入目录，调用 find_docx_files 让用户选择
        original, final = find_docx_files(workdir)
        if not original or not final:
            return
    
    if not original or not final:
        logger.error("无法比较：请检查文件格式")
        return

    logger.info("原稿: %s", os.path.basename(original))
    logger.info("终稿: %s", os.path.basename(final))
    
    output_name = f"差异标注-{os.path.basename(original)}"
    logger.info("开始比较...")
    
    success, result_msg = compare_with_python_inplace(original, final, os.path.join(workdir, output_name))
    
    if success:
        logger.info("方法: %s", result_msg)
        logger.info("比较完成: %s", output_name)
    else:
        logger.error("\n比较失败: %s", result_msg)


if __name__ == "__main__":
    original_path = None
    final_path = None
    
    # 解析命令行参数
    # sys.argv[0] 是脚本路径
    # 如果参数数量 >= 3：第一个文件路径和第二个文件路径
    # 如果参数数量 == 2：工作目录
    if len(sys.argv) >= 3:
        # 传入两个文档路径（脚本路径 + 文件1 + 文件2）
        original_path = sys.argv[1]
        final_path = sys.argv[2]
        workdir = os.path.dirname(original_path)
    elif len(sys.argv) == 2:
        # 传入工作目录（脚本路径 + 工作目录）
        workdir = sys.argv[1]
    else:
        workdir = os.path.dirname(__file__)
    
    main(workdir, original_path, final_path)
