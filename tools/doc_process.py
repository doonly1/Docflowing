import os
import platform
import re
import shutil
import subprocess
import zlib
from typing import Optional

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

from logging_config import setup_logging, get_logger
from mystyle import para_fm, run_fm

setup_logging()
logger = get_logger(__name__)


def find_libreoffice() -> Optional[str]:
    """查找 LibreOffice 可执行文件路径，未找到返回 None

    Returns:
        str or None: soffice 可执行文件路径
    """
    candidates = []
    system = platform.system()
    if system == 'Windows':
        candidates = [
            os.path.join(os.environ.get('PROGRAMFILES', r'C:\Program Files'),
                         'LibreOffice', 'program', 'soffice.exe'),
            os.path.join(os.environ.get('PROGRAMFILES(X86)', r'C:\Program Files (x86)'),
                         'LibreOffice', 'program', 'soffice.exe'),
        ]
        try:
            import winreg
            key = winreg.OpenKey(winreg.HKEY_LOCAL_MACHINE,
                                 r'SOFTWARE\LibreOffice\LibreOffice', 0,
                                 winreg.KEY_READ | winreg.KEY_WOW64_64KEY)
            install_dir, _ = winreg.QueryValueEx(key, 'InstallDir')
            candidates.insert(0, os.path.join(install_dir, 'program', 'soffice.exe'))
        except Exception:
            pass
    elif system == 'Darwin':
        candidates = ['/Applications/LibreOffice.app/Contents/MacOS/soffice']
    else:
        candidates = ['libreoffice', 'soffice',
                      '/usr/bin/libreoffice', '/usr/bin/soffice',
                      '/snap/bin/libreoffice']
    for cmd in candidates:
        try:
            if os.path.isabs(cmd):
                if os.path.isfile(cmd):
                    return cmd
            else:
                resolved = shutil.which(cmd)
                if resolved:
                    return resolved
        except Exception:
            continue
    return None


def libreoffice_install_hint() -> str:
    """返回当前平台的 LibreOffice 安装提示"""
    system = platform.system()
    if system == 'Windows':
        return '请安装 LibreOffice: https://www.libreoffice.org/download/'
    elif system == 'Darwin':
        return '请安装 LibreOffice: brew install --cask libreoffice'
    else:
        return '请安装 LibreOffice: sudo apt install libreoffice  (或 yum/dnf/pacman 等对应命令)'


def doc_to_docx(workdir: str) -> Optional[str]:
    """将workdir目录下的.doc文件批量转换为.docx格式（跨平台）

    记录转换前的 .docx mtime，只认 mtime 有变化的才算本次转换成功。

    Returns:
        None if success, error message string if failed
    """
    doc_files = [f for f in os.listdir(workdir)
                 if f.lower().endswith('.doc') and not f.startswith("~$")]
    if not doc_files:
        return None

    # 记录转换前各 .docx 的 mtime，用于判断本次是否真的生成了新文件
    before_mtimes = {}
    for f in doc_files:
        base = os.path.splitext(f)[0]
        docx_path = os.path.join(workdir, base + '.docx')
        if os.path.exists(docx_path):
            before_mtimes[base] = os.path.getmtime(docx_path)

    def _get_converted():
        """返回本次转换中 .docx 被真正更新的文件列表"""
        converted = []
        for f in doc_files:
            base = os.path.splitext(f)[0]
            docx_path = os.path.join(workdir, base + '.docx')
            if not os.path.exists(docx_path):
                continue
            if not os.path.getsize(docx_path) > 0:
                continue
            old_mtime = before_mtimes.get(base)
            cur_mtime = os.path.getmtime(docx_path)
            # 之前没有 .docx 或 mtime 变了 → 本次真正生成了新文件
            if old_mtime is None or cur_mtime > old_mtime:
                converted.append(base)
        return converted

    def _cleanup(converted):
        """确认转换成功，删除原 .doc"""
        for base in converted:
            doc_path = os.path.join(workdir, base + '.doc')
            try:
                os.remove(doc_path)
                logger.debug('已删除原文件: %s', doc_path)
            except OSError as e:
                logger.warning('删除原文件失败: %s, %s', doc_path, e)

    convert_errors = []

    # ——— win32com 路径 ———
    try:
        from win32com import client
        word = client.Dispatch("Word.Application")
        word.Visible = False
        for file in doc_files:
            logger.info('转化docx：{}'.format(file))
            file_path = os.path.abspath(os.path.normpath(os.path.join(workdir, file)))
            try:
                doc = word.Documents.Open(file_path)
                new_path = os.path.abspath(os.path.normpath(file_path + "x"))
                doc.SaveAs(new_path, 12)
                doc.Close()
            except Exception as e:
                logger.warning('win32com 转换单个文件失败: %s, %s', file, e)
                convert_errors.append(file)
        word.Quit()

        _cleanup(_get_converted())
        if convert_errors:
            err_msg = f"win32com 部分转换失败: {', '.join(convert_errors)}"
            logger.warning(err_msg)
            return err_msg
        return None
    except (ImportError, ModuleNotFoundError):
        logger.info('win32com 不可用，回退到 LibreOffice')
    except Exception as e:
        logger.warning('win32com 转换失败，回退到 LibreOffice: %s', e)

    # ——— LibreOffice 路径 ———
    lo_cmd = find_libreoffice()
    if lo_cmd:
        for file in doc_files:
            logger.info('转化docx(LibreOffice)：{}'.format(file))

        doc_paths = [os.path.join(workdir, f) for f in doc_files]

        cmd = [lo_cmd, '--headless', '--convert-to', 'docx',
               '--outdir', workdir] + doc_paths
        try:
            subprocess.run(cmd, check=True, capture_output=True, timeout=300)
            _cleanup(_get_converted())
            return None
        except subprocess.TimeoutExpired:
            err_msg = 'LibreOffice 转换超时（5分钟）'
            logger.error(err_msg)
            return err_msg
        except subprocess.CalledProcessError as e:
            stderr = (e.stderr or b'').decode('utf-8', errors='replace')[:200]
            err_msg = f'LibreOffice 转换失败: {stderr}'
            logger.error(err_msg)
            return err_msg
        except Exception as e:
            err_msg = f'LibreOffice 转换异常: {e}'
            logger.error(err_msg)
            return err_msg

    err_msg = '未找到 Microsoft Word 或 LibreOffice，无法转换 .doc 文件'
    logger.warning(err_msg)
    logger.warning(libreoffice_install_hint())
    return err_msg


def save_docx(doc: Document, doc_name: str, workdir: Optional[str] = None) -> Optional[str]:
    """
    保存Word文档

    Args:
        doc: Document对象
        doc_name: 文件名（可包含路径）
        workdir: 可选的保存目录，默认为None（当前目录）

    Returns:
        保存的完整文件路径，失败返回None
    """
    if workdir is None:
        if os.path.dirname(doc_name):
            workdir = os.path.dirname(doc_name)
            doc_name = os.path.basename(doc_name)
        else:
            workdir = os.getcwd()

    doc_name = re.sub(r'[\\/:*?"<>|]', '_', doc_name)

    if not doc_name.lower().endswith('.docx'):
        doc_name += '.docx'

    hash_prefix = str(zlib.crc32(doc_name.encode()) % 10000).zfill(4)

    save_path = os.path.normpath(os.path.join(workdir, hash_prefix + "_" + doc_name))
    if os.path.exists(save_path):
        return save_path

    try:
        doc.save(save_path)
        logger.info("已另存: %s", doc_name)
        return save_path
    except Exception as e:
        logger.error("另存失败: %s", e)
        return None


TITLE_KEYWORDS = ['决议', '决定', '命令', '公报', '公告', '通告', '意见',
                  '通知', '通报', '报告', '请示', '批复', '议案', '的函', '纪要',
                  '计划', '总结', '申请', '名单', '制度', '办法', '规定', '方案', '要点']

_H1_PATTERN = re.compile(r'^[一二三四五六七八九十]+、')

_H2_PATTERN = re.compile(r'^（[一二三四五六七八九十]+）')


def set_headings(doc: Document) -> None:
    paras = doc.paragraphs
    for i, para in enumerate(paras):
        text = para.text.rstrip()

        if i < 3 and len(text) < 60 and text[-2:] in TITLE_KEYWORDS:
            para.style = doc.styles['H0']
            para_fm(para, 0, 0, 28.95, 0, 0, 0, 'C')
            for run in para.runs:
                run_fm(run, '方正小标宋简体', 22, 0, 0, 0, 0)

        if 1 < len(text) < 30 and text.endswith('：') and i > 0 and len(paras[i - 1].text) == 0:
            para.style = doc.styles['Fangsong']
            para_fm(para, 0, 0, 28.95, 0, 0, 0, 'J')
            for run in para.runs:
                run_fm(run, '仿宋', 16, 0, 0, 0, 0)

        if _H1_PATTERN.match(text) and len(text) < 60 and '：' not in text:
            para.style = doc.styles['H1']
            for run in para.runs:
                run_fm(run, '黑体', 16, 0, 0, 0, 0)

        if _H2_PATTERN.match(text) and len(text) < 60 and '：' not in text:
            para.style = doc.styles['H2']
            for run in para.runs:
                run_fm(run, '楷体', 16, 0, 0, 0, 0)


def set_appendix(doc: Document) -> None:
    paras = doc.paragraphs

    appendix_indices = []
    for i, para in enumerate(paras):
        if '附件：' in para.text[:5] and len(para.text) > 3:
            appendix_indices.append(i)

    if not appendix_indices:
        return

    first_idx = appendix_indices[0]

    appendix_names = []
    for n in range(len(paras[first_idx + 1:])):
        if str(n + 1) + '.' in paras[first_idx + n].text[:6]:
            name = paras[first_idx + n].text[2:]
            for p in ['：', '.', '。']:
                name = name.replace(p, '')
            appendix_names.append(name)

    if not appendix_names:
        appendix_names.append(paras[first_idx].text.split('：', 1)[-1])
        paras[first_idx].style = doc.styles['Apdix']
        para_fm(paras[first_idx], 0, 0, 28.95, '5ch', 0, '-3ch', 'L')

    count = len(appendix_names)

    if count > 0:
        for i in range(count):
            for j in range(count):
                if str(i + 1) + '.' in paras[first_idx + j].text:
                    paras[first_idx + j].style = doc.styles['Apdix 2']
                    para_fm(paras[first_idx + j], None, None, None, '6ch', None, '-1ch', 'L')
        paras[first_idx].style = doc.styles['Apdix 1']
        para_fm(paras[first_idx], 0, 0, 28.95, '6ch', 0, 0, 'L')
        para_fm(paras[first_idx], 0, 0, 28.95, '6ch', 0, '-4ch', 'L')

    for para in paras[first_idx:]:
        if '附' in para.text and '件' in para.text and len(para.text) < 5:
            if para.text.endswith('：'):
                para.text = para.text[:-1]
            para.style = doc.styles['SimHei']
            para_fm(para, 0, 0, 28.95, 0, 0, 0, 'L')
            for run in para.runs:
                run_fm(run, '黑体', 16, 0, 0, 0, 0)


def set_date(doc: Document) -> None:
    paras = doc.paragraphs
    for i, para in enumerate(paras):
        if '年' in para.text and '月' in para.text and '日' in para.text and len(para.text) < 12:
            para.style = doc.styles['dater']
            para_fm(para, 0, 0, 28.95, 0, '4ch', 0, 'R')

            paras[i - 1].style = doc.styles['Sign']
            prev_fmt = paras[i - 1].paragraph_format
            prev_fmt.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            prev_fmt.first_line_indent = Pt(0)
            signature_len = len(paras[i - 1].text)
            if len(para.text) in {9, 10}:
                prev_fmt.right_indent = Pt((8 - 0.5 * signature_len) * 16)
            elif len(para.text) == 11:
                prev_fmt.right_indent = Pt((8.5 - 0.5 * signature_len) * 16)
