"""
FB 文件库同步 - 文件转换器模块

支持将 doc/docx/pdf/pptx/xlsx/xls/md/txt 转换为 Markdown 格式
使用 MarkItDown 内部转换器（跳过 magika/onnxruntime），保持高质量同时降低内存
"""

import os
import re
from abc import ABC, abstractmethod
from typing import Optional, Dict
from pathlib import Path
from dataclasses import dataclass


MARKITDOWN_EXTENSIONS = {'.pdf', '.docx', '.pptx', '.xlsx', '.xls'}


@dataclass(kw_only=True, frozen=True)
class _StreamInfo:
    """轻量 StreamInfo，替代 markitdown._stream_info.StreamInfo，避免触发 magika 导入"""
    mimetype: Optional[str] = None
    extension: Optional[str] = None
    charset: Optional[str] = None
    filename: Optional[str] = None
    local_path: Optional[str] = None
    url: Optional[str] = None


class BaseConverter(ABC):
    """文件转换器基类"""

    @property
    @abstractmethod
    def file_type(self) -> str:
        """文件类型标识"""
        pass

    @abstractmethod
    def can_convert(self, file_path: str) -> bool:
        """判断是否能转换此文件"""
        pass

    @abstractmethod
    def convert(self, source_path: str) -> Optional[str]:
        """转换文件，返回 Markdown 内容，失败返回 None"""
        pass


class MarkItDownConverter(BaseConverter):
    """
    使用 MarkItDown 内部转换器（按扩展名直调，跳过 magika/onnxruntime）

    按需加载单个转换器，避免全量 MarkItDown 引擎带来的 ~57 MB 内存开销
    以及每文件 ML 推理延迟
    """

    _converter_cache: Dict[str, object] = {}

    _EXTENSION_MAP = {
        '.pdf':  ('markitdown.converters._pdf_converter', 'PdfConverter'),
        '.docx': ('markitdown.converters._docx_converter', 'DocxConverter'),
        '.pptx': ('markitdown.converters._pptx_converter', 'PptxConverter'),
        '.xlsx': ('markitdown.converters._xlsx_converter', 'XlsxConverter'),
        '.xls':  ('markitdown.converters._xlsx_converter', 'XlsConverter'),
    }

    @property
    def file_type(self) -> str:
        return "markitdown"

    def can_convert(self, file_path: str) -> bool:
        ext = os.path.splitext(file_path)[1].lower()
        return ext in MARKITDOWN_EXTENSIONS

    def _get_converter(self, ext: str):
        if ext not in self._converter_cache:
            module_path, class_name = self._EXTENSION_MAP[ext]
            import importlib
            module = importlib.import_module(module_path)
            cls = getattr(module, class_name)
            self._converter_cache[ext] = cls()
        return self._converter_cache[ext]

    def convert(self, source_path: str) -> Optional[str]:
        try:
            ext = os.path.splitext(source_path)[1].lower()
            converter = self._get_converter(ext)

            with open(source_path, 'rb') as f:
                stream_info = _StreamInfo(extension=ext)
                result = converter.convert(f, stream_info=stream_info)

            content = result.markdown if result else None

            # 检测扫描版 PDF（无文本层），跳过不同步
            if ext == '.pdf' and (not content or len(content.strip()) < 20):
                import logging
                logging.getLogger(__name__).info(f"PDF 内容为空（可能是扫描版），跳过: {source_path}")
                return None

            return content
        except Exception as e:
            import logging
            logging.getLogger(__name__).error(
                f"MarkItDown conversion failed: {source_path}, error: {e}"
            )
            return None


class DOCXConverter(BaseConverter):
    """DOCX 文件转换器（MarkItDown 不可用时的备用方案）"""

    @property
    def file_type(self) -> str:
        return "docx"

    def can_convert(self, file_path: str) -> bool:
        return file_path.lower().endswith('.docx')

    def convert(self, source_path: str) -> Optional[str]:
        try:
            from docx import Document
            doc = Document(source_path)

            lines = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                style_name = para.style.name if para.style else ""

                if "Heading" in style_name:
                    level = self._extract_heading_level(style_name)
                    lines.append(f"{'#' * level} {text}")
                elif para.runs and para.runs[0].bold:
                    lines.append(f"**{text}**")
                else:
                    lines.append(text)

            for table in doc.tables:
                lines.append(self._convert_table(table))

            return "\n\n".join(lines)
        except Exception as e:
            import logging
            logging.getLogger(__name__).error(f"DOCX conversion failed: {source_path}, error: {e}")
            return None

    def _extract_heading_level(self, style_name: str) -> int:
        """从样式名提取标题级别"""
        match = re.search(r'\d+', style_name)
        if match:
            return min(int(match.group()), 6)
        return 1

    def _convert_table(self, table) -> str:
        """将表格转换为 Markdown 格式"""
        rows = []
        for i, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")
            if i == 0:
                rows.append("| " + " | ".join(["---"] * len(cells)) + " |")
        return "\n".join(rows)


class PDFConverter(BaseConverter):
    """PDF 文件转换器（MarkItDown 不可用时的备用方案）"""

    @property
    def file_type(self) -> str:
        return "pdf"

    def can_convert(self, file_path: str) -> bool:
        return file_path.lower().endswith('.pdf')

    def convert(self, source_path: str) -> Optional[str]:
        try:
            import pdfplumber

            all_text = []
            with pdfplumber.open(source_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    text = page.extract_text()
                    if text:
                        text = self._clean_pdf_text(text)
                        all_text.append(f"## 第 {page_num} 页\n\n{text}")

            content = "\n\n".join(all_text) if all_text else None

            # 检测扫描版 PDF（无文本层），跳过不同步
            if not content or len(content.strip()) < 20:
                import logging
                logging.getLogger(__name__).info(f"PDF 内容为空（可能是扫描版），跳过: {source_path}")
                return None

            return content
        except Exception as e:
            import logging
            logging.getLogger(__name__).error(f"PDF conversion failed: {source_path}, error: {e}")
            return None

    def _clean_pdf_text(self, text: str) -> str:
        """清理 PDF 提取的文本"""
        lines = []
        for line in text.split('\n'):
            line = line.strip()
            if not line:
                continue

            if re.match(r'^-?\d+-?$', line):
                continue
            if re.match(r'^第\s*\d+\s*页$', line):
                continue
            if line.isdigit() and len(line) <= 3:
                continue
            if '版权所有' in line or '翻印必究' in line:
                continue

            lines.append(line)

        return '\n'.join(lines)


class MDConverter(BaseConverter):
    """Markdown 文件转换器"""

    @property
    def file_type(self) -> str:
        return "md"

    def can_convert(self, file_path: str) -> bool:
        return file_path.lower().endswith('.md')

    def convert(self, source_path: str) -> Optional[str]:
        try:
            encodings = ['utf-8', 'gbk', 'gb2312']
            for encoding in encodings:
                try:
                    with open(source_path, 'r', encoding=encoding) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue
            return None
        except Exception as e:
            import logging
            logging.getLogger(__name__).error(f"MD conversion failed: {source_path}, error: {e}")
            return None


class TXTConverter(BaseConverter):
    """TXT 文件转换器"""

    @property
    def file_type(self) -> str:
        return "txt"

    def can_convert(self, file_path: str) -> bool:
        return file_path.lower().endswith('.txt')

    def convert(self, source_path: str) -> Optional[str]:
        try:
            encodings = ['utf-8', 'gbk', 'gb2312']
            for encoding in encodings:
                try:
                    with open(source_path, 'r', encoding=encoding) as f:
                        content = f.read()
                        return content.strip()
                except UnicodeDecodeError:
                    continue
            return None
        except Exception as e:
            import logging
            logging.getLogger(__name__).error(f"TXT conversion failed: {source_path}, error: {e}")
            return None


CONVERTERS = {}

for ext in MARKITDOWN_EXTENSIONS:
    CONVERTERS[ext] = MarkItDownConverter()

CONVERTERS['.md'] = MDConverter()
CONVERTERS['.txt'] = TXTConverter()


def get_converter(file_path: str) -> Optional[BaseConverter]:
    """根据文件路径获取对应的转换器"""
    ext = os.path.splitext(file_path)[1].lower()
    return CONVERTERS.get(ext)


def can_convert(file_path: str) -> bool:
    """判断文件是否可转换"""
    return get_converter(file_path) is not None


def convert_file(source_path: str, relative_path: str, filebase_id: str) -> Optional[str]:
    """转换文件，返回 Markdown 内容"""
    converter = get_converter(source_path)
    if not converter:
        return None

    return converter.convert(source_path)