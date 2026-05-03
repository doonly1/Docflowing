import os
import uuid
import time
import zipfile
import io
import shutil

from flask import Blueprint, request, jsonify, g, send_file
from functools import wraps

from kb.database import get_db, get_visible_kb_ids, get_user_role, get_storage_path

kb_bp = Blueprint('kb', __name__, url_prefix='/api/kb')

PERMISSION_LEVELS = {'view': 0, 'edit': 1, 'manage': 2}


def _get_user_id():
    auth_header = request.headers.get('Authorization', '')
    token = None
    if auth_header.startswith('Bearer '):
        token = auth_header[7:]
    elif auth_header:
        token = auth_header

    if not token:
        data = request.get_json(silent=True) or {}
        token = data.get('token') or data.get('client_id')

    if not token:
        token = request.args.get('token') or request.args.get('client_id')

    if not token:
        token = request.form.get('token') or request.form.get('client_id')

    if not token:
        return None

    import json
    tokens_path = os.path.join(os.path.expanduser('~'), '.config', 'DocProc', 'auth', 'tokens.json')
    try:
        if os.path.exists(tokens_path):
            with open(tokens_path, 'r', encoding='utf-8') as f:
                tokens = json.load(f)
            return tokens.get(token)
    except Exception:
        pass
    return None


def _is_admin(user_id):
    return get_user_role(user_id) == 'admin'


def _check_kb_permission(kb_id, user_id, required_level):
    if not user_id:
        return False
    if _is_admin(user_id):
        return True
    db = get_db()
    row = db.execute(
        "SELECT permission_level FROM kb_permissions WHERE kb_id = ? AND user_id = ?",
        (kb_id, user_id)
    ).fetchone()
    if row:
        actual = PERMISSION_LEVELS.get(row['permission_level'], -1)
        return actual >= PERMISSION_LEVELS.get(required_level, 0)
    kb_row = db.execute(
        "SELECT owner_id FROM knowledge_bases WHERE id = ?", (kb_id,)
    ).fetchone()
    if kb_row and kb_row['owner_id'] == user_id:
        return True
    return False


def _require_kb_permission(required_level):
    def decorator(f):
        @wraps(f)
        def decorated(*args, **kwargs):
            if request.method == 'OPTIONS':
                return f(*args, **kwargs)

            user_id = _get_user_id()
            if not user_id:
                return jsonify({'success': False, 'message': '未登录，请先登录'}), 401

            kb_id = kwargs.get('kb_id')
            if kb_id and not _check_kb_permission(kb_id, user_id, required_level):
                return jsonify({'success': False, 'message': '权限不足'}), 403

            kwargs['_user_id'] = user_id
            return f(*args, **kwargs)
        return decorated
    return decorator


def _get_kb_storage_dir(kb_id):
    storage_root = get_storage_path()
    kb_dir = os.path.join(storage_root, kb_id)
    os.makedirs(kb_dir, exist_ok=True)
    return kb_dir


# ==================== Task 3: 知识库 CRUD ====================

@kb_bp.route('/list', methods=['POST'])
def create_kb():
    user_id = _get_user_id()
    if not user_id:
        return jsonify({'success': False, 'message': '未登录，请先登录'}), 401

    data = request.get_json()
    name = (data.get('name') or '').strip()
    kb_type = (data.get('kb_type') or 'upload').strip()
    local_path = (data.get('local_path') or '').strip()

    if not name:
        return jsonify({'success': False, 'message': '知识库名称不能为空'})
    if len(name) > 64:
        return jsonify({'success': False, 'message': '知识库名称不能超过64个字符'})

    if kb_type == 'local':
        if not local_path:
            return jsonify({'success': False, 'message': '本地目录路径不能为空'})
        if not os.path.isdir(local_path):
            return jsonify({'success': False, 'message': '指定的本地目录不存在'})
        if not os.path.isabs(local_path):
            return jsonify({'success': False, 'message': '本地目录路径必须是绝对路径'})
    else:
        local_path = ''
        kb_type = 'upload'

    db = get_db()
    kb_id = str(uuid.uuid4())
    now = time.time()
    db.execute(
        "INSERT INTO knowledge_bases (id, name, owner_id, kb_type, local_path, created_at) VALUES (?, ?, ?, ?, ?, ?)",
        (kb_id, name, user_id, kb_type, local_path, now)
    )
    db.execute(
        "INSERT INTO kb_permissions (kb_id, user_id, permission_level) VALUES (?, ?, ?)",
        (kb_id, user_id, 'manage')
    )
    db.commit()

    return jsonify({
        'success': True,
        'kb': {'id': kb_id, 'name': name, 'owner_id': user_id, 'created_at': now, 'kb_type': kb_type, 'local_path': local_path}
    })


@kb_bp.route('/list', methods=['GET'])
def list_kb():
    user_id = _get_user_id()
    if not user_id:
        return jsonify({'success': False, 'message': '未登录，请先登录'}), 401

    is_admin = _is_admin(user_id)
    visible_ids = get_visible_kb_ids(user_id, is_admin)
    db = get_db()

    kbs = []
    for kb_id in visible_ids:
        row = db.execute("SELECT * FROM knowledge_bases WHERE id = ?", (kb_id,)).fetchone()
        if row:
            perm_row = db.execute(
                "SELECT permission_level FROM kb_permissions WHERE kb_id = ? AND user_id = ?",
                (kb_id, user_id)
            ).fetchone()
            permission = 'manage' if row['owner_id'] == user_id or is_admin else (
                perm_row['permission_level'] if perm_row else 'view'
            )

            doc_count = db.execute(
                "SELECT COUNT(*) as cnt FROM documents WHERE kb_id = ?", (kb_id,)
            ).fetchone()['cnt']

            kbs.append({
                'id': row['id'],
                'name': row['name'],
                'owner_id': row['owner_id'],
                'created_at': row['created_at'],
                'permission': permission,
                'document_count': doc_count,
                'kb_type': row['kb_type'] if 'kb_type' in row.keys() else 'upload',
                'local_path': row['local_path'] if 'local_path' in row.keys() else ''
            })

    return jsonify({'success': True, 'kbs': kbs})


@kb_bp.route('/<kb_id>', methods=['PUT'])
@_require_kb_permission('manage')
def rename_kb(kb_id, _user_id=None):
    data = request.get_json()
    name = (data.get('name') or '').strip()
    if not name:
        return jsonify({'success': False, 'message': '名称不能为空'})

    db = get_db()
    db.execute("UPDATE knowledge_bases SET name = ? WHERE id = ?", (name, kb_id))
    db.commit()
    return jsonify({'success': True, 'message': '重命名成功'})


@kb_bp.route('/<kb_id>', methods=['DELETE'])
@_require_kb_permission('manage')
def delete_kb(kb_id, _user_id=None):
    db = get_db()
    kb_row = db.execute("SELECT kb_type FROM knowledge_bases WHERE id = ?", (kb_id,)).fetchone()
    is_local = kb_row and kb_row['kb_type'] == 'local'

    if not is_local:
        doc_rows = db.execute("SELECT id, filename FROM documents WHERE kb_id = ?", (kb_id,)).fetchall()
        storage_dir = _get_kb_storage_dir(kb_id)
        for doc in doc_rows:
            file_path = os.path.join(storage_dir, doc['filename'])
            if os.path.exists(file_path):
                os.remove(file_path)

        if os.path.exists(storage_dir):
            shutil.rmtree(storage_dir, ignore_errors=True)

    db.execute("DELETE FROM documents WHERE kb_id = ?", (kb_id,))
    db.execute("DELETE FROM categories WHERE kb_id = ?", (kb_id,))
    db.execute("DELETE FROM kb_permissions WHERE kb_id = ?", (kb_id,))
    db.execute("DELETE FROM knowledge_bases WHERE id = ?", (kb_id,))
    db.commit()

    return jsonify({'success': True, 'message': '知识库已删除'})


@kb_bp.route('/<kb_id>/transfer', methods=['POST'])
@_require_kb_permission('manage')
def transfer_kb(kb_id, _user_id=None):
    data = request.get_json()
    new_owner_id = (data.get('new_owner_id') or '').strip()
    keep_role = (data.get('keep_role') or 'editor').strip()

    if not new_owner_id:
        return jsonify({'success': False, 'message': '请指定新所有者'})
    if keep_role not in ('view', 'edit', 'manage'):
        keep_role = 'editor'

    db = get_db()
    kb_row = db.execute("SELECT owner_id FROM knowledge_bases WHERE id = ?", (kb_id,)).fetchone()
    if not kb_row:
        return jsonify({'success': False, 'message': '知识库不存在'})

    db.execute("UPDATE knowledge_bases SET owner_id = ? WHERE id = ?", (new_owner_id, kb_id))

    db.execute("DELETE FROM kb_permissions WHERE kb_id = ? AND user_id = ?", (kb_id, new_owner_id))
    db.execute("INSERT INTO kb_permissions (kb_id, user_id, permission_level) VALUES (?, ?, ?)",
               (kb_id, new_owner_id, 'manage'))

    db.execute("DELETE FROM kb_permissions WHERE kb_id = ? AND user_id = ?", (kb_id, _user_id))
    db.execute("INSERT INTO kb_permissions (kb_id, user_id, permission_level) VALUES (?, ?, ?)",
               (kb_id, _user_id, keep_role))

    db.commit()
    return jsonify({'success': True, 'message': '所有权移交成功'})


# ==================== Task 4: ACL 权限 API ====================

@kb_bp.route('/<kb_id>/members', methods=['GET'])
@_require_kb_permission('view')
def list_members(kb_id, _user_id=None):
    import json
    users_path = os.path.join(os.path.expanduser('~'), '.config', 'DocProc', 'auth', 'users.json')
    users = {}
    if os.path.exists(users_path):
        with open(users_path, 'r', encoding='utf-8') as f:
            users = json.load(f)

    db = get_db()
    kb_row = db.execute("SELECT owner_id FROM knowledge_bases WHERE id = ?", (kb_id,)).fetchone()

    members = []
    owner_id = kb_row['owner_id'] if kb_row else None
    if owner_id:
        owner_info = users.get(owner_id, {})
        members.append({
            'user_id': owner_id,
            'username': owner_info.get('username', ''),
            'permission': 'manage',
            'is_owner': True
        })

    perm_rows = db.execute(
        "SELECT user_id, permission_level FROM kb_permissions WHERE kb_id = ? AND user_id != ?",
        (kb_id, owner_id or '')
    ).fetchall()

    for row in perm_rows:
        user_info = users.get(row['user_id'], {})
        members.append({
            'user_id': row['user_id'],
            'username': user_info.get('username', ''),
            'permission': row['permission_level'],
            'is_owner': False
        })

    return jsonify({'success': True, 'members': members})


@kb_bp.route('/<kb_id>/members', methods=['POST'])
@_require_kb_permission('manage')
def add_member(kb_id, _user_id=None):
    import json
    users_path = os.path.join(os.path.expanduser('~'), '.config', 'DocProc', 'auth', 'users.json')
    users = {}
    if os.path.exists(users_path):
        with open(users_path, 'r', encoding='utf-8') as f:
            users = json.load(f)

    data = request.get_json()
    target_username = (data.get('username') or '').strip()
    permission = (data.get('permission') or 'view').strip()

    if permission not in ('view', 'edit', 'manage'):
        return jsonify({'success': False, 'message': '无效的权限级别'})

    target_user_id = None
    for uid, uinfo in users.items():
        if uinfo.get('username') == target_username:
            target_user_id = uid
            break

    if not target_user_id:
        return jsonify({'success': False, 'message': '用户不存在'})

    db = get_db()
    existing = db.execute(
        "SELECT * FROM kb_permissions WHERE kb_id = ? AND user_id = ?",
        (kb_id, target_user_id)
    ).fetchone()
    if existing:
        db.execute(
            "UPDATE kb_permissions SET permission_level = ? WHERE kb_id = ? AND user_id = ?",
            (permission, kb_id, target_user_id)
        )
    else:
        db.execute(
            "INSERT INTO kb_permissions (kb_id, user_id, permission_level) VALUES (?, ?, ?)",
            (kb_id, target_user_id, permission)
        )
    db.commit()

    return jsonify({'success': True, 'message': '成员已添加/更新'})


@kb_bp.route('/<kb_id>/members/<member_id>', methods=['PUT'])
@_require_kb_permission('manage')
def update_member(kb_id, member_id, _user_id=None):
    data = request.get_json()
    permission = (data.get('permission') or 'view').strip()
    if permission not in ('view', 'edit', 'manage'):
        return jsonify({'success': False, 'message': '无效的权限级别'})

    db = get_db()
    db.execute(
        "UPDATE kb_permissions SET permission_level = ? WHERE kb_id = ? AND user_id = ?",
        (permission, kb_id, member_id)
    )
    db.commit()
    return jsonify({'success': True, 'message': '权限已更新'})


@kb_bp.route('/<kb_id>/members/<member_id>', methods=['DELETE'])
@_require_kb_permission('manage')
def remove_member(kb_id, member_id, _user_id=None):
    db = get_db()
    db.execute(
        "DELETE FROM kb_permissions WHERE kb_id = ? AND user_id = ?",
        (kb_id, member_id)
    )
    db.commit()
    return jsonify({'success': True, 'message': '成员已移除'})


# ==================== Task 5: 分类 CRUD API ====================

def _build_category_tree(kb_id):
    db = get_db()
    rows = db.execute(
        "SELECT * FROM categories WHERE kb_id = ? ORDER BY name", (kb_id,)
    ).fetchall()

    cat_map = {}
    tree = []
    for r in rows:
        node = {
            'id': r['id'],
            'name': r['name'],
            'parent_id': r['parent_id'],
            'children': []
        }
        cat_map[r['id']] = node

    for node in cat_map.values():
        if node['parent_id'] and node['parent_id'] in cat_map:
            cat_map[node['parent_id']]['children'].append(node)
        else:
            tree.append(node)

    return tree


@kb_bp.route('/<kb_id>/categories', methods=['POST'])
@_require_kb_permission('edit')
def create_category(kb_id, _user_id=None):
    data = request.get_json()
    name = (data.get('name') or '').strip()
    parent_id = (data.get('parent_id') or '').strip() or None

    if not name:
        return jsonify({'success': False, 'message': '分类名称不能为空'})

    db = get_db()
    cat_id = str(uuid.uuid4())
    db.execute(
        "INSERT INTO categories (id, kb_id, parent_id, name, created_at) VALUES (?, ?, ?, ?, ?)",
        (cat_id, kb_id, parent_id, name, time.time())
    )
    db.commit()

    return jsonify({'success': True, 'category': {'id': cat_id, 'name': name, 'parent_id': parent_id}})


@kb_bp.route('/<kb_id>/categories', methods=['GET'])
@_require_kb_permission('view')
def list_categories(kb_id, _user_id=None):
    tree = _build_category_tree(kb_id)
    return jsonify({'success': True, 'categories': tree})


@kb_bp.route('/<kb_id>/categories/<cat_id>', methods=['PUT'])
@_require_kb_permission('edit')
def rename_category(kb_id, cat_id, _user_id=None):
    data = request.get_json()
    name = (data.get('name') or '').strip()
    if not name:
        return jsonify({'success': False, 'message': '名称不能为空'})

    db = get_db()
    db.execute("UPDATE categories SET name = ? WHERE id = ? AND kb_id = ?", (name, cat_id, kb_id))
    db.commit()
    return jsonify({'success': True, 'message': '重命名成功'})


@kb_bp.route('/<kb_id>/categories/<cat_id>', methods=['DELETE'])
@_require_kb_permission('edit')
def delete_category(kb_id, cat_id, _user_id=None):
    db = get_db()
    db.execute("UPDATE documents SET category_id = NULL WHERE category_id = ? AND kb_id = ?",
               (cat_id, kb_id))
    db.execute("UPDATE categories SET parent_id = NULL WHERE parent_id = ? AND kb_id = ?",
               (cat_id, kb_id))
    db.execute("DELETE FROM categories WHERE id = ? AND kb_id = ?", (cat_id, kb_id))
    db.commit()
    return jsonify({'success': True, 'message': '分类已删除'})


# ==================== Task 6: 文档 CRUD API ====================

def _get_allowed_extensions():
    from kb.database import _load_kb_config
    kb_config = _load_kb_config()
    exts = kb_config.get('allowed_extensions', [])
    if not exts:
        exts = ['.docx', '.doc', '.xlsx', '.xls', '.pptx', '.ppt',
                '.pdf', '.md', '.txt', '.html', '.htm']
    return [e.lower() if e.startswith('.') else f'.{e.lower()}' for e in exts]


def _get_max_file_size():
    from kb.database import _load_kb_config
    kb_config = _load_kb_config()
    return (kb_config.get('max_file_size_mb', 100) or 100) * 1024 * 1024


@kb_bp.route('/<kb_id>/documents', methods=['POST'])
@_require_kb_permission('edit')
def upload_document(kb_id, _user_id=None):
    if 'file' not in request.files:
        return jsonify({'success': False, 'message': '请选择文件'})

    file = request.files['file']
    if not file.filename:
        return jsonify({'success': False, 'message': '文件名不能为空'})

    original_name = file.filename
    ext = os.path.splitext(original_name)[1].lower()
    allowed = _get_allowed_extensions()
    if allowed and ext not in allowed:
        return jsonify({'success': False, 'message': f'不支持的文件类型: {ext}'})

    file_content = file.read()
    file_size = len(file_content)
    max_size = _get_max_file_size()
    if file_size > max_size:
        return jsonify({'success': False, 'message': f'文件超过 {max_size // 1024 // 1024}MB 限制'})
    file.seek(0)

    category_id = (request.form.get('category_id') or '').strip() or None

    doc_id = str(uuid.uuid4())
    storage_filename = f"{doc_id}{ext}"
    storage_dir = _get_kb_storage_dir(kb_id)
    save_path = os.path.join(storage_dir, storage_filename)

    with open(save_path, 'wb') as f:
        f.write(file_content)

    now = time.time()
    db = get_db()
    db.execute(
        """INSERT INTO documents (id, kb_id, category_id, filename, original_name,
           file_size, file_type, created_at, updated_at)
           VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)""",
        (doc_id, kb_id, category_id, storage_filename, original_name,
         file_size, ext, now, now)
    )
    db.commit()

    return jsonify({
        'success': True,
        'document': {
            'id': doc_id,
            'filename': original_name,
            'file_size': file_size,
            'file_type': ext,
            'category_id': category_id,
            'created_at': now
        }
    })


@kb_bp.route('/<kb_id>/documents', methods=['GET'])
@_require_kb_permission('view')
def list_documents(kb_id, _user_id=None):
    category_id = request.args.get('category_id', '').strip() or None
    file_type = request.args.get('type', '').strip() or None

    db = get_db()
    query = "SELECT * FROM documents WHERE kb_id = ?"
    params = [kb_id]

    if category_id:
        query += " AND category_id = ?"
        params.append(category_id)

    if file_type:
        query += " AND file_type = ?"
        params.append(file_type.lower() if file_type.startswith('.') else f'.{file_type.lower()}')

    query += " ORDER BY updated_at DESC"

    rows = db.execute(query, params).fetchall()
    docs = []
    for r in rows:
        docs.append({
            'id': r['id'],
            'filename': r['original_name'],
            'file_size': r['file_size'],
            'file_type': r['file_type'],
            'category_id': r['category_id'],
            'created_at': r['created_at'],
            'updated_at': r['updated_at']
        })

    return jsonify({'success': True, 'documents': docs})


@kb_bp.route('/<kb_id>/documents/<doc_id>', methods=['GET'])
@_require_kb_permission('view')
def get_document(kb_id, doc_id, _user_id=None):
    db = get_db()
    row = db.execute(
        "SELECT * FROM documents WHERE id = ? AND kb_id = ?", (doc_id, kb_id)
    ).fetchone()
    if not row:
        return jsonify({'success': False, 'message': '文档不存在'})

    return jsonify({
        'success': True,
        'document': {
            'id': row['id'],
            'filename': row['original_name'],
            'file_size': row['file_size'],
            'file_type': row['file_type'],
            'category_id': row['category_id'],
            'created_at': row['created_at'],
            'updated_at': row['updated_at']
        }
    })


@kb_bp.route('/<kb_id>/documents/<doc_id>', methods=['PUT'])
@_require_kb_permission('edit')
def update_document(kb_id, doc_id, _user_id=None):
    data = request.get_json()
    db = get_db()
    row = db.execute(
        "SELECT * FROM documents WHERE id = ? AND kb_id = ?", (doc_id, kb_id)
    ).fetchone()
    if not row:
        return jsonify({'success': False, 'message': '文档不存在'})

    new_name = data.get('filename')
    new_category_id = data.get('category_id')

    if new_name is not None:
        new_name = new_name.strip()
        if not new_name:
            return jsonify({'success': False, 'message': '文件名不能为空'})
        db.execute("UPDATE documents SET original_name = ?, updated_at = ? WHERE id = ?",
                   (new_name, time.time(), doc_id))

    if new_category_id is not None:
        db.execute("UPDATE documents SET category_id = ?, updated_at = ? WHERE id = ?",
                   (new_category_id, time.time(), doc_id))

    db.commit()
    return jsonify({'success': True, 'message': '更新成功'})


@kb_bp.route('/<kb_id>/documents/<doc_id>', methods=['DELETE'])
@_require_kb_permission('edit')
def delete_document(kb_id, doc_id, _user_id=None):
    db = get_db()
    row = db.execute(
        "SELECT filename FROM documents WHERE id = ? AND kb_id = ?", (doc_id, kb_id)
    ).fetchone()
    if not row:
        return jsonify({'success': False, 'message': '文档不存在'})

    storage_dir = _get_kb_storage_dir(kb_id)
    file_path = os.path.join(storage_dir, row['filename'])
    if os.path.exists(file_path):
        os.remove(file_path)

    db.execute("DELETE FROM documents WHERE id = ? AND kb_id = ?", (doc_id, kb_id))
    db.commit()

    return jsonify({'success': True, 'message': '文档已删除'})


@kb_bp.route('/<kb_id>/documents/<doc_id>/download', methods=['GET'])
@_require_kb_permission('view')
def download_document(kb_id, doc_id, _user_id=None):
    db = get_db()
    row = db.execute(
        "SELECT filename, original_name FROM documents WHERE id = ? AND kb_id = ?",
        (doc_id, kb_id)
    ).fetchone()
    if not row:
        return jsonify({'success': False, 'message': '文档不存在'})

    storage_dir = _get_kb_storage_dir(kb_id)
    file_path = os.path.join(storage_dir, row['filename'])
    if not os.path.exists(file_path):
        return jsonify({'success': False, 'message': '文件不存在'})

    return send_file(file_path, as_attachment=True, download_name=row['original_name'])


@kb_bp.route('/<kb_id>/documents/<doc_id>/replace', methods=['PUT'])
@_require_kb_permission('edit')
def replace_document(kb_id, doc_id, _user_id=None):
    if 'file' not in request.files:
        return jsonify({'success': False, 'message': '请选择文件'})

    db = get_db()
    old_row = db.execute(
        "SELECT filename FROM documents WHERE id = ? AND kb_id = ?", (doc_id, kb_id)
    ).fetchone()
    if not old_row:
        return jsonify({'success': False, 'message': '文档不存在'})

    file = request.files['file']
    if not file.filename:
        return jsonify({'success': False, 'message': '文件名不能为空'})

    original_name = file.filename
    ext = os.path.splitext(original_name)[1].lower()
    allowed = _get_allowed_extensions()
    if allowed and ext not in allowed:
        return jsonify({'success': False, 'message': f'不支持的文件类型: {ext}'})

    file_content = file.read()
    file_size = len(file_content)
    max_size = _get_max_file_size()
    if file_size > max_size:
        return jsonify({'success': False, 'message': f'文件超过 {max_size // 1024 // 1024}MB 限制'})

    storage_dir = _get_kb_storage_dir(kb_id)
    old_path = os.path.join(storage_dir, old_row['filename'])
    if os.path.exists(old_path):
        os.remove(old_path)

    new_storage_name = f"{doc_id}{ext}"
    new_path = os.path.join(storage_dir, new_storage_name)
    with open(new_path, 'wb') as f:
        f.write(file_content)

    now = time.time()
    db.execute(
        """UPDATE documents SET filename = ?, original_name = ?, file_size = ?,
           file_type = ?, updated_at = ? WHERE id = ? AND kb_id = ?""",
        (new_storage_name, original_name, file_size, ext, now, doc_id, kb_id)
    )
    db.commit()

    return jsonify({'success': True, 'message': '替换成功'})


@kb_bp.route('/<kb_id>/documents/batch-download', methods=['POST'])
@_require_kb_permission('view')
def batch_download(kb_id, _user_id=None):
    data = request.get_json()
    doc_ids = data.get('doc_ids', [])
    if not doc_ids:
        return jsonify({'success': False, 'message': '请选择文档'})

    db = get_db()
    storage_dir = _get_kb_storage_dir(kb_id)

    memory_file = io.BytesIO()
    with zipfile.ZipFile(memory_file, 'w', zipfile.ZIP_DEFLATED) as zf:
        for doc_id in doc_ids:
            row = db.execute(
                "SELECT filename, original_name FROM documents WHERE id = ? AND kb_id = ?",
                (doc_id, kb_id)
            ).fetchone()
            if row:
                file_path = os.path.join(storage_dir, row['filename'])
                if os.path.exists(file_path):
                    zf.write(file_path, row['original_name'])

    memory_file.seek(0)
    return send_file(memory_file, mimetype='application/zip',
                     as_attachment=True, download_name='documents.zip')


# ==================== Task 7: Markdown 内容读写 API ====================

@kb_bp.route('/<kb_id>/documents/<doc_id>/content', methods=['GET'])
@_require_kb_permission('view')
def get_document_content(kb_id, doc_id, _user_id=None):
    db = get_db()
    row = db.execute(
        "SELECT filename, file_type FROM documents WHERE id = ? AND kb_id = ?",
        (doc_id, kb_id)
    ).fetchone()
    if not row:
        return jsonify({'success': False, 'message': '文档不存在'})

    storage_dir = _get_kb_storage_dir(kb_id)
    file_path = os.path.join(storage_dir, row['filename'])
    if not os.path.exists(file_path):
        return jsonify({'success': False, 'message': '文件不存在'})

    ext = (row['file_type'] or '').lower()
    if ext in ('.md', '.txt', '.html', '.htm', '.xml', '.json', '.csv', '.yaml', '.yml', '.py', '.js', '.css'):
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            return jsonify({'success': True, 'content': content, 'file_type': ext})
        except UnicodeDecodeError:
            return jsonify({'success': False, 'message': '无法以文本方式读取此文件'})

    if ext == '.docx':
        try:
            from docx import Document
            doc = Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs]
            return jsonify({'success': True, 'content': '\n'.join(paragraphs), 'file_type': ext})
        except Exception:
            return jsonify({'success': False, 'message': '无法读取 docx 内容'})

    if ext in ('.xlsx', '.xls'):
        try:
            import openpyxl
            wb = openpyxl.load_workbook(file_path, data_only=True)
            lines = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                lines.append(f'# {sheet_name}')
                for row in ws.iter_rows(values_only=True):
                    lines.append('\t'.join(str(c) if c is not None else '' for c in row))
            return jsonify({'success': True, 'content': '\n'.join(lines), 'file_type': ext})
        except Exception:
            return jsonify({'success': False, 'message': '无法读取 xlsx 内容'})

    return jsonify({'success': False, 'message': '不支持在线查看此文件类型的内容'})


@kb_bp.route('/<kb_id>/documents/<doc_id>/content', methods=['PUT'])
@_require_kb_permission('edit')
def save_document_content(kb_id, doc_id, _user_id=None):
    data = request.get_json()
    content = data.get('content', '')
    if content is None:
        content = ''

    db = get_db()
    row = db.execute(
        "SELECT filename FROM documents WHERE id = ? AND kb_id = ?",
        (doc_id, kb_id)
    ).fetchone()
    if not row:
        return jsonify({'success': False, 'message': '文档不存在'})

    storage_dir = _get_kb_storage_dir(kb_id)
    file_path = os.path.join(storage_dir, row['filename'])

    with open(file_path, 'w', encoding='utf-8') as f:
        f.write(content)

    db.execute("UPDATE documents SET file_size = ?, updated_at = ? WHERE id = ?",
               (len(content.encode('utf-8')), time.time(), doc_id))
    db.commit()

    return jsonify({'success': True, 'message': '保存成功'})


# ==================== Task 8: 全文搜索 API ====================

@kb_bp.route('/search', methods=['GET'])
def search_documents():
    user_id = _get_user_id()
    if not user_id:
        return jsonify({'success': False, 'message': '未登录，请先登录'}), 401

    q = (request.args.get('q') or '').strip()
    if not q:
        return jsonify({'success': True, 'results': []})

    is_admin = _is_admin(user_id)
    visible_ids = get_visible_kb_ids(user_id, is_admin)
    if not visible_ids:
        return jsonify({'success': True, 'results': []})

    db = get_db()
    results = []
    keywords = q.lower().split()

    for kb_id in visible_ids:
        kb_row = db.execute("SELECT name, kb_type, local_path FROM knowledge_bases WHERE id = ?", (kb_id,)).fetchone()
        kb_name = kb_row['name'] if kb_row else ''
        kb_type = (kb_row['kb_type'] if 'kb_type' in kb_row.keys() else 'upload') if kb_row else 'upload'
        local_path = (kb_row['local_path'] if 'local_path' in kb_row.keys() else '') if kb_row else ''

        if kb_type == 'local' and local_path and os.path.isdir(local_path):
            results.extend(_search_local_dir(local_path, kb_id, kb_name, keywords))
        else:
            doc_rows = db.execute(
                "SELECT * FROM documents WHERE kb_id = ? ORDER BY updated_at DESC", (kb_id,)
            ).fetchall()

            for doc in doc_rows:
                matched = False
                match_type = ''
                original_lower = doc['original_name'].lower()

                for kw in keywords:
                    if kw in original_lower:
                        matched = True
                        match_type = 'filename'
                        break

                if not matched:
                    ext = (doc['file_type'] or '').lower()
                    if ext in ('.md', '.txt', '.html', '.htm', '.xml', '.json', '.csv'):
                        storage_dir = _get_kb_storage_dir(kb_id)
                        file_path = os.path.join(storage_dir, doc['filename'])
                        if os.path.exists(file_path):
                            try:
                                with open(file_path, 'r', encoding='utf-8') as f:
                                    content = f.read().lower()
                                for kw in keywords:
                                    if kw in content:
                                        matched = True
                                        match_type = 'content'
                                        break
                            except Exception:
                                pass

                    if not matched and ext in ('.docx',):
                        storage_dir = _get_kb_storage_dir(kb_id)
                        file_path = os.path.join(storage_dir, doc['filename'])
                        if os.path.exists(file_path):
                            try:
                                from docx import Document
                                d = Document(file_path)
                                text = '\n'.join(p.text for p in d.paragraphs).lower()
                                for kw in keywords:
                                    if kw in text:
                                        matched = True
                                        match_type = 'content'
                                        break
                            except Exception:
                                pass

                    if not matched and ext in ('.xlsx', '.xls'):
                        storage_dir = _get_kb_storage_dir(kb_id)
                        file_path = os.path.join(storage_dir, doc['filename'])
                        if os.path.exists(file_path):
                            try:
                                import openpyxl
                                wb = openpyxl.load_workbook(file_path, data_only=True)
                                text_parts = []
                                for sn in wb.sheetnames:
                                    ws = wb[sn]
                                    for row in ws.iter_rows(values_only=True):
                                        text_parts.append(' '.join(
                                            str(c) if c is not None else '' for c in row))
                                text = ' '.join(text_parts).lower()
                                for kw in keywords:
                                    if kw in text:
                                        matched = True
                                        match_type = 'content'
                                        break
                            except Exception:
                                pass

                if matched:
                    results.append({
                        'document_id': doc['id'],
                        'kb_id': kb_id,
                        'kb_name': kb_name,
                        'filename': doc['original_name'],
                        'file_type': doc['file_type'],
                        'file_size': doc['file_size'],
                        'updated_at': doc['updated_at'],
                        'match_type': match_type
                    })

    return jsonify({'success': True, 'results': results, 'query': q})


def _search_local_dir(base_path, kb_id, kb_name, keywords):
    results = []
    try:
        for root, dirs, files in os.walk(base_path):
            dirs[:] = [d for d in dirs if not d.startswith('~$')]
            for fname in files:
                if fname.startswith('~$'):
                    continue
                matched = False
                match_type = ''
                fname_lower = fname.lower()

                for kw in keywords:
                    if kw in fname_lower:
                        matched = True
                        match_type = 'filename'
                        break

                if not matched:
                    ext = os.path.splitext(fname)[1].lower()
                    if ext in ('.md', '.txt', '.html', '.htm', '.xml', '.json', '.csv'):
                        file_path = os.path.join(root, fname)
                        try:
                            with open(file_path, 'r', encoding='utf-8') as f:
                                content = f.read().lower()
                            for kw in keywords:
                                if kw in content:
                                    matched = True
                                    match_type = 'content'
                                    break
                        except Exception:
                            pass

                if matched:
                    full_path = os.path.join(root, fname)
                    rel_path = os.path.relpath(full_path, base_path).replace('\\', '/')
                    stat = os.stat(full_path)
                    results.append({
                        'document_id': rel_path,
                        'kb_id': kb_id,
                        'kb_name': kb_name,
                        'filename': fname,
                        'file_type': os.path.splitext(fname)[1],
                        'file_size': stat.st_size,
                        'updated_at': stat.st_mtime,
                        'match_type': match_type,
                        'rel_path': rel_path
                    })
    except PermissionError:
        pass
    return results


@kb_bp.route('/<kb_id>/local-files', methods=['GET'])
@_require_kb_permission('view')
def list_local_files(kb_id, _user_id=None):
    db = get_db()
    kb_row = db.execute("SELECT local_path, kb_type FROM knowledge_bases WHERE id = ?", (kb_id,)).fetchone()
    if not kb_row:
        return jsonify({'success': False, 'message': '知识库不存在'})
    if kb_row['kb_type'] != 'local':
        return jsonify({'success': False, 'message': '仅支持本地目录知识库'})

    local_path = kb_row['local_path']
    if not os.path.isdir(local_path):
        return jsonify({'success': False, 'message': '本地目录不存在', 'files': [], 'categories': []})

    subdir = request.args.get('subdir', '').strip()
    target_path = os.path.join(local_path, subdir) if subdir else local_path
    target_path = os.path.normpath(target_path)

    if not target_path.startswith(os.path.normpath(local_path)):
        return jsonify({'success': False, 'message': '不允许访问上级目录'})

    if not os.path.isdir(target_path):
        return jsonify({'success': False, 'message': '目录不存在'})

    files = []
    categories = []
    try:
        for entry in os.scandir(target_path):
            if entry.name.startswith('~$'):
                continue
            stat = entry.stat()
            if entry.is_dir():
                categories.append({
                    'name': entry.name,
                    'path': os.path.relpath(entry.path, local_path).replace('\\', '/')
                })
            elif entry.is_file():
                _, ext = os.path.splitext(entry.name)
                files.append({
                    'name': entry.name,
                    'path': os.path.relpath(entry.path, local_path).replace('\\', '/'),
                    'size': stat.st_size,
                    'mtime': stat.st_mtime,
                    'ext': ext.lower()
                })
    except PermissionError:
        return jsonify({'success': False, 'message': '没有权限访问此目录'})

    categories.sort(key=lambda x: x['name'].lower())
    files.sort(key=lambda x: x['name'].lower())

    return jsonify({
        'success': True,
        'files': files,
        'categories': categories,
        'current_path': subdir.replace('\\', '/') if subdir else ''
    })


@kb_bp.route('/<kb_id>/local-categories', methods=['GET'])
@_require_kb_permission('view')
def list_local_categories(kb_id, _user_id=None):
    db = get_db()
    kb_row = db.execute("SELECT local_path, kb_type FROM knowledge_bases WHERE id = ?", (kb_id,)).fetchone()
    if not kb_row:
        return jsonify({'success': False, 'message': '知识库不存在'})
    if kb_row['kb_type'] != 'local':
        return jsonify({'success': False, 'message': '仅支持本地目录知识库'})

    local_path = kb_row['local_path']
    if not os.path.isdir(local_path):
        return jsonify({'success': False, 'categories': []})

    recursive = request.args.get('recursive', '0') == '1'

    if recursive:
        categories = _scan_categories_recursive(local_path, local_path)
        return jsonify({'success': True, 'categories': categories})

    subdir = request.args.get('subdir', '').strip()
    target_path = os.path.join(local_path, subdir) if subdir else local_path
    target_path = os.path.normpath(target_path)

    if not target_path.startswith(os.path.normpath(local_path)):
        return jsonify({'success': False, 'categories': []})

    categories = _scan_categories(target_path, local_path)
    return jsonify({'success': True, 'categories': categories})


def _scan_categories(target_path, base_path):
    categories = []
    try:
        for entry in os.scandir(target_path):
            if entry.name.startswith('~$') or not entry.is_dir():
                continue
            rel_path = os.path.relpath(entry.path, base_path).replace('\\', '/')
            child_categories = []
            try:
                for child in os.scandir(entry.path):
                    if not child.name.startswith('~$') and child.is_dir():
                        child_categories.append({
                            'name': child.name,
                            'path': os.path.relpath(child.path, base_path).replace('\\', '/')
                        })
            except PermissionError:
                pass
            child_categories.sort(key=lambda x: x['name'].lower())
            categories.append({
                'name': entry.name,
                'path': rel_path,
                'children': child_categories
            })
    except PermissionError:
        pass
    categories.sort(key=lambda x: x['name'].lower())
    return categories


def _scan_categories_recursive(target_path, base_path):
    categories = []
    try:
        for entry in os.scandir(target_path):
            if entry.name.startswith('~$') or not entry.is_dir():
                continue
            rel_path = os.path.relpath(entry.path, base_path).replace('\\', '/')
            children = _scan_categories_recursive(entry.path, base_path)
            categories.append({
                'name': entry.name,
                'path': rel_path,
                'children': children
            })
    except PermissionError:
        pass
    categories.sort(key=lambda x: x['name'].lower())
    return categories


@kb_bp.route('/<kb_id>/local-files/download', methods=['GET'])
@_require_kb_permission('view')
def download_local_file(kb_id, _user_id=None):
    db = get_db()
    kb_row = db.execute("SELECT local_path, kb_type FROM knowledge_bases WHERE id = ?", (kb_id,)).fetchone()
    if not kb_row or kb_row['kb_type'] != 'local':
        return jsonify({'success': False, 'message': '仅支持本地目录知识库'})

    local_path = kb_row['local_path']
    rel_path = request.args.get('path', '').strip()
    if not rel_path:
        return jsonify({'success': False, 'message': '未指定文件路径'})

    file_path = os.path.normpath(os.path.join(local_path, rel_path))
    if not file_path.startswith(os.path.normpath(local_path)):
        return jsonify({'success': False, 'message': '不允许访问上级目录'})

    if not os.path.isfile(file_path):
        return jsonify({'success': False, 'message': '文件不存在'})

    return send_file(file_path, as_attachment=True, download_name=os.path.basename(file_path))


@kb_bp.route('/<kb_id>/local-files/content', methods=['GET'])
@_require_kb_permission('view')
def get_local_file_content(kb_id, _user_id=None):
    db = get_db()
    kb_row = db.execute("SELECT local_path, kb_type FROM knowledge_bases WHERE id = ?", (kb_id,)).fetchone()
    if not kb_row or kb_row['kb_type'] != 'local':
        return jsonify({'success': False, 'message': '仅支持本地目录知识库'})

    local_path = kb_row['local_path']
    rel_path = request.args.get('path', '').strip()
    if not rel_path:
        return jsonify({'success': False, 'message': '未指定文件路径'})

    file_path = os.path.normpath(os.path.join(local_path, rel_path))
    if not file_path.startswith(os.path.normpath(local_path)):
        return jsonify({'success': False, 'message': '不允许访问上级目录'})

    if not os.path.isfile(file_path):
        return jsonify({'success': False, 'message': '文件不存在'})

    ext = os.path.splitext(file_path)[1].lower()

    if ext in ('.md', '.txt', '.html', '.htm', '.xml', '.json', '.csv', '.yaml', '.yml', '.py', '.js', '.css'):
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            return jsonify({'success': True, 'content': content, 'file_type': ext})
        except UnicodeDecodeError:
            return jsonify({'success': False, 'message': '无法以文本方式读取此文件'})

    if ext == '.docx':
        try:
            from docx import Document
            doc = Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs]
            return jsonify({'success': True, 'content': '\n'.join(paragraphs), 'file_type': ext})
        except Exception:
            return jsonify({'success': False, 'message': '无法读取 docx 内容'})

    if ext in ('.xlsx', '.xls'):
        try:
            import openpyxl
            wb = openpyxl.load_workbook(file_path, data_only=True)
            lines = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                lines.append(f'# {sheet_name}')
                for row in ws.iter_rows(values_only=True):
                    lines.append('\t'.join(str(c) if c is not None else '' for c in row))
            return jsonify({'success': True, 'content': '\n'.join(lines), 'file_type': ext})
        except Exception:
            return jsonify({'success': False, 'message': '无法读取 xlsx 内容'})

    return jsonify({'success': False, 'message': '不支持在线查看此文件类型的内容'})


@kb_bp.route('/<kb_id>/local-files/open', methods=['GET'])
@_require_kb_permission('view')
def open_local_file(kb_id, _user_id=None):
    db = get_db()
    kb_row = db.execute("SELECT local_path, kb_type FROM knowledge_bases WHERE id = ?", (kb_id,)).fetchone()
    if not kb_row or kb_row['kb_type'] != 'local':
        return jsonify({'success': False, 'message': '仅支持本地目录知识库'})

    local_path = kb_row['local_path']
    rel_path = request.args.get('path', '').strip()
    if not rel_path:
        return jsonify({'success': False, 'message': '未指定文件路径'})

    file_path = os.path.normpath(os.path.join(local_path, rel_path))
    if not file_path.startswith(os.path.normpath(local_path)):
        return jsonify({'success': False, 'message': '不允许访问上级目录'})

    if not os.path.isfile(file_path):
        return jsonify({'success': False, 'message': '文件不存在'})

    try:
        import platform
        import subprocess
        system = platform.system()
        if system == 'Windows':
            os.startfile(file_path)
        elif system == 'Darwin':
            subprocess.Popen(['open', file_path])
        else:
            subprocess.Popen(['xdg-open', file_path])
        return jsonify({'success': True, 'message': '已打开文件'})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})
